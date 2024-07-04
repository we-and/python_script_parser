#from PyPDF2 import PdfReader
import re
import os
import math
import sys
import io
import pandas as pd
from docx import Document
import csv
import platform
from pdfminer.high_level import extract_pages
from pdfminer.layout import LTTextContainer, LTChar
from pdfminer.layout import LTPage
from utils import get_log_file_path
import pdfplumber
import logging
import chardet
from utils import get_file_extension
from utils_parser import count_nonempty_lines_in_file ,is_characterline_TIMECODE_HYPHEN_UPPERCASECHARACTER_SPACE_SEMICOLON_DOUBLESPACE_TEXT_NEWLINE_TEXT,extract_character_name_TIMECODE_HYPHEN_UPPERCASECHARACTER_SPACE_SEMICOLON_DOUBLESPACE_TEXT_NEWLINE_TEXT,is_text_with_brackets_pattern,extract_text_after_brackets,extract_text_between_brackets,extract_TIMECODE_ARROW_TIMECODE_NEWLINE_CHARACTER_SEMICOLON_DIALOG_NEWLINE_DIALOG,matches_number_parenthesis_timecode,extract_scene_name,is_scene_line,isSeparatorEmptyLinesTimecode,isSeparatorNameParenthesisTimecode,isSeparatorParenthesisNameTimecode,matches_format_parenthesis_name_timecode,matches_format_parenthesis_name_timecode,matches_format_parenthesis_name_timecode,extract_matches,matches_scenestart_sceneno,count_lines_in_file,count_matches_NAME_NEWLINE_DIALOG_NEWLINE_NEWLINE,count_matches_TIMECODE_NEWLINE_CHARACTERINBRACKETS_DIALOG_NEWLINE_NEWLINE,count_matches_TIMECODE_HYPHEN_TIMECODE_NEWLINE_CHARACTER_SEMICOLON_NEWLINE_DIALOGNEWLINE,count_matches_LINE_NEWLINE_TIMECODE_ARROW_TIMECODE_NEWLINE_TEXT_ITAG,getCharacterSepType,extract_scene_name1,extract_scene_name2,extract_character_name,ensure_dialog_starts_with_uppercase,extract_charactername_CHARACTERUPPERCASE_DIALOG,extract_charactername_CHARACTERUPPERCASE_DIALOG_regex,extract_charactername_NAME_ATLEAST1TAB_TEXT,extract_charactername_NAME_ATLEAST8SPACES_TEXT,extract_charactername_NAME_SEMICOLON_DIALOG,extract_charactername_NAME_SEMICOLON_OPTSPACES_TAB_TEXT,extract_speech,extract_speech_NAME_SEMICOLON_OPTSPACES_TAB_TEXT,is_character_speaking,is_matching_character_speaking,getSceneSeparator,detectCharacterSeparator,matches_CHARACTERUPPERCASE_DIALOG,matches_charactername_NAME_SEMICOLON_DIALOG,countMethods
from  constants import cellLayoutModes, action_verbs,characterSeparators,countMethods,multilineCharacterSeparators
from utils_filters import filter_character_name
from utils_regex import is_celllayout_CHARACTERUPPERCASE_NEWLINE_DIALOG,extract_dialog_NUM_SPACE_CHARACTERUPPERCASE_SEMICOLON_DIALOG,extract_character_NUM_SPACE_CHARACTERUPPERCASE_SEMICOLON_DIALOG,is_celllayout_NUM_SPACE_CHARACTERUPPERCASE_SEMICOLON_DIALOG,is_timecode_arrow_timecode_format,is_NUM_TIMECODE_ARROW_TIMECODE,is_timecode_simple,is_TIMECODE_HYPHEN_TIMECODE,is_NUM_SEMICOLON_TIMECODE_SPACE_TIMECODE_SPACE_TIME
#script_path="scripts2/YOU CAN'T RUN FOREVER_SCRIPT_VO.txt"
#output_path="YOU CANT RUN FOREVER_SCRIPT_VOc/"
#script_name="YOU CANT RUN FOREVER_SCRIPT_VO"

#script_path="scripts2/ZERO11.txt"
#output_path="ZERO11c/"
#script_name="ZERO11b"



app_log_path=get_log_file_path()
logging.basicConfig(filename=app_log_path,level=logging.DEBUG,encoding='utf-8',filemode='w')
logging.debug("Script starting...")



def myprint1(s):
    logging.debug(s)
    #myprint1(s)



#################################################################
#ENCODING
def detect_file_encoding(file_path):
    myprint1("detect encoding of file "+str(file_path))
    with open(file_path, 'rb') as file:  # Open the file in binary mode
        raw_data = file.read(10000)  # Read the first 10000 bytes to guess the encoding
        result = chardet.detect(raw_data)
        return result

def test_encoding(script_path):
    encodings=['windows-1252', 'iso-8859-1', 'utf-16','ascii','utf-8']
    res={}
    firstok=None
    for enc in encodings:
        myprint1("  > Try encoding  : "+enc)
        try:
            with open(script_path, 'r', encoding=enc) as file:
                myprint1("  > Testing encoding  : "+enc)
                for line in file:
                    line = line.strip()  # Remove any leading/trailing whitespace
            res[enc]=True
            if firstok==None:
                firstok=enc
            return enc
        except UnicodeDecodeError:
            myprint1(f"  > Failed decoding with {enc}")
            res[enc]=False
    myprint1("Res="+str(res))
    if firstok!=None:
        return firstok
    return "?"    









#################################################################
# UTILS


def convert_csv_to_xlsx(csv_file_path, xlsx_file_path, script_name,encoding_used):
    myprint1("convert_csv_to_xlsx > 0")
    # Read the CSV file
    df = pd.read_csv(csv_file_path,header=None,encoding=encoding_used)

    # Write the DataFrame to an Excel file
    #myprint1("convert_csv_to_xlsx > Write to "+xlsx_file_path)

    header_rows = pd.DataFrame([
        [None, 'Header 1', None, 'Header Information Across Columns'],  # Merge cells will be across 1 & 4
        ['Role', 'Prises de parole', 'Caractères', 'Lignes']
    ])
    #myprint1("convert_csv_to_xlsx > 1")
    
    # Concatenate the header rows and the original data
    # The ignore_index=True option reindexes the new DataFrame
    df = pd.concat([header_rows, df], ignore_index=True)
    #myprint1("convert_csv_to_xlsx > 2")

    # Write the DataFrame to an Excel file
    with pd.ExcelWriter(xlsx_file_path, engine='openpyxl') as writer:
        #myprint1("convert_csv_to_xlsx > 3")

        df.to_excel(writer, index=False, sheet_name='Sheet1')
        #myprint1("convert_csv_to_xlsx > 4")

        # Load the workbook and sheet for modification
        workbook = writer.book
        sheet = workbook['Sheet1']
        #myprint1("convert_csv_to_xlsx > 5")

        # Merge cells in the first and second new rows
        # Assuming you want to merge from the first to the last column
        sheet.merge_cells('A1:D1')  # Modify range according to your number of columns
        sheet.merge_cells('A2:D2')  # Modify this as needed
        sheet['A1'] = script_name
        sheet['A2'] = "Length: "
        #myprint1("convert_csv_to_xlsx > 6")

    myprint1("convert_csv_to_xlsx > done")

#    df.to_excel(xlsx_file_path, index=False, engine='openpyxl')

def write_character_map_to_file(character_map, filename):
    myprint1(" > Write map to "+filename)
    """Writes the character to scene map to a specified file."""
    with open(filename, 'w', encoding='utf-8') as file:
        for character, scenes in character_map.items():
            file.write(f"{character}: {scenes}\n")

def find_first_uppercase_sequence(line):
    """Finds the first sequence of contiguous uppercase words in a line."""
    # Regex pattern to match the first sequence of contiguous uppercase words separated by spaces
    pattern = re.compile(r'\b([A-Z]+(?:\s[A-Z]+)*)\b')
    match = re.search(pattern, line)
    if match:
        return match.group(0)
    return None  # Return None if no uppercase sequence is found



def sort_dict_values(d):
    sorted_dict = {}
    for key, value_set in d.items():
        try:
            # Attempt to sort assuming all values are numeric strings
            sorted_list = sorted(value_set, key=int)
        except ValueError:
            # Handle the case where values are not all numeric
            myprint1(f"Non-numeric values found in the set for key '{key}'. Values: {value_set}")
            # Optionally sort only numeric values or handle differently
            numeric_values = [val for val in value_set if val.isdigit()]
            sorted_list = sorted(numeric_values, key=int)
        sorted_dict[key] = sorted_list
    return sorted_dict

def compute_length(line,method):
    if method=="ALL":
        return len(line)
    elif method=="ALL_NOSPACE":
        return len(line.replace(" ",""))
    else:
        return len(line)
    

def read_docx(file_path):
    # Load the document
    myprint1("OPEN DOCX 4")

    doc = Document(file_path)
    
    # Iterate through each paragraph in the document
    for para in doc.paragraphs:
        myprint1(para.text)
    return doc

def is_supported_extension(ext):
    ext=ext.lower()
    return ext==".txt" or ext==".docx" or ext==".doc" or ext==".rtf" or ext==".pdf" or ext==".xlsx"



def convert_docx_combined_continuity(file,table):
    myprint1("convert_docx_combined_continuity")
    current_character=""
    mode = "linear"
    titles=[]
    for row in table.rows[1:]:
        n=len(row.cells)
        if n>7:
            title=row.cells[n-1].text
            dur=row.cells[n-2].text
            is_title= len(dur)>0 and len(title)>0
            is_speech= len(dur)==0 and len(title)>0


            s=""
            if is_title:
                if "/" in title:
                    mode="split"
                    titles=title.split("/")
                else:
                    mode="linear"
                    title=clean_character_name(title)
                    current_character=title
            if is_speech:
                if mode=="split":
                    speeches=title.split("\n")

                    idx=0
  #                  myprint1("split mode speech"+title+str(len(speeches)))

                    for i in speeches:
                        current_character=titles[idx]
                        speech=speeches[idx].strip()
                        current_character=clean_character_name(current_character)
                        if speech.startswith("-"):
                            speech=speech[1:].strip()
                        s=current_character+"\t"+speech+"\n"
                        file.write(s)
                        idx=idx+1
                    mode="linear"

                elif mode=="linear":
                    current_character=clean_character_name(current_character)
                    s=current_character+"\t"+title+"\n"
                    #myprint1("linear mode add "+s)
                        
                    file.write(s)
def match_uppercase_semicolon(text):
    # Regex pattern to match text in uppercase ending with a semicolon
    pattern = r'^[A-Z\d\s]+:$'
    
    # Use re.match to check if the text matches the pattern
    if re.match(pattern, text):
        return True
    else:
        return False
def remove_semicolon(s):
    """
    Remove a semicolon at the end of a string if it exists.
    
    Args:
    s (str): The string from which to remove the semicolon if present at the end.

    Returns:
    str: The string without the trailing semicolon.
    """
    # Check if the string ends with a semicolon and remove it
    if s.endswith(':'):
        return s[:-1]  # Slice off the last character
    return s  # Return the original string if there's no semicolon at the end

def get_text_without_parentheses(input_string):
    pattern = r'\([^()]*\)'
    # Use re.sub() to replace the occurrences with an empty string
    result_string = re.sub(pattern, '', input_string)
    return result_string
def remove_text_in_brackets(text):
    """
    Remove all text within square brackets from the provided string.

    Args:
    text (str): The string from which to remove the text in brackets.

    Returns:
    str: The string with text in brackets removed.
    """
    # Regular expression to find text in square brackets
    pattern = r'\[.*?\]'
    # Use re.sub() to replace all occurrences of the pattern with an empty string
    cleaned_text = re.sub(pattern, '', text)
    return cleaned_text
def filter_speech(input):
    s=get_text_without_parentheses(input)
    s=remove_text_in_brackets(s)
    s=s.replace("â€™","'")
    s=s.replace("<i>","")
    s=s.replace("</i>","")
    s=s.replace("♪","").replace("Â§","").replace("§","")
    s=s.replace("â€¦ ",".")
    if s.startswith("- "):
        s=s.lstrip("- ")
    if s.startswith(": "):
        s=s.lstrip(": ")
    if s.startswith("<"):
        s=s.lstrip("<")
    if s.startswith(">"):
        s=s.lstrip(">")
    return s.strip()

def filter_speech_keepbrackets(input):
    s=get_text_without_parentheses(input)
    s=s.replace("â€™","'")
    s=s.replace("â€¦ ",".")
    return s

def convert_docx_dialogwithspeakerid(file,table,contentColIdx):
    
    myprint1("convert_docx_dialogwithspeakerid")
    idx=1
    row_idx=1
    current_character=""
    cumulated_speech=""
    ndisp=300
    
    for row in table.rows[1:]:
            show=True#row_idx<200 and row_idx>160
            if show:
                myprint1("---------------------------------------------")
                myprint1("Row "+str(row_idx))
                myprint1("Idx "+str(idx))
                myprint1("nCells "+str(len(row.cells)))
                myprint1("contentColIdx "+str(contentColIdx))
            row_idx=row_idx+1
            col=contentColIdx
            if len(row.cells) >2 and contentColIdx>len(row.cells)-1:
                col=len(row.cells)-1
            scenedesc=row.cells[col].text.strip()        
            if show:
                myprint1(scenedesc)
            parts=scenedesc.split("\n")
            upperpart=""
            nonupperpart=""
            for k in parts:

                if remove_text_in_brackets(k).isupper():
                    if show:
                        myprint1('+UP '+str(k))
                    upperpart=upperpart+" "+k
                else:
                    if show:
                        myprint1('+LO '+str(k))

                    nonupperpart=nonupperpart+" "+k
            part_idx=1
            upperpart=upperpart.strip()
            nonupperpart=nonupperpart.strip()
            if show:
                myprint1("upperpart:  "+upperpart)
                myprint1("lowerpart: " +nonupperpart)
            
            if is_character_name_valid(upperpart):
                if show:
                        myprint1("valid character candidate"+upperpart)
                upperpart=remove_text_in_brackets(upperpart).strip()
                if match_uppercase_semicolon(upperpart) or upperpart.isupper():
                        # if  :
                            if show:
                                    myprint1("IS CHARACTER YES")
                            if len(current_character)>0 and len(cumulated_speech)>0:
                                s=current_character+"\t"+cumulated_speech+"\n"
                                file.write(s)
                                myprint1(">>>>>>>>>>> "+s)
                                myprint1("reset cum")
                                cumulated_speech=""
                            current_character=remove_semicolon(upperpart) 
                            current_character=clean_character_name(current_character)
                            if show:
                                myprint1("new character "+current_character)
                else:
                    if show:
                        myprint1("Not a character")

                if len(nonupperpart)>0:
                    if show:
                        myprint1("has noupper"+nonupperpart)
                # if  idx<90:
                    #    myprint1("part "+str(part_idx)+":"+part)
                    part=filter_speech_keepbrackets(nonupperpart)
                    speech=part.strip() 
                    myprint1("Add to cummulated"+speech)
                    if len(current_character)>0:
                        cumulated_speech=cumulated_speech+" "+speech

                    if show:
                        myprint1("final"+nonupperpart)
            else:
                    if show:
                        myprint1("Not a valid character"+upperpart)
                    
    if len(cumulated_speech)>0:
            s=current_character+"\t"+cumulated_speech+"\n"
            file.write(s)
            idx=idx+1
            myprint1(">>>>>>>>>>> "+s)
def find_split_sep(cellcontent):
    nNewlines=cellcontent.count("\n")
    nHyphen=cellcontent.count("-")
    if nNewlines>0:
        return "\n"
    if nNewlines==0:
        if nHyphen>0:
            return "- "
    return "?"
    
def convert_docx_combined(file,table,bothCol):
    myprint1(f"convert_docx_combined {bothCol}")
    mode="LINEAR"
    for row in table.rows[1:]:
            myprint1(f"---")
            
            isglobalmergecell=True
            comp=row.cells[0].text
            myprint1("convert_docx_combined comp= "+ comp)
            for colidx in range(1,len(table.columns)):
                myprint1("convert_docx_combined comp with"+ row.cells[colidx].text)

                if comp!=row.cells[colidx].text:
                    isglobalmergecell=False
            myprint1("convert_docx_combined ismerged"+ str(isglobalmergecell))

            if isglobalmergecell:
                continue

            
            cellcontent=row.cells[bothCol].text.strip()      
            mode =detect_split_or_linear_mode_both(cellcontent)
            lines=cellcontent.split("\n")
            if mode=="LINEAR" :   
                current_character=lines[0].strip()
                speech=" ".join(lines[1:]).strip()

                current_character=filter_character_name(current_character)
                speech=filter_speech(speech)
                myprint1("convert_docx_combined Add "+ current_character + " "+ speech)
                res=current_character+"\t"+speech+"\n"  # New line after each row
                myprint1("convert_docx_combined Add "+ current_character + " "+ speech)
                file.write(res)   
            elif mode=="SPLIT":
                sep=find_split_sep(cellcontent)
                                
                speakers=extract_speakers(lines[0])
                myprint1("convert_docx_combined MODE SPLIT sep='"+str(sep)+"'")
                if sep=="\n":
                    dialoguespl=lines[1:]
                elif sep=='- ':
                    dialogue=" ".join(lines[1:])
                    dialoguespl=dialogue.split(sep)

                myprint1("convert_docx_combined characters"+str(speakers))
                filtered_array = [element for element in dialoguespl if element]
                myprint1("convert_docx_combined dialogue"+str(filtered_array))

                for index,k in enumerate(speakers):
                    myprint1("convert_docx_combined index"+str(index)+" k="+str(k))

                    current_character=filter_character_name(k)
                    speech= filter_speech(lines[index+1])
                    res=current_character+"\t"+speech+"\n"  # New line after each row
                    myprint1("convert_docx_combined Add ch="+ current_character + " sp="+ speech)
                    file.write(res)    
           

def convert_docx_scenedescription(file,table,sceneDescriptionIdx,titlesIdx):
    myprint1("convert_docx_scenedescription")
    current_character=""
    idx=1
    cumulated_speech=""
    row_idx=1

    for row in table.rows[1:]:
        if row_idx<35:
                myprint1("---------------------------------------------")
                myprint1("Row "+str(row_idx))
                myprint1("nCells "+str(len(row.cells)))
                scenedesc=row.cells[titlesIdx].text.strip()        
                myprint1("Row content"+str(scenedesc))
        row_idx=row_idx+1
    row_idx=1
    for row in table.rows[1:]:
            if idx<10:
                myprint1("---------------------------------------------")
                myprint1("nCells "+str(len(row.cells)))
                myprint1("Row "+str(row_idx))
            row_idx=row_idx+1
            scenedesc=row.cells[titlesIdx].text.strip()        
            if idx<100:
                myprint1(scenedesc)
            parts=scenedesc.split("\n")
            part_idx=1
            for part in parts:
                if idx<10:
                    myprint1("part "+str(part_idx)+":"+part)
                
                part=part.strip()
                part_idx=part_idx+1
                if len(part)>0:
#                    if idx<10:
 #                       myprint1("part "+str(part_idx)+":"+part)
                    part=filter_speech(part)
                    if match_uppercase_semicolon(part):
                        #flush cunulated speech
                        if len(cumulated_speech)>0:
                            s=current_character+"\t"+cumulated_speech+"\n"
                            file.write(s)
                            if idx<10:
                                myprint1(">>>>>>>>> "+current_character+"\t"+cumulated_speech)
                            idx=idx+1
                            cumulated_speech=""
                        current_character=remove_semicolon(part) 
                        if idx<10:
                            myprint1("new character "+current_character)
                    else:
                        if not part.isupper():                    
                            speech=part.strip()
                            current_character=clean_character_name(current_character)
                            if len(current_character)>0 and len(speech)>0:
                                cumulated_speech=cumulated_speech+" "+speech
                                #s=current_character+"\t"+speech+"\n"
                                #file.write(s)
    if len(cumulated_speech)>0:
        s=current_character+"\t"+cumulated_speech+"\n"
        file.write(s)
                                                    
def clean_character_name(title):
    title = title.replace("to herself","")
    title = title.replace(" said","")
    return title.strip().upper()

def convert_rtf_to_txt(file_path,currentOutputFolder,encoding):
    if platform.system() != 'Windows':
        return 
    pythoncom.CoInitialize()

    try:
        word = win32com.client.Dispatch("Word.Application")
        doc = word.Documents.Open(file_path)
        doc.Activate()

        # Extract text
        text = doc.Content.Text
         # Rename path with .docx
        new_file_abs = os.path.abspath(file_path)
        base=(os.path.basename(new_file_abs).replace(".rtf", ".converted.txt"))
        new_file_abs = os.path.join(currentOutputFolder, base)

        # Save and Close
#        word.ActiveDocument.SaveAs(new_file_abs, FileFormat=16)  # FileFormat=16 for .docx, not .doc
        doc.Close(False)
        txt_encoding="utf-8"
        with open(new_file_abs, 'w', encoding=txt_encoding) as file:
            file.write(text)

        return new_file_abs,txt_encoding

    finally:
        # Make sure to uninitialize COM
        pythoncom.CoUninitialize()
    return ""

def convert_docx_plain_text(file,doc):
    myprint1("convert_docx_plain_text")
    myprint1("convert_docx_plain_text "+str(len(doc.paragraphs))+" paragraphs")
    for para in doc.paragraphs:
        # Write the text of each paragraph to the file followed by a newline
        myprint1("convert_docx_plain_text write "+para.text)
        myprint1("                        isupper="+str(para.text.isupper()))
        myprint1("                        para= "+str(para))
        file.write(para.text + '\n')

def convert_docx_indented_plaintext(file,doc):
    myprint1("convert_docx_indented_plaintext")
    myprint1("convert_docx_indented_plaintext "+str(len(doc.paragraphs))+" paragraphs")
    for para in doc.paragraphs:
        if len(para.text)>0:
            indent = para.paragraph_format.left_indent
            myprint1("convert_docx_indented_plaintext text="+para.text)
            myprint1("convert_docx_indented_plaintext indent="+str(indent))
            
            if indent is not None:
                # Write the text of each paragraph to the file followed by a newline
                myprint1("                        write "+str(para.text.isupper()))
                myprint1("                        write "+str(para))
                file.write(para.text + '\n')

def extract_speakersNO(conversation, character_list):
    # Helper function to validate if a name is in the character list
    def is_valid_character(name):
        return name in character_list

    # Split the string by " TO " and process parts
    parts = conversation.split(" TO ")
    names = []
    seen = set()
    
    i = 0
    while i < len(parts) - 1:
        speaker = parts[i]
        listener = parts[i + 1]
        
        # Add speaker if it's a valid character and not already seen
        if is_valid_character(speaker) and speaker not in seen:
            names.append(speaker)
            seen.add(speaker)
        
        # Add listener if it's a valid character and not already seen
        if is_valid_character(listener) and listener not in seen:
            names.append(listener)
            seen.add(listener)
        
        i += 2
    
    # Handle the case where conversation does not end in a proper " TO " sequence
    if len(parts) % 2 == 1:
        last_part = parts[-1]
        if is_valid_character(last_part) and last_part not in seen:
            names.append(last_part)
    
    # Handle case where there's only one speaker or no valid conversations
    if not names and parts:
        if is_valid_character(parts[0]):
            return [parts[0]]
    
    return names
def normalize_spaces(text):
    return re.sub(r'\s+', ' ', text).strip()
def extract_speakers(conversation):
    # Split the string by " TO "
    parts = conversation.split(" TO ")
    nb_counts=conversation.count(" TO ")
    if nb_counts==0:
        return [conversation]
    if nb_counts==1:
        return [parts[0]]
    # Check if there are exactly two parts
    if nb_counts == 2:
        # If there is a reciprocal conversation, both parts should have the same speaker at the end
        first_speaker=parts[0]
        last_dest=parts[2]
        mid=parts[1]
        mid=normalize_spaces(mid)
        myprint1("case first= "+str(first_speaker)+" mid="+str(mid)+" last "+str(last_dest) )

        nb_words=mid.count(" ")
        midspl=mid.split(" ")
        myprint1("case mid= "+str(midspl) )
        if len(midspl)==2:
            second_speaker=midspl[1]
            return [first_speaker,second_speaker]
        elif len(midspl)==3:
            second_speaker=midspl[0]
            return [first_speaker,second_speaker]
        else:
            second_speaker=midspl[0]
            return [first_speaker,second_speaker]
    
    
            myprint1("error splitting characters within "+str(conversation) )
    
    
    
    
    else:
        # For more complex conversations, extract unique names
        seen = set()
        names = []
        for name in parts:
            if name not in seen:
                names.append(name)
                seen.add(name)
        return names




    
def extract_speakers1(conversation):
    # Split the string by " TO "
    parts = conversation.split(" TO ")

    # Extract names while preserving order
    seen = set()
    names = []
    for name in parts:
        if name not in seen:
            names.append(name)
            seen.add(name)
    
    # If there is only one unique name, return just that name
    if len(names) == 1:
        return names[0]
    
    # Otherwise, return the list of unique names
    return names
def detect_split_or_linear_mode_both(text):
    lines=text.split("\n")
    character=lines[0]
    character=filter_character_name(character)
    myprint1("character filtered"+character)
    
    myprint1("extract speakers for"+character)
    speakers=extract_speakers(character)
    myprint1("speakers="+str(speakers))

    nb_lines= len(speakers)
    if nb_lines>1:
        return "SPLIT"
    else:
        return "LINEAR"
def detect_split_or_linear_mode_separated(character):
    character=filter_character_name(character)
    myprint1("character filtered"+character)
    ##if "(O.S)" in character:
        #  character=character.replace("(O.S)","")
    #if "(O.S.)" in character:
        #   character=character.replace("(O.S.)","")

    myprint1("extract speakers for"+character)
    speakers=extract_speakers(character)

    #spl=character.split("\n")
    myprint1("speakers="+str(speakers))
    nb_lines= len(speakers)
    if nb_lines>1:
        return "SPLIT"
    else:
        return "LINEAR"
def detectCellLayoutMode(univtable,bothCol):
    best=None
    bestscore=0
    cells=univtable['cells']
    myprint1("detectCellLayoutMode col= "+ str(bothCol) )
    
    myprint1("detectCellLayoutMode ntests= "+ str(len(cellLayoutModes)) )
    for s in cellLayoutModes:
        myprint1("detectCellLayoutMode test "+ s )

        score=0
        for row_idx,row in enumerate(cells):
            t=row[bothCol]
            if s=="NUM_SPACE_CHARACTERUPPERCASE_SEMICOLON_DIALOG":
                ismatch=is_celllayout_NUM_SPACE_CHARACTERUPPERCASE_SEMICOLON_DIALOG(t)
                if ismatch:
                    score=score+1
            elif s=="CHARACTERUPPERCASE_NEWLINE_DIALOG":
                ismatch=is_celllayout_CHARACTERUPPERCASE_NEWLINE_DIALOG(t)
                if ismatch:
                    score=score+1
                
        myprint1("detectCellLayoutMode "+ s + " score="+ str(score)+"/"+str(len(cells)))

        if score>bestscore:
            best=s
            bestscore=score
    if bestscore<1:
        return "UNKNOWN_LAYOUT"
    myprint1("detectCellLayoutMode best="+ str(best) )
    return best
def convert_docx_separated(file,table,start,dialogueCol,characterIdCol):
    myprint1("convert_docx_characterid_and_dialogue")
    # Iterate through each row in the table
    current_character=""
    is_song_sung_by_character=False
    character_list=set()
    for row in table.rows[start:]:
        cell_texts = [cell.text for cell in row.cells]
            # Join all cell text into a single string
        row_text = ' | '.join(cell_texts)
        
        
        dialogue=row.cells[dialogueCol].text.strip()
        character=row.cells[characterIdCol].text.strip()

        mode="LINEAR"
        dialogue=dialogue.replace("\n","")
        myprint1("----------")
        myprint1("row"+row_text)
        myprint1("dialogue"+dialogue)
        myprint1("character"+character)
        mode =detect_split_or_linear_mode_separated(character)
        

        if mode=="LINEAR" :   
        
            if len(character)>0:
                current_character=character
            if len(dialogue)>0:  
                is_didascalie=dialogue.startswith("(")
                is_song=dialogue.startswith("♪") 
                is_song_sung_by_character=is_song and len(character)>0 #song starts but has character name on it
                if is_song:
                    dialogue=filter_speech(dialogue)             
                    force_current_character="__SONG"
                    if is_song_sung_by_character:
                         force_current_character=current_character
                    s=force_current_character+"\t"+dialogue+"\n"  # New line after each row
                    myprint1("Add "+ force_current_character + " "+ dialogue)
                    file.write(s)
                else:
                    is_song_sung_by_character=False
                if not is_didascalie and not is_song: 
                    dialogue=filter_speech(dialogue)             
                    if current_character=="":
                        current_character="__VOICEOVER"
                    s=current_character+"\t"+dialogue+"\n"  # New line after each row
                    myprint1("Add "+ current_character + " "+ dialogue)
                    file.write(s)
        elif mode=="SPLIT":
            speakers=extract_speakers(character)
            myprint1("MODE SPLIT")
            myprint1("characters"+str(speakers))
            dialoguespl=dialogue.split("- ")
            filtered_array = [element for element in dialoguespl if element]
            myprint1("dialogue"+str(filtered_array))
            for index,k in enumerate(speakers):
                myprint1("index"+str(index)+" k="+str(k))

                current_character=k
                speech=filter_speech( filtered_array[index])
                s=current_character+"\t"+speech+"\n"  # New line after each row
                myprint1("Add "+ current_character + " "+ dialogue)
                file.write(s)    


def convert_doc_to_docx(doc_path,currentOutputFolder):
    """Converts a .doc file to .docx"""
    
    if platform.system() != 'Windows':
        return 

    pythoncom.CoInitialize()

    try:
        word = win32com.client.Dispatch("Word.Application")
        doc = word.Documents.Open(doc_path)
        doc.Activate()

        # Rename path with .docx
        new_file_abs = os.path.abspath(doc_path)
        base=(os.path.basename(new_file_abs).replace(".doc", ".docx"))
        new_file_abs = os.path.join(currentOutputFolder, base)

        # Save and Close
        word.ActiveDocument.SaveAs(new_file_abs, FileFormat=16)  # FileFormat=16 for .docx, not .doc
        doc.Close(False)

        return new_file_abs

    finally:
        # Make sure to uninitialize COM
        pythoncom.CoUninitialize()
    return ""


def is_in_left_margin(block_left,  threshold):
    return block_left<threshold
def is_in_top_margin(block_left,  threshold):
    return block_left>threshold
def is_in_bottom_margin(block_left,  threshold):
    return block_left<threshold
def is_in_right_margin(block_left, min_left):
    return block_left>min_left

def is_centered(block_left, min_left, max_left, threshold):
    center = (min_left + max_left) / 2
    return abs(block_left - center) <= threshold
       
def get_pdf_text_elements(file_path,page_idx, page_start,page_end,progress_bar):
    return get_pdf_text_elements_alt(file_path,page_idx,page_start,page_end,progress_bar)
def group_words_into_lines(self, words, y_tolerance=3):
        #myprint7("group_words_into_lines")
        myprint1(f"group_words_into_lines N input ={len(words)}")
        
        # Dictionary to hold words grouped by their y position
        lines_map = {}
        
        for word in sorted(words, key=lambda w: (w['bbox'][1], w['bbox'][0])):
            text = word['text']
            x0, y0, x1, y1 = word['bbox']

            # Find the appropriate y position group
            added = False
            for y in list(lines_map.keys()):
                if abs(y - y0) <= y_tolerance:
                    lines_map[y].append(word)
                    added = True
                    break
            
            if not added:
                lines_map[y0] = [word]
        
  
        myprint1(f"group_words_into_lines map={lines_map.keys()}")
        for idx,k in enumerate(lines_map):
            myprint1(f"{idx} : {k.text}")

        # Merge the words in the map and output as a sorted list of lines

        lines = []
        for y in sorted(lines_map.keys()):
            sorted_line = sorted(lines_map[y], key=lambda w: w['bbox'][0])
            full_text = ' '.join([word['text'] for word in sorted_line])
            bbox = sorted_line[0]['bbox']
            lines.append({
                'text': full_text,
                'bbox': bbox
            })
     #   myprint7(f"group_words_into_lines N out ={len(lines)}")
      #  myprint7(f"group_words_into_lines res={lines}")
     #   myprint7(f"group_words_into_lines map={lines_map}")
        return lines

def get_pdf_text_elements_alt(file_path,page_idx, page_start,page_end,progress_bar):
        myprint1("---------- get_pdf_page_blocks -----------------")
        mode="FLOW"
        text_elements=[]
        page_split_elements=[]
        minboxleft=100000
        if page_idx<0:
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    progress_bar['value'] = page_num
                    progress_bar.update_idletasks()
                    current_page_elements=[]
                    myprint1(f"\nPage {page_num}")
                    if page_num<page_start:
                            continue
                    if page_num>page_end:
                            break
                                    # Get page dimensions
                    page_width = page.width
                    page_height = page.height
                    table = page.extract_table()
                    if table:
                    
                        page_layout = page.extract_words()
                        mode="TABLE"
                        myprint1("ELEMENTS: "+str(len(page_layout)))
                        for element in page_layout:
                                
                                # Loop over the text blocks within the text container
                                text_block = element['text']
                                bbox =   [element['x0'], page_height - element['top'], element['x1'],page_height -  element['bottom']]                  
                                
                                tt=text_block.strip().replace("\n","")
                                stripped_text=tt
                                if len(stripped_text)>0:
                                    # Print the text block and its bounding box
                                    myprint1(f"{bbox} {stripped_text}")
                                    if bbox[0] < minboxleft:
                                        minboxleft=bbox[0]
                                    text_elements.append({
                                        'text':stripped_text,
                                        'bbox':bbox,
                                    })
                                current_page_elements.append({
                                        'text':stripped_text,
                                        'bbox':bbox,
                                    })
                        current_page_elements = sorted(current_page_elements, key=lambda x: x['bbox'][1], reverse=True)
                        page_split_elements.append(current_page_elements)

    
                    else:
                        page_elements=[]
                        
                        # Text
                        text_settings = {
                        "x_tolerance": 3,  # Adjust this value to fine-tune line detection
                        "y_tolerance": 3,  # Adjust this value to fine-tune line detection
                        }
                        page_layout = page.extract_text_lines(text_settings)
                        for element in page_layout:
                                    myprint1(f"  Word: '{element['text']}', bbox: {element['x0']}, {element['top']}, {element['x1']}, {element['bottom']}")
                                    tt=element['text'].strip().replace("\n","")
                                    bbox =   [element['x0'],page_height -  element['top'], element['x1'], page_height - element['bottom']]                  
                                    stripped_text=tt
                                    if len(stripped_text)>0:
                                        # Print the text block and its bounding box
                                        myprint1(f"{bbox} {tt}")
                                        if bbox[0] < minboxleft:
                                            minboxleft=bbox[0]
                                        page_elements.append({
                                            'text':stripped_text,
                                            'bbox':bbox,
                                        })
                        myprint1("NON EMPTY ELEMENTS: "+str(len(page_elements)))
                    
                        page_elements = sorted(page_elements, key=lambda x: x['bbox'][1], reverse=True)
                        for k in page_elements:
                            text_elements.append(k)

        else:        
            with pdfplumber.open(file_path) as pdf:
                page = pdf.pages[page_idx]
                page_layout=page.extract_words()
                
                # Get page dimensions
                page_width = page.width
                page_height = page.height

                table = page.extract_table()

                # Convert the extracted table to a DataFrame
                if table:
                    mode="TABLE"
                    page_layout = page.extract_words()

                    myprint1("ELEMENTS: "+str(len(page_layout)))
                    for element in page_layout:
                            
                            # Loop over the text blocks within the text container
                            text_block = element['text']
                            bbox =   [element['x0'], page_height - element['top'], element['x1'],page_height -  element['bottom']]                  
                            
                            tt=text_block.strip().replace("\n","")
                            stripped_text=tt
                            if len(stripped_text)>0:
                                # Print the text block and its bounding box
                                myprint1(f"{bbox} {stripped_text}")
                                if bbox[0] < minboxleft:
                                    minboxleft=bbox[0]
                                text_elements.append({
                                    'text':stripped_text,
                                    'bbox':bbox,
                                })
                    myprint1("NON EMPTY ELEMENTS: "+str(len(text_elements)))
                    if False:
                        # Loop through each row in the table
                        print("TABLE")
                        for row_idx, row in enumerate(table):
                            for col_idx, cell in enumerate(row):
                                if cell:
                                    # Extract the bounding box of each cell
                                    cell_texts = page.extract_text_lines()
                                    for cell_text in cell_texts:
                                        print(f"CELL test {row_idx} {col_idx}")

                                        # Check if the text is in the current cell's row and column
                                        if cell_text['top'] < page.bbox[3] - (row_idx + 1) * (page.bbox[3] / len(table)) and cell_text['bottom'] > page.bbox[3] - row_idx * (page.bbox[3] / len(table)):
                                            print(f"CELL test {row_idx} {col_idx} IN CELL")

                                            stripped_text = cell_text['text'].strip()
                                            bounding_box_of_the_block = cell_text['bbox']
                                            text_elements.append({
                                                'text': stripped_text,
                                                'bbox': bounding_box_of_the_block,
                                                'row': row_idx,
                                                'column': col_idx
                                            })
                                            myprint1(f"table {row_idx} {col_idx} {stripped_text}")
                else:
                    
                    text_settings = {
                        "x_tolerance": 3,  # Adjust this value to fine-tune line detection
                        "y_tolerance": 3,  # Adjust this value to fine-tune line detection
                    }
                    page_layout = page.extract_text_lines(text_settings)

                    myprint1("ELEMENTS: "+str(len(page_layout)))
                    for element in page_layout:
                            
                            # Loop over the text blocks within the text container
                            text_block = element['text']
                            bbox =   [element['x0'], page_height - element['top'], element['x1'],page_height -  element['bottom']]                  
                            
                            tt=text_block.strip().replace("\n","")
                            stripped_text=tt
                            if len(stripped_text)>0:
                                # Print the text block and its bounding box
                                myprint1(f"{bbox} {stripped_text}")
                                if bbox[0] < minboxleft:
                                    minboxleft=bbox[0]
                                text_elements.append({
                                    'text':stripped_text,
                                    'bbox':bbox,
                                })
                    myprint1("NON EMPTY ELEMENTS: "+str(len(text_elements)))
                    
                    text_elements = sorted(text_elements, key=lambda x: x['bbox'][1], reverse=True)

        myprint1("return get_pdf_page_blocks "+str(mode))
        return mode,text_elements,page_split_elements

def get_pdf_text_elements_pdfminer(file_path,page_idx, page_start,page_end):
        myprint1("---------- get_pdf_page_blocks -----------------")
        text_elements=[]
        minboxleft=100000

        if page_idx<0:
            for page_layout in extract_pages(file_path):
                    myprint1("----------------------------")
                    myprint1(f"Page number: {page_layout.pageid}")
                    if page_layout.pageid<page_start:
                        continue
                    if page_layout.pageid>page_end:
                        break

                    # Get the page bounding box coordinates and dimensions
                    if isinstance(page_layout, LTPage):
                        page_bbox = page_layout.bbox
                        page_x0, page_y0, page_x1, page_y1 = page_bbox
#                        page_width = page_x1 - page_x0
 #                       page_height = page_y1 - page_y0
                    else:
                        # If page bounding box is not available, skip processing the page
                        continue
                    # Loop over the elements in the page
                    myprint1(str(page_layout))
                    myprint1("ELEMENTS: "+str(len(page_layout)))
                    page_elements=[]
                    for element in page_layout:
                        
                        # Check if the element is a text container
                        if isinstance(element, LTTextContainer):
                            # Loop over the text blocks within the text container
                            text_block = element.get_text()
                            bbox = element.bbox  # (x0, y0, x1, y1)                        
                            is_inside=    bbox[0] >= page_x0 and bbox[1] >= page_y0 and                bbox[2] <= page_x1 and bbox[3] <= page_y1
                            if is_inside:
                                tt=text_block.strip().replace("\n","")
                                stripped_text=tt
                                if len(stripped_text)>0:
                                # Print the text block and its bounding box
                                    myprint1(f"{bbox} {tt}")
                                    if bbox[0] < minboxleft:
                                        minboxleft=bbox[0]
                                    page_elements.append({
                                        'text':text_block.strip(),
                                        'bbox':bbox,
                                    })
                    myprint1("NON EMPTY ELEMENTS: "+str(len(page_elements)))
                
                    page_elements = sorted(page_elements, key=lambda x: x['bbox'][1], reverse=True)
                    for k in page_elements:
                        text_elements.append(k)
        else:        
            for page_layout in extract_pages(file_path, page_numbers=[page_idx ]):
                myprint1("----------------------------")
                myprint1(f"Page number: {page_layout.pageid}")
                myprint1("-- PAGE "+str(page_idx))
            
                
                # Get the page bounding box coordinates and dimensions
                if isinstance(page_layout, LTPage):
                    page_bbox = page_layout.bbox
                    page_x0, page_y0, page_x1, page_y1 = page_bbox
                    page_width = page_x1 - page_x0
                    page_height = page_y1 - page_y0
                else:
                    # If page bounding box is not available, skip processing the page
                    continue
                # Loop over the elements in the page
                myprint1(str(page_layout))
                myprint1("ELEMENTS: "+str(len(page_layout)))
                for element in page_layout:
                    
                    # Check if the element is a text container
                    if isinstance(element, LTTextContainer):
                        # Loop over the text blocks within the text container
                        text_block = element.get_text()
                        bbox = element.bbox  # (x0, y0, x1, y1)                        
                        is_inside=    bbox[0] >= page_x0 and bbox[1] >= page_y0 and                bbox[2] <= page_x1 and bbox[3] <= page_y1
                        if is_inside:
                            tt=text_block.strip().replace("\n","")
                            stripped_text=tt
                            if len(stripped_text)>0:
                                # Print the text block and its bounding box
                                myprint1(f"{bbox} {stripped_text}")
                                if bbox[0] < minboxleft:
                                    minboxleft=bbox[0]
                                text_elements.append({
                                    'text':stripped_text,
                                    'bbox':bbox,
                                })
                myprint1("NON EMPTY ELEMENTS: "+str(len(text_elements)))
                
                text_elements = sorted(text_elements, key=lambda x: x['bbox'][1], reverse=True)

    
        return text_elements
def get_pdf_page_blocks(file_path,page_idx):
    myprint1("---------- get_pdf_page_blocks -----------------")
    text_elements=[]
    minboxleft=100000
    for page_layout in extract_pages(file_path, page_numbers=[page_idx ]):
                myprint1("----------------------------")
                myprint1(f"Page number: {page_layout.pageid}")
                myprint1("-- PAGE "+str(page_idx))
            
                # Get the page bounding box coordinates and dimensions
                if isinstance(page_layout, LTPage):
                    page_bbox = page_layout.bbox
                    page_x0, page_y0, page_x1, page_y1 = page_bbox
                    page_width = page_x1 - page_x0
                    page_height = page_y1 - page_y0
                else:
                    # If page bounding box is not available, skip processing the page
                    continue
                # Loop over the elements in the page
                myprint1(str(page_layout))

                for element in page_layout:
                    
                    # Check if the element is a text container
                    if isinstance(element, LTTextContainer):
                        # Loop over the text blocks within the text container
                        text_block = element.get_text()
                        #for text_line in element:
                        #   text_block = ""
                        #  text_position = []

                        bbox = element.bbox  # (x0, y0, x1, y1)
                        
                        is_inside=    bbox[0] >= page_x0 and bbox[1] >= page_y0 and                bbox[2] <= page_x1 and bbox[3] <= page_y1
                        #myprint1(                bbox[0] >= page_x0) 
                        #myprint1(         bbox[1] >= page_y0 )
                        #myprint1(         bbox[2] <= page_x1 )
                        #myprint1(          bbox[3] <= page_y1)
         
                        if is_inside:
                            # Print the text block and its bounding box
                            myprint1(f"{bbox} {text_block.strip()}")
                            if bbox[0] < minboxleft:
                                minboxleft=bbox[0]
                            text_elements.append({
                                'text':text_block.strip(),
                                'bbox':bbox,
                            })
    xvalues={}
    for k,el in enumerate(text_elements):
        left=el['bbox'][0]
        if left in xvalues:
            xvalues[left]=xvalues[left]+1
        else:
            xvalues[left]=1

    myprint1("XVALUES"+str(xvalues))

    xval=list(xvalues.keys())
    # Find the minimum and maximum left positions
    min_left = min(xval)
    max_left = max(xval)
    myprint1("MIN_LEFT"+str(min_left))
    myprint1("MAX_LEFT"+str(max_left))


    left_margin=110
    top_margin=740
    right_margin=500
    # Categorize blocks as left-aligned or centered
    top_aligned_blocks = []
    left_aligned_blocks = []
    right_aligned_blocks = []
    centered_blocks = []

    myprint1("-------------------------------")
    myprint1("TEST")
    for k,el in enumerate(text_elements):
        left_pos=el['bbox'][0]
        bottom_pos=el['bbox'][1]
        myprint1("test ["+str(left_pos) +"]"+str(el['text']))
        if is_in_top_margin(bottom_pos, top_margin):
            myprint1("test left")
            top_aligned_blocks.append(el)
        elif is_in_left_margin(left_pos, left_margin):
            myprint1("test left")
            left_aligned_blocks.append(el)
        elif is_in_right_margin(left_pos,right_margin):
            myprint1("test right")
            right_aligned_blocks.append(el)
        else:#if is_centered(left_pos, min_left, max_left, threshold):
            myprint1("test centre")
            centered_blocks.append(el)
    return {
        'left':left_aligned_blocks,
        'right':right_aligned_blocks,
        'top':top_aligned_blocks,
        'center':centered_blocks
    }
def split_elements(text_elements,left_margin,top_margin,right_margin,bottom_margin):
    top_blocks = []
    left_blocks = []
    right_blocks = []
    bottom_blocks = []
    centered_blocks = []
    myprint1("SPLIT")
    if text_elements != None:
        myprint1("NB EL"+str(len(text_elements)))
        myprint1(f" > elements {len(text_elements)}")
    myprint1(f" > margins left={left_margin} top={top_margin} right={right_margin} bottom={bottom_margin}")
    #myprint1("TEST")
    if text_elements!=None:
        for k,el in enumerate(text_elements):
            if len(el['text'])>0:
                left_pos=el['bbox'][0]
                bottom_pos=el['bbox'][1]
            #   myprint1("test ["+str(left_pos) +"]"+str(el['text']))
                if is_in_top_margin(bottom_pos, top_margin):
                    #myprint1(f"    - is_top bottompos={bottom_pos} topmargin={top_margin}")
                    top_blocks.append(el)
                elif is_in_left_margin(left_pos, left_margin):
            #     myprint1("test left")
                    left_blocks.append(el)
                elif is_in_right_margin(left_pos,right_margin):
                #    myprint1("test right")
                    right_blocks.append(el)
                elif is_in_bottom_margin(bottom_pos,bottom_margin):
                #    myprint1("test right")
                    bottom_blocks.append(el)
                else:#if is_centered(left_pos, min_left, max_left, threshold):
                #   myprint1("test centre")
                    centered_blocks.append(el)
        myprint1(f" > groups left={len(left_blocks)} top={len(top_blocks)} right={len(right_blocks)} center={len(centered_blocks)} bottom={len(bottom_blocks)}")

    return {
        'bottom':bottom_blocks,
        'left':left_blocks,
        'right':right_blocks,
        'top':top_blocks,
        'center':centered_blocks
    }

def convert_pdf_to_txt(file_path,absCurrentOutputFolder,encoding):
    myprint1("convert_pdf_to_txt")
    myprint1("currentOutputFolder             :"+absCurrentOutputFolder)
    myprint1("Input              :"+file_path)
    converted_file_path=""
    if ".pdf" in file_path.lower() :
        converted_file_path=os.path.join(absCurrentOutputFolder, os.path.basename(file_path).lower().replace(".pdf",".converted.txt"))
    if ".pdf" in file_path.lower() :
        converted_raw_file_path=os.path.join(absCurrentOutputFolder, os.path.basename(file_path).lower().replace(".pdf",".raw.txt"))
    minboxleft=100000
    npagesmax=10000
    firstpage=3
    page_idx=0
    
    with open(converted_raw_file_path, 'w', encoding='utf-8') as fileraw:              
        
        text_elements=[]
        for page_layout in extract_pages(file_path):
                page_idx=page_idx+1
                if page_idx<firstpage:
                    continue
                if page_idx>npagesmax:
                    break
                myprint1("----------------------------")
                myprint1(f"Page number: {page_layout.pageid}")
                myprint1("-- PAGE "+str(page_idx))
            
                # Get the page bounding box coordinates and dimensions
                if isinstance(page_layout, LTPage):
                    page_bbox = page_layout.bbox
                    page_x0, page_y0, page_x1, page_y1 = page_bbox
                    page_width = page_x1 - page_x0
                    page_height = page_y1 - page_y0
                else:
                    # If page bounding box is not available, skip processing the page
                    continue
                # Loop over the elements in the page
                for element in page_layout:
                    # Check if the element is a text container
                    if isinstance(element, LTTextContainer):
                        # Loop over the text blocks within the text container
                        text_block = element.get_text()
                        #for text_line in element:
                        #   text_block = ""
                        #  text_position = []

                        bbox = element.bbox  # (x0, y0, x1, y1)
                        
                        is_inside=    bbox[0] >= page_x0 and bbox[1] >= page_y0 and                bbox[2] <= page_x1 and bbox[3] <= page_y1
                        #myprint1(                bbox[0] >= page_x0) 
                        #myprint1(         bbox[1] >= page_y0 )
                        #myprint1(         bbox[2] <= page_x1 )
                        #myprint1(          bbox[3] <= page_y1)
         
                        if is_inside:
                            # Print the text block and its bounding box
                            if len(text_block.strip())>0:
                                myprint1(f"{bbox} {text_block.strip()}")
                            if bbox[0] < minboxleft:
                                minboxleft=bbox[0]
                            text_elements.append({
                                'text':text_block.strip(),
                                'bbox':bbox,
                            })
    xvalues={}
    for k,el in enumerate(text_elements):
        left=el['bbox'][0]
        if left in xvalues:
            xvalues[left]=xvalues[left]+1
        else:
            xvalues[left]=1

    myprint1("XVALUES"+str(xvalues))

    xval=list(xvalues.keys())
    # Find the minimum and maximum left positions
    min_left = min(xval)
    max_left = max(xval)
    myprint1("MIN_LEFT"+str(min_left))
    myprint1("MAX_LEFT"+str(max_left))

    threshold = 50
    # Categorize blocks as left-aligned or centered
    top_aligned_blocks = []
    left_aligned_blocks = []
    right_aligned_blocks = []
    centered_blocks = []

    myprint1("-------------------------------")
    myprint1("TEST")
    for k,el in enumerate(text_elements):
        left_pos=el['bbox'][0]
        bottom_pos=el['bbox'][1]
        #myprint1("test ["+str(left_pos) +"]"+str(el['text']))
        if is_in_top_margin(bottom_pos,  threshold):
         #   myprint1("test left")
            top_aligned_blocks.append(el)
        elif is_in_left_margin(left_pos, min_left, threshold):
          #  myprint1("test left")
            left_aligned_blocks.append(el)
        elif is_in_right_margin(left_pos, max_left, threshold):
           # myprint1("test right")
            right_aligned_blocks.append(el)
        else:#if is_centered(left_pos, min_left, max_left, threshold):
            #myprint1("test centre")
            centered_blocks.append(el)

    
    myprint1("-------------------------------")
    myprint1("LEFT")
    for k,el in enumerate(left_aligned_blocks):
        left_pos=el['bbox'][0]
        text=el['text']
        myprint1(str(text)+" "+str(left_pos))
    myprint1("-------------------------------")
    myprint1("RIGHT")
    for k,el in enumerate(right_aligned_blocks):
        left_pos=el['bbox'][0]
        text=el['text']
        myprint1(str(text)+" "+str(left_pos))
    myprint1("-------------------------------")
    myprint1("CENTER")
    for k,el in enumerate(centered_blocks):
        left_pos=el['bbox'][0]
        text=el['text']
        myprint1(str(text)+" "+str(left_pos))

    myprint1("-------------------------------")
    myprint1("OUTPUT")
    current_character=None
    is_after_character=False
    with open(converted_file_path, 'w', encoding='utf-8') as file:
        for k,el in enumerate(centered_blocks):
            left_pos=el['bbox'][0]
            text=el['text']
            parts=text.split("\n")
            Nparts=len(parts)
            if Nparts==1:    
                if text.isupper() and not is_after_character and not text.startswith("¡"):
                    myprint1("OUT  ["+str(left_pos)+"]   "+text+"    -->  CHAR   ")
                    current_character=filter_character_name(text)
                    is_after_character=True
                else:
                    myprint1("OUT  ["+str(left_pos)+"]   "+text+"    -->  DIALOG   ")
                    dialog=filter_speech(text)
                    is_after_character=False
                    if current_character!=None:
                        s=current_character+"\t"+dialog+"\n"  # New line after each row
                        file.write(s)
            else:
                for part in parts:
                    if part.isupper() and not is_after_character and not text.startswith("¡"):
                        myprint1("OUT  ["+str(left_pos)+"]   "+part+"    -->  CHAR   ")
                        current_character=filter_character_name(part)
                        is_after_character=True
                    else:
                        myprint1("OUT  ["+str(left_pos)+"]   "+part+"    -->  DIALOG   ")
                        dialog=filter_speech(part)
                        is_after_character=False
                        if current_character!=None:
                            s=current_character+"\t"+dialog+"\n"  # New line after each row
                            file.write(s)


    myprint1("-------------------------------")
    myprint1("FINISHED")
    myprint1("Converted")                       
    myprint1(converted_file_path)
    return converted_file_path,encoding

def run_convert_pdf_to_txt(file_path,absCurrentOutputFolder,centered_blocks,encoding):
    converted_file_path=""
    if ".pdf" in file_path.lower() :
        converted_file_path=os.path.join(absCurrentOutputFolder, os.path.basename(file_path).lower().replace(".pdf",".converted.txt"))
    if ".pdf" in file_path.lower() :
        converted_raw_file_path=os.path.join(absCurrentOutputFolder, os.path.basename(file_path).lower().replace(".pdf",".raw.txt"))

    current_character=None
    is_after_character=False
    with open(converted_file_path, 'w', encoding='utf-8') as file:
        for k,el in enumerate(centered_blocks):
            left_pos=el['bbox'][0]
            text=el['text']
            parts=text.split("\n")
            Nparts=len(parts)
            if Nparts==1:    
                if text.isupper() and not is_after_character and not text.startswith("¡"):
                    myprint1("OUT1  ["+str(left_pos)+"]   "+text+"        --> CHAR   ")
                    current_character=filter_character_name(text)
                    is_after_character=True
                else:
                    myprint1("OUT1  ["+str(left_pos)+"]   "+text+"        --> DIALOG   ")
                    dialog=filter_speech(text)
                    is_after_character=False
                    if current_character!=None:
                        s=current_character+"\t"+dialog+"\n"  # New line after each row
                        file.write(s)
            else:
                for part in parts:
                    if part.isupper() and not is_after_character and not text.startswith("¡"):
                        myprint1("OUT1  ["+str(left_pos)+"]   "+part+"        --> CHAR   ")
                        current_character=filter_character_name(part)
                        is_after_character=True
                    else:
                        myprint1("OUT1  ["+str(left_pos)+"]   "+part+"        --> DIALOG   ")
                        dialog=filter_speech(part)
                        is_after_character=False
                        if current_character!=None:
                            s=current_character+"\t"+dialog+"\n"  # New line after each row
                            file.write(s)


    myprint1("-------------------------------")
    myprint1("FINISHED")
    myprint1("Converted")                       
    myprint1(converted_file_path)
    return converted_file_path,encoding
            
def convert_pdf_to_txt_pdfplum(file_path,absCurrentOutputFolder,encoding):
    myprint1("convert_pdf_to_txt")
    myprint1("currentOutputFolder             :"+absCurrentOutputFolder)
    myprint1("Input              :"+file_path)
    converted_file_path=""
    if ".pdf" in file_path.lower() :
        converted_file_path=os.path.join(absCurrentOutputFolder, os.path.basename(file_path).lower().replace(".pdf",".converted.txt"))
    if ".pdf" in file_path.lower() :
        converted_raw_file_path=os.path.join(absCurrentOutputFolder, os.path.basename(file_path).lower().replace(".pdf",".raw.txt"))
    with open(converted_file_path, 'w', encoding='utf-8') as file:
        with open(converted_raw_file_path, 'w', encoding='utf-8') as fileraw:
                pdf = PdfReader(file_path)
                page_idx=1  
                for page_num in range(len(pdf.pages)):
                    page = pdf.pages[page_num]
                    page_idx=page_idx+1
                    if page_idx<40000:
                        # Extract tables from the page
                        myprint1("########################################################")
                        myprint1("convert_pdf_to_txt page"+str(page_idx))
                        #crop_coords = [0.68,0.20,0.95,0.90]
                        #my_width = page.width
                        #my_height = page.height
                        #my_bbox = (crop_coords[0]*float(my_width), crop_coords[1]*float(my_height), crop_coords[2]*float(my_width), crop_coords[3]*float(my_height))
                        #page = page.crop(bbox=my_bbox)
                        myprint1("extract")
                        extr=page.extract_text(TextAlignment.NONE).split("\n\n")
                        myprint1(extr)
                        text = str(extr)
                        myprint1(text)
                        myprint1("write")
                        fileraw.write(text)
                        myprint1("split")
                        text_blocks = extr.split('\n')

                        for block_num, block in enumerate(text_blocks, start=1):
                            # Skip empty blocks
                            if block.strip() == "":
                                continue
                            
                            # Print the page number, block number, and block text
                            myprint1(f"Page: {page_num + 1}")
                            myprint1(f"Block: {block_num}")
                            myprint1(f"Text: {block.strip()}")
                            myprint1("---")
                            block_position = page.get_text_block_position(block)
                            if block_position is not None:
                                # Extract the coordinates of the block position
                                x0, y0, x1, y1 = block_position

                                myprint1(f"Position: ({x0}, {y0}) - ({x1}, {y1})")
                            else:

                                myprint1(f"Position: none")

                        textlines=text.split("\n")
                        current_character=""
                        current_characters_split=[]
                        speech=""
                        mode="linear"
                        for line in textlines:
                            if line.isupper():
                                if " THEN " in line or "," in line:
                                    mode="split"
                                    if " THEN " in line:
                                        current_characters_split=line.split(" THEN ")   
                                    elif "," in line: 
                                        current_characters_split=line.split(",")    
                                    
                                    current_character=line
                                else:
                                    current_character=line
                                    mode="linear"
                            else:
                                speech=line
                                if mode=="split":
                                    speeches=speech.split("\n")
                                    charidx=0
                                    for sp in speeches:
                                        sp=sp.replace("- ","")
                                        current_character=current_characters_split[charidx]
                                        s=current_character+"\t"+sp+"\n"  # New line after each row
                                        myprint1("Add "+ current_character + " "+ speech)
                                        file.write(s)
                                        charidx=charidx+1

                                if mode=="linear":
                                    if current_character=="":
                                        current_character="__VOICEOVER"
                                    s=current_character+"\t"+speech+"\n"  # New line after each row
                                    myprint1("Add "+ current_character + " "+ speech)
                                    file.write(s)

    myprint1("Converted")                       
    myprint1(converted_file_path)
    return converted_file_path,encoding

def convert_pdf_to_txt_pdfplumber(file_path,absCurrentOutputFolder,encoding):
    myprint1("convert_pdf_to_txt")
    myprint1("currentOutputFolder             :"+absCurrentOutputFolder)
    myprint1("Input              :"+file_path)
    converted_file_path=""
    if ".pdf" in file_path.lower() :
        converted_file_path=os.path.join(absCurrentOutputFolder, os.path.basename(file_path).lower().replace(".pdf",".converted.txt"))
    if ".pdf" in file_path.lower() :
        converted_raw_file_path=os.path.join(absCurrentOutputFolder, os.path.basename(file_path).lower().replace(".pdf",".raw.txt"))
    with open(converted_file_path, 'w', encoding='utf-8') as file:
        with open(converted_raw_file_path, 'w', encoding='utf-8') as fileraw:
            with pdfplumber.open(file_path) as pdf:
                page_idx=1  
                for page in pdf.pages[1:]:

                    page_idx=page_idx+1
                    if page_idx<40000:
                        # Extract tables from the page
                        myprint1("########################################################")
                        myprint1("convert_pdf_to_txt page"+str(page_idx))
                        #crop_coords = [0.68,0.20,0.95,0.90]
                        #my_width = page.width
                        #my_height = page.height
                        #my_bbox = (crop_coords[0]*float(my_width), crop_coords[1]*float(my_height), crop_coords[2]*float(my_width), crop_coords[3]*float(my_height))
                        #page = page.crop(bbox=my_bbox)
                        myprint1("extract")
                        extr=page.extract_text()
                        myprint1(extr)
                        text = str(extr)
                        myprint1(text)
                        myprint1("write")
                        fileraw.write(text)
                        myprint1("split")

                        text_blocks = extr.split('\n')
                        myprint1("Blocks"+str(len(text_blocks)))
                        # Iterate over each text block
                        for block in text_blocks:
                            myprint1("block"+str(block))
                            # Find the position and dimensions of the text block
                            bbox = page.bbox_for_text(block)
                            myprint1("bbox")
                            if bbox:
                                x0, top, x1, bottom = bbox
                                width = x1 - x0
                                
                                # Print the text block information
                                myprint1(f"Page: {page_idx}")
                                myprint1(f"Text: '{block}'")
                                myprint1(f"Position: ({x0}, {top}) - ({x1}, {bottom})")
                                myprint1(f"Width: {width}")
                                myprint1("---")
                            else:
                                myprint1("no bbox")

                        textlines=text.split("\n")
                        current_character=""
                        current_characters_split=[]
                        speech=""
                        mode="linear"
                        for line in textlines:
                            if line.isupper():
                                if " THEN " in line or "," in line:
                                    mode="split"
                                    if " THEN " in line:
                                        current_characters_split=line.split(" THEN ")   
                                    elif "," in line: 
                                        current_characters_split=line.split(",")    
                                    
                                    current_character=line
                                else:
                                    current_character=line
                                    mode="linear"
                            else:
                                speech=line
                                if mode=="split":
                                    speeches=speech.split("\n")
                                    charidx=0
                                    for sp in speeches:
                                        sp=sp.replace("- ","")
                                        current_character=current_characters_split[charidx]
                                        s=current_character+"\t"+sp+"\n"  # New line after each row
                                        myprint1("Add "+ current_character + " "+ speech)
                                        file.write(s)
                                        charidx=charidx+1

                                if mode=="linear":
                                    if current_character=="":
                                        current_character="__VOICEOVER"
                                    s=current_character+"\t"+speech+"\n"  # New line after each row
                                    myprint1("Add "+ current_character + " "+ speech)
                                    file.write(s)

                       
                       
    myprint1("Converted")                       
    myprint1(converted_file_path)
    return converted_file_path,encoding



def convert_pdf_title(file,table,titleIdx):
    myprint1("ROW LEN="+str(len(table[0]))+str(table[0]))
    for row in table[1:]:  # Skip header
        lastIdx=-1
        for i in range(len(row)):
            if row[i]:
                myprint1("has data"+str(row[i]))
                lastIdx=i
        myprint1("---------------------------------"+str(lastIdx))
        title = row[lastIdx]
        myprint1("item "+str(i)+" "+str(row[i]))
        myprint1("ROW LEN="+str(len(row))+str(row))
        myprint1(str(row))
        myprint1(">>>>>> TITLE"+str(title))
        if title:
            parts=title.split("\n")
            character=""
            speech=""
            for part in parts:
                if part.isupper():
                    character=character+" " +part
                else:
                    hascharacter=True
                    speech=speech+speech

            file.write(f"{character}\t{speech}\n")

def test_pdf_header(file,table,header):
            myprint1("test_pdf_headers"+str(header))
            dialogueCol=-1
            characterIdCol=-1
            titlesCol=-1
    
            myprint1("headers"+str(header))
            if 'CHARACTER' in header and 'DIALOGUE' in header:
                pdf_mode="PDF_CHARACTER_DIALOGUE"
            if 'Title' in header:
                pdf_mode="PDF_TITLE"
            if "CHARACTER" in header:
                characterIdCol = header.index('CHARACTER')
            if "Title" in header:
                titlesCol = header.index('Title')
            if "DIALOGUE" in header:
                dialogueCol = header.index('DIALOGUE')

            pdf_mode_title=titlesCol>-1
            if pdf_mode_title:
                myprint1("test_pdf_header pdf_mode_title")
                convert_pdf_title(file,table,titlesCol)
                return True

def detect_timecodes(text):
    # Regular expression pattern to match timecodes in the format [00:01:13.15]
    pattern = re.compile(r'\[\d{2}:\d{2}:\d{2}\.\d{2}\]')
    
    # Find all occurrences of the pattern in the text
    timecodes = pattern.findall(text)
    
    return timecodes
def detect_timecodes_raw(text):
    # Regular expression pattern to match timecodes in the format [00:01:13.15]
    pattern = re.compile(r'\d{2}:\d{2}:\d{2}\.\d{2} -->\n \d{2}:\d{2}:\d{2}\.\d{2}')
    
    # Find all occurrences of the pattern in the text
    timecodes = pattern.findall(text)
    
    return timecodes
def detect_word_table(table,forceMode="",forceCols={}):
    myprint1(f"detect_word_table {len(table.rows)} x {len(table.columns)}")
    for i in range(6):
        myprint1("try header "+str(i))
        header=table.rows[i]
        myprint1("header read ")        
        success, mode, character,dialog,map_= detect_word_header(header,forceMode,forceCols)
        myprint1("header success= "+str(success)+" "+str(map_))        
        if success:
            return success, mode, character,dialog,map_
    

    idx=0
    for cell in header.cells:
        t='NONE'        
        idx=idx+1
        map_[idx]={'type':t}
    return False,"",-1,-1,map_


def detect_universal_table(table_idx,table,forceMode="",forceCols={}):
    myprint1(f"detect_universal_table {table_idx} {table['row_count']} x {table['col_count']}")
    cells=table['cells']
    row_count=table['row_count']
    col_count=table['col_count']
    myprint1(f"detect_universal_table size= {row_count} x {col_count}")
    myprint1(f"detect_universal_table size= {cells}")
    
    
    for i in range(6):
        if i<row_count:
            myprint1("try header "+str(i))
            row=cells[i]
            myprint1("header read ")        
            success, mode, character,dialog,map_= detect_universal_header(table_idx,table,row,forceMode,forceCols)
            myprint1("header success= "+str(success)+" "+str(map_))        
            if success:
                return success, mode, character,dialog,map_
    

    idx=0
    for col_idx in range(0,table['col_count']):
        t='NONE'        
        idx=idx+1
        map_[idx]={'type':t}
    return False,"",-1,-1,map_

def isTableColumnCharacter(t):
    t=t.lower().strip()
    return t=="speaker id" or t=="character" or t=="charakter" or t=="personnage" or t=="character id" or t=="role" 
def isTableDialogColumn(t):
    t=t.lower().strip()
    return  t=="dialogue" or t=="dialogue list" or t=="dialouge" or t=="english" or t=="en ut" or  t=="dialog - deutsch"
def isTableColumnTitle(t):
    t=t.lower().strip()
    return  t=="title" or t=="titles"  
def isTableColumnBoth(t):
    t=t.lower().strip()
    return  t=="dialog with \nspeaker id" or t=='subtitle/ spotting' or t=='combined continuity' or t=='scene description'
def detect_word_header(header,forceMode="",forceCols={}):
    myprint1("-------------- detect_word_header -----------------")
    myprint1("forceMode"+str(forceMode))
    myprint1("forceCols"+str(forceCols))
    dialogueCol=-1
    characterIdCol=-1
    combinedContinuityCol=-1
    titlesCol=-1
    bothCol=-1
    sceneDescriptionCol=-1
    idx=0

    docx_mode_dialogue_characterid= False
    docx_mode_combined_continuity= False
    docx_mode_scenedescription= False
    docx_mode_dialogwithspeakerid=False

    if len(forceMode)>0:
        if forceMode=="DETECT_CHARACTER_DIALOG":
            dialogueCol=forceCols['DIALOG']
            characterIdCol=forceCols['CHARACTER']
            docx_mode_dialogue_characterid=True
            myprint1("CHARACTER "+str(characterIdCol))
            myprint1("DIALOG "+str(dialogueCol))
        else:
            myprint1("UNKNOWN FORCE MODE")
    else:
        
        myprint1(f"Header nCols={len(header.cells)}:")
        for cell in header.cells:
            t=cell.text.strip()
            myprint1("   * "+str(t))
            if isTableColumnCharacter(t):
                characterIdCol=idx
            elif isTableDialogColumn(t):
                dialogueCol=idx
            elif isTableColumnTitle(t):
                titlesCol=idx
            elif isTableColumnBoth(t):
                bothCol=idx
            idx=idx+1

    myprint1(f"dialogCol {dialogueCol}")
    myprint1(f"characterIdCol {characterIdCol}")
    myprint1(f"combinedContinuityCol {combinedContinuityCol}")
    myprint1(f"titlesCol {titlesCol}")
    myprint1(f"sceneDescriptionCol {sceneDescriptionCol}")
    myprint1(f"dialogWithSpeakerId {bothCol}")
    myprint1("detect_word_header assigned")
    docx_mode_dialogue_characterid= dialogueCol>-1 and characterIdCol>-1
    #docx_mode_combined_continuity= combinedContinuityCol>-1
    docx_mode_scenedescription= sceneDescriptionCol>-1 and titlesCol>-1
    docx_mode_dialogwithspeakerid=bothCol>-1
    myprint1("detect_word_header assigned mode")

    if dialogueCol>-1 and characterIdCol==-1:
        bothCol=dialogueCol

    mode=None
    character=None
    dialog=None
    map_={}
    idx=0
    myprint1("detect_word_headermap")

    if dialogueCol>-1 and characterIdCol>-1:
        mode="SPLIT"
        character=characterIdCol
        dialog=dialogueCol
    if sceneDescriptionCol>-1 and titlesCol>-1:
        mode="COMBINED"
        character=titlesCol
        dialog=titlesCol
    if   bothCol>-1:
        mode="COMBINED"
        character=bothCol
        dialog=bothCol

    myprint1(f"detect_word_header gen map mode={mode} d={dialog} c={character}")
    for cell in header.cells:
        myprint1(f"  test idx={idx} d={dialog}")
        t=cell.text.strip()
        if idx==dialog and idx==character:
            t='BOTH'
        elif idx==dialog:
            t='DIALOG'
        elif idx==character:
            t='CHARACTER'
        else:
            t='NONE'
        
        map_[idx]={'type':t}
        idx=idx+1

    myprint1("detect_word_header done")
    myprint1("map_"+str(map_))
    myprint1("mode"+str(mode))
    myprint1("dialog"+str(dialog))
    myprint1("character"+str(character))
    

    if docx_mode_dialogue_characterid:
        return True, mode, character,dialog,map_
    elif docx_mode_scenedescription:
        return True, mode, character,dialog,map_
    elif docx_mode_dialogwithspeakerid:
        return True, mode, character,dialog,map_
    elif bothCol>-1:
        return True, mode, character,dialog,map_
   # elif docx_mode_combined_continuity:
    #    return True, mode, character,dialog,map_
    else:
        myprint1("Tables but no automatic header match")
        return False, mode, character,dialog,map_

def convert_selection_to_char_dialog(selections):
    myprint1("sel"+str(selections))
    res=[]
    characterCol=-1
    dialogCol=-1
    for idx,selection in enumerate(selections):
        myprint1(f"table_idx={idx} valobj={selection}")

        for col_idx,valobj in enumerate(selection):
            myprint1(f"col={col_idx} valobj={selection[col_idx]}")
            val=selection[col_idx]['type']
            if val=="CHARACTER" or val=="LES DEUX":
                characterCol=col_idx
            if val=="DIALOG" or val=="LES DEUX":
                dialogCol=col_idx
    res={}
    res['character']=characterCol
    res['dialog']=dialogCol
    return res

def convert_universaltables_combined(file,univtable,bothCol):
    myprint1(f"convert_universaltables_combined {bothCol}")
    cells=univtable['cells']
    mode="LINEAR"
    cellLayoutMode=detectCellLayoutMode(univtable,bothCol)
    myprint1(f"convert_universaltables_combined {bothCol}")
    for row_idx,row in enumerate(cells):
        myprint1(f"---")
        
        isglobalmergecell=True
        comp=row[0]
        myprint1("convert_universaltables_combined comp= "+ comp)
        for colidx,cell in enumerate(row):
            myprint1("convert_universaltables_combined comp with"+ row[colidx])
            if comp!=cell:
                isglobalmergecell=False
        myprint1("convert_universaltables_combined ismerged"+ str(isglobalmergecell))

        if isglobalmergecell:
            continue

        
        t=row[bothCol].strip()

        if len(t)>0: 
            lines=t.split("\n")
            for k in lines:
                file.write(k+"\n")
            file.write("\n")
            if False:
                myprint1(f"convert_universaltables_combined row={row_idx} t={t} mode={cellLayoutMode}")     
                if cellLayoutMode=="NUM_SPACE_CHARACTERUPPERCASE_SEMICOLON_DIALOG":
                    myprint1(f"convert_universaltables_combined num space")     
                    current_character=extract_character_NUM_SPACE_CHARACTERUPPERCASE_SEMICOLON_DIALOG(t)
                    speech=extract_dialog_NUM_SPACE_CHARACTERUPPERCASE_SEMICOLON_DIALOG(t)
                    if current_character!=None and speech!=None:
                        myprint1("convert_universaltables_combined Add char="+ str(current_character) + " sp="+ str(speech))
                        res=current_character+"\t"+speech+"\n"  # New line after each row
                        file.write(res)
                if cellLayoutMode=="UNKNOWN_LAYOUT":
                    file.write("\n".join(lines))
                else:    
                    mode =detect_split_or_linear_mode_both(t)
                    myprint1(f"convert_universaltables_combined other")     
                    if mode=="LINEAR" :   
                        nlines=len(lines)
                        combinedCellMode="CHARACTER1_DIALOG1_CHARACTER2_DIALOG2"
                        if combinedCellMode=="CHARACTER1_DIALOG1_CHARACTER2_DIALOG2":
                            for line_idx,line in enumerate(lines):
                                if line.isupper() and line_idx+1<nlines :
                                    current_character=lines[line_idx].strip()
                                    speech=lines[line_idx+1].strip()
                                    myprint1("convert_universaltables_combined Add char="+ current_character + " sp="+ speech)
                                    res=current_character+"\t"+speech+"\n"  # New line after each row
                                    file.write(res)   
            #            elif combinedCellMode=="CHARACTER1_CHARACTER2_DIALOG1_DIALOG2":

                        
                        if False:
                            current_character=lines[0].strip()
                            speech=" ".join(lines[1:]).strip()

                            current_character=filter_character_name(current_character)
                            speech=filter_speech(speech)
                            myprint1("convert_docx_combined Add char="+ current_character + " sp="+ speech)
                            res=current_character+"\t"+speech+"\n"  # New line after each row
                            myprint1("convert_docx_combined Add ch="+ current_character + " sp="+ speech)
                            file.write(res)   
                    elif mode=="SPLIT":
                        sep=find_split_sep(t)
                                        
                        speakers=extract_speakers(lines[0])
                        myprint1("convert_docx_combined MODE SPLIT sep='"+str(sep)+"'")
                        if sep=="\n":
                            dialoguespl=lines[1:]
                        elif sep=='- ':
                            dialogue=" ".join(lines[1:])
                            dialoguespl=dialogue.split(sep)

                        myprint1("convert_docx_combined characters"+str(speakers))
                        filtered_array = [element for element in dialoguespl if element]
                        myprint1("convert_docx_combined dialogue"+str(filtered_array))

                        for index,k in enumerate(speakers):
                            myprint1("convert_docx_combined index"+str(index)+" k="+str(k))

                            current_character=filter_character_name(k)
                            speech= filter_speech(lines[index+1])
                            res=current_character+"\t"+speech+"\n"  # New line after each row
                            myprint1("convert_docx_combined Add ch="+ current_character + " sp="+ speech)
                            file.write(res)    
            
def convert_universaltables_separated(file,table,characterIdCol,dialogueCol):
    myprint1("convert_docx_characterid_and_dialogue characterIdCol="+str(characterIdCol)+" dialogueCol="+str(dialogueCol))
    # Iterate through each row in the table
    current_character=""
    is_song_sung_by_character=False
    cells=table['cells']
    start=0
    for row_idx,row in enumerate(cells):
        if row_idx<=start:
            continue
        
        dialogue=row[dialogueCol].strip()
        character=row[characterIdCol].strip()

        mode="LINEAR"
        dialogue=dialogue.replace("\n","")
        myprint1("----------")
        myprint1("dialogue"+dialogue)
        myprint1("character"+character)
        mode =detect_split_or_linear_mode_separated(character)
        

        if mode=="LINEAR" :   
        
            if len(character)>0:
                current_character=character
            if len(dialogue)>0:  
                is_didascalie=dialogue.startswith("(")
                is_song=dialogue.startswith("♪") 
                is_song_sung_by_character=is_song and len(character)>0 #song starts but has character name on it
                if is_song:
                    dialogue=filter_speech(dialogue)             
                    force_current_character="__SONG"
                    if is_song_sung_by_character:
                         force_current_character=current_character
                    s=force_current_character+"\t"+dialogue+"\n"  # New line after each row
                    myprint1("Add "+ force_current_character + " "+ dialogue)
                    file.write(s)
                else:
                    is_song_sung_by_character=False
                if not is_didascalie and not is_song: 
                    dialogue=filter_speech(dialogue)             
                    if current_character=="":
                        current_character="__VOICEOVER"
                    s=current_character+"\t"+dialogue+"\n"  # New line after each row
                    myprint1("Add "+ current_character + " "+ dialogue)
                    file.write(s)
        elif mode=="SPLIT":
            speakers=extract_speakers(character)
            myprint1("MODE SPLIT")
            myprint1("characters"+str(speakers))
            dialoguespl=dialogue.split("- ")
            filtered_array = [element for element in dialoguespl if element]
            myprint1("dialogue"+str(filtered_array))
            for index,k in enumerate(speakers):
                myprint1("index"+str(index)+" k="+str(k))

                current_character=k
                speech=filter_speech( filtered_array[index])
                s=current_character+"\t"+speech+"\n"  # New line after each row
                myprint1("Add "+ current_character + " "+ dialogue)
                file.write(s) 


def convert_universaltables_to_txt(file,univtables,params):
    if not "selection" in params:
        return False
    for k, table in enumerate(univtables):
        myprint1(" -------------  convert_universaltables_to_txt -----------")
        myprint1("convert_universaltables_to_txt params"+str(params))
        res=        convert_selection_to_char_dialog(params['selection'])
        char=res['character']
        dial=res['dialog']
        myprint1("convert_universaltables_to_txt params"+str(char)+" d="+str(dial))
        
        if char==dial:
            convert_universaltables_combined(file,table,char)
            return True
        else:
            convert_universaltables_separated(file,table,char,dial)
            return True

def detect_universal_header(table_idx, table,header,forceMode="",forceCols={}):
    myprint1("-------------- detect_universal_header "+str(table_idx)+" -----------------")
    myprint1("forceMode"+str(forceMode))
    myprint1("forceCols"+str(forceCols))
    dialogueCol=-1
    characterIdCol=-1
    combinedContinuityCol=-1
    titlesCol=-1
    bothCol=-1
    sceneDescriptionCol=-1
    idx=0

    docx_mode_dialogue_characterid= False
    docx_mode_combined_continuity= False
    docx_mode_scenedescription= False
    docx_mode_dialogwithspeakerid=False

    if len(forceMode)>0:
        if forceMode=="DETECT_CHARACTER_DIALOG":
            dialogueCol=forceCols['DIALOG']
            characterIdCol=forceCols['CHARACTER']
            docx_mode_dialogue_characterid=True
            myprint1("CHARACTER "+str(characterIdCol))
            myprint1("DIALOG "+str(dialogueCol))
        else:
            myprint1("UNKNOWN FORCE MODE")
    else:
        
        myprint1(f"Header nCols={len(header)}:")
        for col_idx,cell in enumerate(header):
            if cell!=None:
                t=cell.strip()
                myprint1("   * "+str(t))
                if isTableColumnCharacter(t):
                    characterIdCol=idx
                elif isTableDialogColumn(t):
                    dialogueCol=idx
                elif isTableColumnTitle(t):
                    titlesCol=idx
                elif isTableColumnBoth(t):
                    bothCol=idx
            idx=idx+1

    myprint1(f"dialogCol {dialogueCol}")
    myprint1(f"characterIdCol {characterIdCol}")
    myprint1(f"combinedContinuityCol {combinedContinuityCol}")
    myprint1(f"titlesCol {titlesCol}")
    myprint1(f"sceneDescriptionCol {sceneDescriptionCol}")
    myprint1(f"dialogWithSpeakerId {bothCol}")
    myprint1("detect_word_header assigned")
    docx_mode_dialogue_characterid= dialogueCol>-1 and characterIdCol>-1
    #docx_mode_combined_continuity= combinedContinuityCol>-1
    docx_mode_scenedescription= sceneDescriptionCol>-1 and titlesCol>-1
    docx_mode_dialogwithspeakerid=bothCol>-1
    myprint1("detect_word_header assigned mode")

    if dialogueCol>-1 and characterIdCol==-1:
        bothCol=dialogueCol

    mode=None
    character=None
    dialog=None
    map_={}
    idx=0
    myprint1("detect_word_headermap")

    if dialogueCol>-1 and characterIdCol>-1:
        mode="SPLIT"
        character=characterIdCol
        dialog=dialogueCol
    if sceneDescriptionCol>-1 and titlesCol>-1:
        mode="COMBINED"
        character=titlesCol
        dialog=titlesCol
    if   bothCol>-1:
        mode="COMBINED"
        character=bothCol
        dialog=bothCol

    myprint1(f"detect_word_header gen map mode={mode} d={dialog} c={character}")
    for col_idx,cell in enumerate(header):
        if cell!=None:
            t=cell.strip()
            if idx==dialog and idx==character:
                t='BOTH'
            elif idx==dialog:
                t='DIALOG'
            elif idx==character:
                t='CHARACTER'
            else:
                t='NONE'
        else:
            t='NONE'
            
        map_[idx]={'type':t}
        idx=idx+1

    myprint1("detect_word_header done")
    myprint1("map_"+str(map_))
    myprint1("mode"+str(mode))
    myprint1("dialog"+str(dialog))
    myprint1("character"+str(character))
    

    if docx_mode_dialogue_characterid:
        return True, mode, character,dialog,map_
    elif docx_mode_scenedescription:
        return True, mode, character,dialog,map_
    elif docx_mode_dialogwithspeakerid:
        return True, mode, character,dialog,map_
    elif bothCol>-1:
        return True, mode, character,dialog,map_
   # elif docx_mode_combined_continuity:
    #    return True, mode, character,dialog,map_
    else:
        myprint1("Tables but no ")
        return False, mode, character,dialog,map_


def test_word_header_and_convert(file,table,header,forceMode="",forceCols={}):
            myprint1(" -------------  test word header -----------")
            myprint1("test_word_header_and_convert forceMode"+str(forceMode))
            myprint1("test_word_header_and_convert forceCols"+str(forceCols))
            start=1
            dialogueCol=-1
            characterIdCol=-1
            combinedContinuityCol=-1
            titlesCol=-1
            bothCol=-1
            dialogWithSpeakerId=-1
            sceneDescriptionCol=-1
            idx=0

            docx_mode_dialogue_characterid= False
            docx_mode_combined_continuity= False
            docx_mode_scenedescription= False
            docx_mode_dialogwithspeakerid=False


            if len(forceMode)>0:
                myprint1("test_word_header_and_convert forceMode")
                if forceMode=="DETECT_CHARACTER_DIALOG":
                    dialogueCol=forceCols['DIALOG']
                    characterIdCol=forceCols['CHARACTER']
                    if dialogueCol==characterIdCol:
                        bothCol=dialogueCol
                    docx_mode_dialogue_characterid=True
                    myprint1("test_word_header_and_convert CHARACTER "+str(characterIdCol))
                    myprint1("test_word_header_and_convert DIALOG "+str(dialogueCol))
                else:
                    myprint1("test_word_header_and_convert UNKNOWN FORCE MODE")
            else:
                for cell in header.cells:
                    t=cell.text.strip()
                    myprint1("test_word_header_and_convert Header cell"+str(t))
                    if isTableColumnCharacter(t):
                        characterIdCol=idx
                    elif isTableDialogColumn(t):
                        dialogueCol=idx
                    elif t=="Scene Description":
                        sceneDescriptionCol=idx
                    elif isTableColumnTitle(t):
                        titlesCol=idx
                    elif t=="Dialog With \nSpeaker Id" or t=='Subtitle/ Spotting':
                        dialogWithSpeakerId=idx
                    elif t=="COMBINED CONTINUITY":
                        combinedContinuityCol=idx
                    idx=idx+1

                num_rows=len(table.rows)
                num_cols=len(table.columns)
        
                #FORCE IF NO HEADER BUT LIKE KOKON
                if num_cols==1:
                    characterIdCol=0
                    dialogueCol=0
                    start=0

                if num_cols==3:
                    t=table.rows[0].cells[0].text
                    myprint1("Test timecode"+str(t))
                    isTimecodeArrow=is_timecode_arrow_timecode_format(t)         
                    isTimecode=is_timecode_simple(t)         
                    myprint1("Test timecode"+str(isTimecode)+" "+str(isTimecodeArrow))
                    
                    if isTimecode or isTimecodeArrow:
                        characterIdCol=1
                        dialogueCol=2
                        start=0

                if dialogueCol>-1 and characterIdCol==-1:
                    bothCol=dialogueCol

                docx_mode_dialogue_characterid= dialogueCol>-1 and characterIdCol>-1
                docx_mode_combined_continuity= combinedContinuityCol>-1
                docx_mode_scenedescription= sceneDescriptionCol>-1 and titlesCol>-1
                docx_mode_dialogwithspeakerid=dialogWithSpeakerId>-1 
                if docx_mode_dialogue_characterid or docx_mode_combined_continuity or docx_mode_scenedescription or docx_mode_dialogwithspeakerid or bothCol>-1:
                    myprint1("Headers found")
                else:
                    myprint1("Headers not found")
                    myprint1(" CharacterId"+str(characterIdCol))
                    myprint1(" sceneDescriptionCol"+str(sceneDescriptionCol))
                    myprint1(" dialogueCol"+str(dialogueCol))
                    myprint1(" titlesCol"+str(titlesCol))
                    myprint1(" dialogWithSpeakerId"+str(dialogWithSpeakerId))
                    myprint1(" combinedContinuityCol"+str(combinedContinuityCol))
                    return False

            if bothCol>-1:
                convert_docx_combined(file,table,bothCol)
                return True
            elif docx_mode_dialogue_characterid:
                convert_docx_separated(file,table,start,dialogueCol,characterIdCol)
                return True
            elif docx_mode_scenedescription:
                convert_docx_scenedescription(file,table,sceneDescriptionCol,titlesCol)
                return True
            elif docx_mode_dialogwithspeakerid :
                convert_docx_dialogwithspeakerid(file,table,dialogWithSpeakerId)
                return True
            elif docx_mode_combined_continuity:
                convert_docx_combined_continuity(file,table)
                return True
            else:
                myprint1("Tables but no ")
                return False

def access_cell(df,i, j):
    try:
        cell_value = df.iloc[i, j]
        
        return cell_value
    except IndexError:
        myprint1(f"ERR Index out of range {i} {j}")
        return "Index out of range"
def test_xlsx_header_and_convert(file,df,rowidx,absCurrentOutputFolder,forceMode="",forceCols={}):
    myprint1(f"test_xlsx_header_and_convert {rowidx} ")
    dialogueCol=-1
    characterIdCol=-1
    combinedContinuityCol=-1
    bothCol=-1
    titlesCol=-1
    dialogWithSpeakerId=-1
    sceneDescriptionCol=-1
    idx=0

    docx_mode_dialogue_characterid= False
    docx_mode_combined_continuity= False
    docx_mode_scenedescription= False
    docx_mode_dialogwithspeakerid=False

    # Get the number of rows and columns
    num_rows = df.shape[0]
    num_cols = df.shape[1]

    myprint1(f"Size {num_rows} x {num_cols} ")


    if len(forceMode)>0:
        if forceMode=="DETECT_CHARACTER_DIALOG":
            dialogueCol=forceCols['DIALOG']
            characterIdCol=forceCols['CHARACTER']
            docx_mode_dialogue_characterid=True
            myprint1("CHARACTER "+str(characterIdCol))
            myprint1("DIALOG "+str(dialogueCol))
        else:
            myprint1("UNKNOWN FORCE MODE")
    else:
        row_data = df.iloc[rowidx, :]

        # Loop over the cells of the row
        for colidx, cell in enumerate(row_data):   
            t=access_cell(df,rowidx, colidx)
            myprint1("testcol"+str(rowidx)+" "+str(colidx)+" = "+str(t))
            if not pd.isna(cell):
                t=cell.strip()                
                myprint1("Header "+str(rowidx)+","+str(colidx)+str(t))
                if isTableColumnCharacter(t):
                    characterIdCol=idx
                elif isTableDialogColumn(t):
                    dialogueCol=idx
                elif t=="Scene Description":
                    sceneDescriptionCol=idx
                elif isTableColumnTitle(t):
                    titlesCol=idx
                elif isTableColumnBoth(t):
                    bothCol=idx
            idx=idx+1

        #FORCE IF NO HEADER BUT LIKE KOKON
        if num_cols==3:
            t=access_cell(df,0, 0)
            myprint1("Test timecode"+str(t))
            isTimecode=detect_timecodes_raw(t)         
            if isTimecode:
                characterIdCol=1
                dialogueCol=2
        

        if dialogueCol>-1 and characterIdCol ==-1:
            bothCol=dialogueCol


        docx_mode_dialogue_characterid= dialogueCol>-1 and characterIdCol>-1
        docx_mode_combined_continuity= combinedContinuityCol>-1
        docx_mode_scenedescription= sceneDescriptionCol>-1 and titlesCol>-1
        docx_mode_dialogwithspeakerid=dialogWithSpeakerId>-1
        if docx_mode_dialogue_characterid or docx_mode_combined_continuity or docx_mode_scenedescription or docx_mode_dialogwithspeakerid:
            myprint1("Headers found")
        else:
            myprint1("Headers not found")
            myprint1(" CharacterId"+str(characterIdCol))
            myprint1(" sceneDescriptionCol"+str(sceneDescriptionCol))
            myprint1(" dialogueCol"+str(dialogueCol))
            myprint1(" titlesCol"+str(titlesCol))
            myprint1(" bothCol"+str(bothCol))
            myprint1(" dialogWithSpeakerId"+str(dialogWithSpeakerId))
            myprint1(" combinedContinuityCol"+str(combinedContinuityCol))
      
    start=rowidx+1
    if dialogueCol>-1 and characterIdCol>-1:
        return convert_xlsx_split(file,df,characterIdCol,dialogueCol, absCurrentOutputFolder,start,forceMode="",forceCols={})
    elif bothCol>-1:
        return convert_xlsx_both(file,df,bothCol, absCurrentOutputFolder,start,forceMode="",forceCols={})
    else:
        myprint1("Unknown format ")
        return ""
            

def convert_xlsx_both(file_path,df,bothCol,absCurrentOutputFolder,start,forceMode="",forceCols={}):
    # Extract the columns "Character" and "English"
    both_column = df.iloc[:, bothCol]
    myprint1(f"START {start}")
    # Prepare the text file content
    lines = []
    lineidx=1
    for cell in both_column:
        if lineidx>start:
            if not pd.isna(cell):
                lines.append(f"{cell}")
        lineidx=lineidx+1

    converted_file_path= os.path.join(absCurrentOutputFolder,os.path.basename(file_path).replace(".xlsx",".converted.txt"))

    # Write to a text file
    with open(converted_file_path, 'w') as f:
        for line in lines:
            f.write(line + '\n')
    myprint1("Done")
    return converted_file_path

def convert_xlsx_split(file_path,df,charCol,dialCol, absCurrentOutputFolder,start,forceMode="",forceCols={}):
    myprint1("convert_xlsx_split")


    # Extract the columns "Character" and "English"
     # Access the columns by index
    character_column = df.iloc[:, charCol]
    dialog_column = df.iloc[:, dialCol]
    
    # Prepare the text file content
    lines = []
    rowidx=0
    for char, dia in zip(character_column, dialog_column):
        if rowidx>start:
            if pd.notna(char) and pd.notna(dia):
                myprint1("eng="+str(rowidx)+" dia='"+str(dia)+"'")
                dia=str(dia).replace("\n"," ")
                myprint1("Add line"+str(dia))
                myprint1("Add line linear")
                if dia.startswith("- "):
                        dia=dia.lstrip("- ")

                if dia.count("- ")>0:
                    spl=dia.split("- ")
                    chars=char.split("-")
                    if len(spl) == len(chars):
                        for index,k in enumerate(spl):
                            lines.append(f"{chars[index].upper()}\t{k}")        
                    else:
                        lines.append(f"{char.upper()}\t{dia}")                            
                elif dia.count("-")>0:
                    spl=dia.split("-")
                    chars=char.split("-")
                    if len(spl) == len(chars):
                        for index,k in enumerate(spl):
                            lines.append(f"{chars[index].upper()}\t{k}")        
                    else:
                        lines.append(f"{char.upper()}\t{dia}")                            


                else:      
                    lines.append(f"{char.upper()}\t{dia}")
        rowidx=rowidx+1

    converted_file_path= os.path.join(absCurrentOutputFolder,os.path.basename(file_path).replace(".xlsx",".converted.txt"))

    # Write to a text file
    with open(converted_file_path, 'w') as f:
        for line in lines:
            f.write(line + '\n')
    myprint1("Done")
    return converted_file_path

def convert_xlsx_to_txt(file_path,absCurrentOutputFolder,forceMode="",forceCols={}):
    myprint1("convert_xlsx_to_txt")
    myprint1("currentOutputFolder             :"+absCurrentOutputFolder)
    myprint1("Input             :"+file_path)
    df = pd.read_excel(file_path,header=None)
    # Load the Excel file

    # Get the number of rows and columns
    num_rows = df.shape[0]
    num_cols = df.shape[1]


    headerSuccess=False
    for i in range(0,10):
        myprint1(f"test {i}")
        converted_file_path=test_xlsx_header_and_convert(file_path,df,i,absCurrentOutputFolder,forceMode=forceMode,forceCols=forceCols)
        if len(converted_file_path)>0:
            return converted_file_path
    return ""



def get_paragraph_indentation(paragraph):
    """
    Returns the effective left indentation of the paragraph, considering the style hierarchy.
    """
    indent = paragraph.paragraph_format.left_indent
    style = paragraph.style
    while indent is None and style is not None:
        indent = style.paragraph_format.left_indent
        style = style.base_style
    return indent

def get_paragraph_style_hierarchy(paragraph):
    """
    Returns the style hierarchy of a paragraph.
    """
    styles = []
    style = paragraph.style
    while style is not None:
        styles.append(style.name)
        style = style.base_style
    return styles

def inspect_paragraphs_with_style_hierarchy(para):
        text = para.text.strip()
        indent = get_paragraph_indentation(para)
        styles = get_paragraph_style_hierarchy(para)
        myprint1(str({
            'text': text,
            'indent': indent,
            'style_hierarchy': ' > '.join(styles)
        }))

def word_has_paragraph_style_dialog(para):
    styles=get_paragraph_style_hierarchy(para)
    styles = ' > '.join(styles)
    return "DIALOG" in styles

def word_has_paragraph_style_character(para):
    styles=get_paragraph_style_hierarchy(para)
    styles = ' > '.join(styles)
    return "CHARACTER" in styles

def convert_pdftables_to_txt(listtables,converted_file_path):
    
    with open(converted_file_path, 'w', encoding='utf-8') as file:
        myprint1("conv open")
        for table in listtables:
            row_count=table['row_count']
            col_count=table['col_count']
            myprint1(f"conv size= {row_count} x {col_count}")
            cells=table['cells']
            for rowidx,row in enumerate(cells):
                myprint1(f"conv row {rowidx} {row}")

                character=None
                dial=None
                
                for idx,k in enumerate(row):
                    myprint1(f"conv cell {k} ")

                    if idx ==2:
                        character=k
                    if idx==3:
                        dial=k
                        dial=dial.replace("\n"," ")
                        dial=filter_speech(dial)
                        character=filter_character_name(character)
                        myprint1(f"conv add {character} {dial} ")
                        if len(character)>0 and len(dial)>0:
                            file.write(character +"\t"+dial +'\n')
    return converted_file_path


def convert_word_withstyles_to_plaintext(doc,file_path,absCurrentOutputFolder):
    converted_file_path=get_docx_to_txt_converted_filepath(file_path,absCurrentOutputFolder)
    current_character=None
    current_dialog=None
    with open(converted_file_path, 'w', encoding='utf-8') as file:
        for para in doc.paragraphs: 
            if word_has_paragraph_style_character(para):
                char=filter_character_name(para.text.upper())
                current_character=char
            if word_has_paragraph_style_dialog(para):
                speech=para.text
                current_dialog=filter_speech(speech)
            if current_character != None and current_dialog!=None:
                s=current_character+"\t"+current_dialog+"\n"  # New line after each row
                myprint1("Add "+ current_character + " "+ current_dialog)
                file.write(s)
                current_dialog=None
    return converted_file_path
def detect_word_styles_character_dialog(doc):
    has_dialog_style=False
    has_character_style=False
    for para in doc.paragraphs:
        if word_has_paragraph_style_dialog(para):
            has_dialog_style=True
        if word_has_paragraph_style_character(para):
            has_character_style=True
    if has_dialog_style and has_character_style:
        return True
    return False

def detect_plaintext_indented(doc):
    isIndented=False
    nParagraphs=0
    nIndented=0
    # Iterate over paragraphs
    for para in doc.paragraphs:
        myprint1('-------')
        if len(para.text)>0:
            myprint1(str(para.text)+" ")
            nParagraphs=nParagraphs+1
            indent = para.paragraph_format.left_indent
            rightindent=para.paragraph_format.right_indent
            first= para.paragraph_format.first_line_indent
            for  run in para.runs:
            
                myprint1("RUN UPPER"+str(run.text)+" "+ "allcaps="+str(run.font.all_caps))
                if run.font.all_caps:
                    myprint1("UPPER")
            myprint1("leftindent="+str(indent)+" rightindent="+str(rightindent)+" first="+str(first))
            inspect_paragraphs_with_style_hierarchy(para)
            if indent is not None:
                myprint1("INDENTED "+str(para.text))
                nIndented=nIndented+1
        
    indentedPc=nIndented/nParagraphs
    myprint1("Plaintext detect indented")
    myprint1("Plaintext "+str(nIndented)+" / "+str(nParagraphs)+" paragraphs= "+str(indentedPc)+" indented")
    if indentedPc>0.1:
        return True
    else:
        return False        

def get_universal_converted_path(file_path,absCurrentOutputFolder):
    fileext=get_file_extension(file_path)
    if fileext==".docx":
        return    get_docx_to_txt_converted_filepath(file_path,absCurrentOutputFolder)    
    if fileext==".pdf":
        return    get_pdf_to_txt_converted_filepath(file_path,absCurrentOutputFolder)    
    if fileext==".xlsx":
        return    get_xlsx_to_txt_converted_filepath(file_path,absCurrentOutputFolder)    
    return ""
def get_docx_to_txt_converted_filepath(file_path,absCurrentOutputFolder):
    converted_file_path=os.path.join(absCurrentOutputFolder, os.path.basename(file_path).replace(".docx",".converted.txt"))
    return converted_file_path
def get_pdf_to_txt_converted_filepath(file_path,absCurrentOutputFolder):
    converted_file_path=os.path.join(absCurrentOutputFolder, os.path.basename(file_path).replace(".pdf",".converted.txt"))
    return converted_file_path
def get_xlsx_to_txt_converted_filepath(file_path,absCurrentOutputFolder):
    converted_file_path=os.path.join(absCurrentOutputFolder, os.path.basename(file_path).replace(".xlsx",".converted.txt"))
    return converted_file_path
def get_doc_to_txt_converted_filepath(file_path,absCurrentOutputFolder):
    converted_file_path=os.path.join(absCurrentOutputFolder, os.path.basename(file_path).replace(".doc",".converted.txt"))
    return converted_file_path
def convert_word_to_txt(file_path,absCurrentOutputFolder,forceMode="",forceCols={}):
    myprint1("convert_docx_to_txt")
    myprint1("currentOutputFolder : "+absCurrentOutputFolder)
    myprint1("Input               : "+file_path)
    converted_file_path=""
    if ".docx" in file_path:
        myprint1(".docx")
        converted_file_path=get_docx_to_txt_converted_filepath(file_path,absCurrentOutputFolder)
    elif ".doc" in file_path:
        myprint1(".doc")
        converted_file_path=get_docx_to_txt_converted_filepath(file_path,absCurrentOutputFolder)

#        myprint1("Convert doc to docx  ")
 #       docx_file_path = convert_doc_to_docx(file_path,absCurrentOutputFolder)
  #      myprint1("Output         :"+os.path.abspath(docx_file_path))
   #     converted_file_path=os.path.abspath(docx_file_path).replace(".docx",".converted.txt")
    #    file_path=docx_file_path
    else:
        myprint1("other extension")

    myprint1("Converted file path : "+converted_file_path)    
    myprint1("Doc opening         : "+file_path)
    myprint1("OPEN DOCX 3")

    doc = Document(file_path)
    myprint1("Doc opened")

    with open(converted_file_path, 'w', encoding='utf-8') as file:
        # Check if there are any tables in the document
        myprint1("Table count             : "+str(len(doc.tables)))
       # Iterate through the tables in the document and get their dimensions
        table_dimensions = []
        if len(doc.tables) > 0:

            for table in doc.tables:
                num_rows = len(table.rows)
                num_columns = len(table.columns)
                table_dimensions.append((num_rows, num_columns))

            # Print the dimensions of each table
            for i, (rows, cols) in enumerate(table_dimensions):
                myprint1(f"Table {i+1}: {rows} rows, {cols} columns")

            # Get the first table   
            table = doc.tables[0]
            
            headerSuccess=False
            for i in range(3):
                header=table.rows[i]
                success=test_word_header_and_convert(file,table,header,forceMode=forceMode,forceCols=forceCols)
                if success:
                    headerSuccess=True
                    break
            if not headerSuccess:
                return ""

        else:
            hasWordStyles=detect_word_styles_character_dialog(doc)
            if hasWordStyles:
                convert_word_withstyles_to_plaintext(doc,file_path,absCurrentOutputFolder)
            else:
                myprint1("No tables, convert to plain text")
                isIndented=detect_plaintext_indented(doc)
                convert_docx_plain_text(file,doc)
                myprint1("Converted file path : "+converted_file_path)  
                myprint1("converted to plain text")
        
    return converted_file_path

def get_all_characters(breakdown):
    #myprint1("get_all_characters")
    all_characters=[]
    for item in breakdown:
        if item["type"]=="SPEECH":
            character=item["character"]
            if character==None:
                character="ERR ???"
                
            #myprint1("test  CHARACTER"+str(character)+" "+str(all_characters))

            if not character in all_characters:
                myprint1("ADD NEW CHARACTER"+character)
                all_characters.append(character)
    return all_characters

splitables=[" TALK TO "," TO "]
def hasSplitable(character):
    for i in splitables:
        if i in character:
            return i
    return ""

def has_duplicates(lst):
    return len(lst) != len(set(lst))

def indices_of_duplicates(lst):
    from collections import defaultdict
    index_map = defaultdict(list)  # Stores list of indices for each item
    for index, item in enumerate(lst):
        index_map[item].append(index)
    
    # Filter items that appear more than once and collect their indices
    return {item: indices for item, indices in index_map.items() if len(indices) > 1}

def map_semi_duplicates(names):
    normalized_map = {}  # Maps normalized names to their first occurrence
    duplicates = {}  # Stores mappings of semi-duplicate entries

    for name in names:
        normalized = name.replace(" ", "")  # Remove spaces to normalize
        if normalized in normalized_map:
            # Map current name to the first occurrence of this normalized form
            duplicates[name] = normalized_map[normalized]
        else:
            # Store the first occurrence of this normalized form
            normalized_map[normalized] = name

    return duplicates

def is_action_verb_charactername(charactername):
    # List of action verbs
   
  
    # Regex to match "<NAME> <ACTION>"
    pattern = r"(\b\w+\b) (" + '|'.join(action_verbs) + r")\b"

    # Search for matches
    match = re.search(pattern, charactername)

    if match:
        myprint1("Match found:", match.groups())
    else:
        myprint1("No match found.")

def extract_character_and_action(charactername, action_verbs):
    # Create a regex pattern that matches a word followed by any of the action verbs
    pattern = r"(\b\w+\b) (" + '|'.join(action_verbs) + r")\b"

    # Search for matches in the provided line
    match = re.search(pattern, charactername)
    if match:
        # Returns the character's name and the action verb
        return match.groups()
    else:
        return None

def merge_breakdown_character_by_replacelist(breakdown,replace_list):
    myprint1("merge_breakdown_character_by_replacelist")
    checkIfAlreadyNamed=False
    for item in breakdown:
        if item["type"]=="SPEECH":
            character=item["character"]
            if character in replace_list:
                firstchar=replace_list[character]
                item['character']=firstchar                   

    return breakdown

def is_multiple_character(char):
    if char==None:
        return False
    return " AND " in char
def count_pattern_occurrences2(line):
    # Define the regular expression pattern
    pattern = re.compile(
        r'(.*?): (.*)',  # Text line with : and <i></i> tags
    )
    
    # Find all occurrences of the pattern
    matches = pattern.findall(line)

    # Return the number of matches
    return len(matches)

def split_text_by_uppercase(text):
    """
    Splits the text by uppercase words which are assumed to be speaker names,
    avoiding splitting character names like "DR WALSH".

    :param text: The text to split.
    :return: A list of strings, each starting with a speaker name.
    """
    # Split the text by spaces to get words
    words = text.split()

    if "H2O," in words:
        dodebug=True
        for word in words:
            sel=  word.isupper() and word[:-1] != "," and len(word) > 1
            print(f"word {word} {sel} last={word[-1]}")
    

    # Find indices of all uppercase words longer than one character
    uppercase_indices = [i for i, word in enumerate(words) if word.isupper() and word[-1]!="," and len(word) > 1]
    

    # Remove indices if the previous word is also uppercase
    filtered_indices = []
    for i in range(len(uppercase_indices)):
        if i == 0 or (uppercase_indices[i] - 1 != uppercase_indices[i - 1]):
            filtered_indices.append(uppercase_indices[i])
    

    # Create segments based on the filtered indices
    segments = []
    start_index = 0
    for index in filtered_indices:
        segment = " ".join(words[start_index:index])
        if segment:  # Add non-empty segments
            segments.append(segment.strip())
        start_index = index
    # Add the last segment
    segments.append(" ".join(words[start_index:]).strip())
    
    return segments


def split_AND_character(breakdown):
    myprint1("merge_breakdown_character_by_replacelist")
    for item in breakdown:
        if item["type"]=="SPEECH":
            character=item["character"]
            if is_multiple_character(character):
                characters=character.split(" AND ")
                first=characters[0]
                chidx=0
                for ch in characters:
                    if chidx==0:
                        item['character']=ch
                    else:
                        item2={
                            'character':ch,
                            'type':item['type'],
                            'scene_id':item['scene_id'],
                            'line_idx':item['line_idx'],
                            'character_raw':character,
                            'speech':item['speech']
                        }
                        breakdown.append(item2)

                    chidx=chidx+1

    return breakdown

def merge_breakdown_character_talking_to(breakdown,all_characters):
    myprint1("merge_breakdown_character_talking_to")
    replaceList={}
    checkIfAlreadyNamed=False
    for item in breakdown:
        #smyprint1(str(item))
        if item["type"]=="SPEECH":
            character=item["character"]
            if character==None:
                character="ERR CHAR"
            if character!=character.strip():
                character=character.strip()

            splitable=hasSplitable(character)
            if splitable!="":
                #myprint1(" has to"+character)
                characters=character.split(splitable)
                #myprint1(" split"+str(characters))

                if checkIfAlreadyNamed:
                    are_parts_characternames=True
                    for k in characters:
                        if k in all_characters:
                            are_parts_characternames=True
                        else:
                            are_parts_characternames=False
                            break
                    if are_parts_characternames:
                        firstchar=characters[0].strip()
                        #myprint1("REPLACE "+character+" with "+str(firstchar))
                        replaceList[character]=firstchar
                        item['character']=firstchar                   
                else:
                    firstchar=characters[0].strip()
                    #myprint1("REPLACE"+character+" with "+str(firstchar))
                    replaceList[character]=firstchar
                    item['character']=firstchar                   

    return breakdown,replaceList
def filter_speech2(s):
    res=s.replace("♪","").replace("Â§","").replace("§","")
    #filter songs
    return res

def save_string_to_file(text, filename):
    """Saves a given string `text` to a file named `filename`."""
    myprint1(" > Write to "+filename)
    with open(filename, 'w', encoding='utf-8') as file:
        file.write(text)


def is_character_name_valid(char):
    isNote= "NOTE D'AUTEUR" in char
    isEnd = ("END CREDITS" in char )
    isNar=("NARRATIVE TITLE" in char)
    isOST =("ON-SCREEN TEXT" in char )
    isMusic =( char.startswith("♪") )
    
    isMain =("MAIN TITLE" in char )
    isOpen=("OPENING CREDITS" in char)
    return (not isNote) and (not isEnd) and (not isNar) and (not isOST) and (not isOpen) and (not isMain) and (not isMusic)
def detect_text_in_brackets(line):
    # Define the pattern to match text in brackets followed by any text
    pattern = re.compile(r'^\[\w+\].*')
    
    # Check if the line matches the pattern
    match = pattern.match(line)
    
    # Return True if a match is found, otherwise False
    return match is not None

def extract_text_in_brackets(line):
    # Define the pattern to match text in brackets
    pattern = re.compile(r'\[(\w+)\]')
    
    # Search for the pattern in the line
    match = pattern.search(line)
    
    # If a match is found, return the captured group (text within brackets)
    if match:
        return match.group(1)
    else:
        return None
def extract_text_not_in_brackets(line):
        # Define the pattern to match text in brackets
    pattern = re.compile(r'\[.*?\]')
    
    # Replace the text in brackets with an empty string
    result = pattern.sub('', line)
    
    # Strip any leading or trailing whitespace
    return result.strip()
def extract_pattern_occurrences2(content):
    myprint1("extract in "+str(content))
    # Define the regular expression pattern
    pattern = re.compile(
    
    r'^([A-Z]+): (.*)$',
    #   r'^(.*?): (.*)$',  # Text line with : and <i></i> tags
        re.MULTILINE
    )
    

    # Find all occurrences of the pattern
    matches = pattern.findall(content)

    # Extract the required fields into a list of dictionaries
    extracted_data = [
        {
            "character": match[0].strip(),
            "dialog": match[1].strip()
        }
        for match in matches
    ]
    print("res="+str(extracted_data))

    return extracted_data
def is_character_didascalie(c):
    c=c.lower()
    return c=="sigh" or c=="sighs" or c=="cries" or c=="laughs"
#################################################################
# PROCESS
def process_script(script_path,output_path,script_name,countingMethod,encoding,forceMode="",forceCols={},ignoreBeginning=0,ignoreEnd=0,forceCharacterMode=None):
    myprint1("  > -----------------------------------")
    myprint1("  > SCRIPT PARSER version 1.22")
    myprint1("  > Script path       : "+script_path)
    myprint1("  > Output folder     : "+output_path)
    myprint1("  > Script name       : "+script_name) 
    myprint1("  > Forced encoding   : "+encoding)
    myprint1("  > Forced mode       : "+str(forceCharacterMode))
    myprint1("  > Counting method   : "+countingMethod)

    if not os.path.exists(output_path):
        os.mkdir(output_path)

    file_name = os.path.basename(script_path)
    name, extension = os.path.splitext(file_name)
    myprint1("  > File name         : "+file_name)
    myprint1("  > Extension         : "+extension)
    if not is_supported_extension(extension):
        myprint1("  > File type "+extension+" not supported.")
        return

    if extension.lower()==".docx":

        if "CLEAR CUT" in file_name:
            myprint1(" !!!!!!!!!! FORCE")
            forceMode="DETECT_CHARACTER_DIALOG"
            forceCols={
                "CHARACTER":5,
                "DIALOG":6
            }
        converted_file_path=convert_word_to_txt(script_path,forceMode=forceMode,forceCols=forceCols)
        if len(converted_file_path)==0:
            myprint1("  > Conversion failed 1")
            info ={
                "success":False,
                "fail_desc":"convert_word_to_txt failed"
            }
            return info

  #          return None,None,None,None,None,None
        return process_script(converted_file_path,output_path,script_name,countingMethod,forceMode=forceMode,forceCols=forceCols)




    if extension.lower()==".xlsx":

        converted_file_path=convert_xlsx_to_txt(script_path,forceMode=forceMode,forceCols=forceCols)
        if len(converted_file_path)==0:
            myprint1("  > Conversion failed 1")
            info ={
                "success":False,
                "fail_desc":"convert_word_to_txt failed"
            }
            return info
 #           return None,None,None,None,None,None
        return process_script(converted_file_path,output_path,script_name,countingMethod,forceMode=forceMode,forceCols=forceCols)





    if extension.lower()==".pdf":
        converted_file_path,encoding=convert_pdf_to_txt(script_path)
        myprint1("process 1")
        if len(converted_file_path)==0:
            myprint1("  > Conversion failed 1")
            info ={
                "success":False,
                "fail_desc":"convert_word_to_txt failed"
            }
            return info

#            return None,None,None,None,None,None
        myprint1("process")
        return process_script(converted_file_path,output_path,script_name,countingMethod,forceMode=forceMode,forceCols=forceCols)

    uppercase_lines=[]
    current_scene_id=""
    wasEmptyLine=False
    scene_characters_map={}
    character_linecount_map={}
    character_order_map={}
    character_order={}
    character_count=1
    character_textlength_map={}
    character_scene_map={}
    current_scene_count=1
    breakdown=[]

    is_verbose=False
    encoding_info = detect_file_encoding(script_path)
    myprint1("  > Encoding info     : "+str(encoding_info['encoding']))
    encoding_used=encoding_info['encoding']
    myprint1("  > Encoding used     : "+str(encoding_used))
    encoding_tested=test_encoding(script_path)
    encoding_used=encoding_tested
    if not (encoding==""):
        myprint1("  > Force encoding    :"+encoding)
        encoding_used=encoding
    if encoding_used==None:
        encoding_used='utf-8'


    scene_separator=getSceneSeparator(script_path,encoding_used)
    myprint1("  > Scene separator   : "+scene_separator)

    if forceCharacterMode!=None:
        character_mode=forceCharacterMode
    else:
        character_mode=detectCharacterSeparator(script_path,encoding_used)
    myprint1("  > Character mode    : "+str(character_mode))
    
    character_sep_type=getCharacterSepType(character_mode)
    myprint1("  > Character sep type    : "+str(character_sep_type))
    
    if scene_separator=="EMPTYLINES_SCENE_SEPARATOR":
        current_scene_id="Scene 1"
    # Open the file and process each line
    line_idx=1
    isEmptyLine=False
    multiline_current_character_text = None
    multiline_current_lines_of_dialog = []
    multiline_in_pattern = False
    totalLines=count_lines_in_file(script_path,encoding)
    firstLineIdxToTreat=ignoreBeginning+1
    lastLineIdxToTreat=totalLines-ignoreEnd+1
    wasDialog=None
    wasChar=False
    wasTimecode=False
    isTranslation=False
    myprint1(f"  > Treat lines from {firstLineIdxToTreat} to {lastLineIdxToTreat} ")

    nNonEmptyLines=count_nonempty_lines_in_file(script_path,encoding)
    if nNonEmptyLines==0:
        return 
    with open(script_path, 'r', encoding=encoding_used) as file:
        myprint1("  > Opened    : "+str(script_path))
        for line in file:
            if line_idx<firstLineIdxToTreat or line_idx>lastLineIdxToTreat:
                myprint1("  > IGNORE LINE    : "+str(line_idx))
                line_idx=line_idx+1
                continue 
            myprint1("  -----------------")
            myprint1("  > Line    : "+str(line))
            #line = line  # Remove any leading/trailing whitespace
            trimmed_line = line.strip()  # Remove any leading/trailing whitespace
    
            isNewEmptyLine=len(trimmed_line)==0
            myprint1("  > Sep")

            if scene_separator=="EMPTYLINES_SCENE_SEPARATOR":
                if (not isNewEmptyLine) and  (isEmptyLine and wasEmptyLine):
                    current_scene_count=current_scene_count+1
                    current_scene_id = extract_scene_name(line,scene_separator,current_scene_count)
                    if is_verbose:
                        myprint1("  > ---------------------------------------")
                    #myprint1(f"Scene Line: {line}")
    
            isEmptyLine=len(trimmed_line)==0
            if character_sep_type=="CHARACTER_MODE_SINGLELINE":
                if is_verbose:
                    myprint1("  > Line "+str(line_idx))
                if len(trimmed_line)>0:
                    myprint1("  > trimmed")   
                    if is_scene_line(line) or (isEmptyLine and wasEmptyLine):
                        current_scene_count=current_scene_count+1
                                        
                        current_scene_id = extract_scene_name(line,scene_separator,current_scene_count)
                        breakdown.append({"line_idx":line_idx,"scene_id":current_scene_id,"type":"SCENE_SEP" })    
                        if is_verbose:
                            myprint1("  > --------------------------------------")
                        myprint1(f"  > Scene Line: {current_scene_id}")
                    else:
                            myprint1("  > not scene")   
                            if True:#current_scene_id!=1:

                                if character_mode=="CHARACTERUPPERCASE_DIALOG":
                                    split_text = split_text_by_uppercase(trimmed_line)
                                    for k in split_text:
                                        myprint1("SUBSEC part="+k)
                                        trimmed_line=k
                                        is_speaking=is_character_speaking(trimmed_line,character_mode)
                                        if is_speaking:
                                            character_name=extract_character_name(trimmed_line,character_mode)
                                            myprint1("      > speaking 2b    char="+str(character_name))   
                                            character_name=filter_character_name(character_name)
                                            myprint1("      > speaking 2b filchar="+str(character_name))   
                                            if is_character_name_valid(character_name):
                                                myprint1("      > speaking 2b valid")   
                                                spoken_text=extract_speech(trimmed_line,character_mode,character_name)
                                                spoken_text=filter_speech(spoken_text)
                                                character_name,spoken_text=ensure_dialog_starts_with_uppercase(character_name,spoken_text)
                                                myprint1("      > speaking 2b spoken="+str(spoken_text))   
                                                breakdown.append({"scene_id":current_scene_id,
                                                        "character_raw":character_name,
                                                        "line_idx":line_idx,"speech":spoken_text,"type":"SPEECH", "character":character_name })    

                                else:



                                    is_speaking=is_character_speaking(trimmed_line,character_mode)
                                    myprint1("  > speaking"+str(is_speaking))   

                                    if is_verbose:
                                        myprint1("    IsSpeaking "+str(is_speaking)+" "+trimmed_line)
                                    if is_speaking:
                                        myprint1("  > speaking1"+str(trimmed_line)+" "+str(character_mode))   

                                        character_name=extract_character_name(trimmed_line,character_mode)
                                        myprint1("  > speaking 2b char="+str(character_name))   
                                        character_name=filter_character_name(character_name)
                                        myprint1("  > speaking 2")   
                                        if is_verbose:
                                            myprint1("   name="+str(character_name))
                                        if not character_name == None:
                                            if is_character_name_valid(character_name):
                                                #remove character name for stats
                                                spoken_text=extract_speech(trimmed_line,character_mode,character_name)
                                                spoken_text=filter_speech(spoken_text)
                                                myprint1(f"Add amend char={character_name} d={spoken_text} ")
                                                breakdown.append({"scene_id":current_scene_id,
                                                                "character_raw":character_name,
                                                                "line_idx":line_idx,"speech":spoken_text,"type":"SPEECH", "character":character_name })    
                                                if is_verbose:
                                                    myprint1("   text="+str(spoken_text))
                                        
                                    else:
                                        breakdown.append({"line_idx":line_idx,"text":trimmed_line,"type":"NONSPEECH" })    
            elif character_sep_type=="CHARACTER_MODE_MULTILINE":
                if is_scene_line(line):
                        
                        current_scene_count=current_scene_count+1
                                        
                        current_scene_id = extract_scene_name(line,scene_separator,current_scene_count)
                        breakdown.append({"line_idx":line_idx,"scene_id":current_scene_id,"type":"SCENE_SEP" })    
                        if is_verbose:
                            myprint1("  > --------------------------------------")
                        myprint1(f"  > Scene Line: {current_scene_id}")
                    # Check for uppercase text
                else:
                    if character_mode=="CHARACTER_NEWLINE_DIALOG_NEWLINE_NEWLINE":
                        myprint1(f"  > Not a scene Line: {trimmed_line}")
                        if trimmed_line.isupper() and wasEmptyLine:
                            myprint1("  > upper and not in pattern, set uppercase = "+trimmed_line)

                            multiline_current_character_text = trimmed_line
                            multiline_current_lines_of_dialog = []
                            multiline_in_pattern = True
                        else:
                            myprint1("  > in pattern")

                            if isEmptyLine:  # Empty line
                                myprint1("  > empty line")

                                # Check if we have hit the double newline
    #                            if not multiline_current_lines_of_text or multiline_current_lines_of_text[-1] == '':
                                if True:
                                    myprint1("  > double line")
                                    # Double newline indicates end of current pattern
                                    if multiline_current_character_text:
                                        character_name=filter_character_name(multiline_current_character_text)
                                        for k in multiline_current_lines_of_dialog:
                                            #speech='\n'.join(multiline_current_lines_of_text).strip()       
                                            speech=filter_speech(k)           
                                            if k=="TRANSLATION":
                                                break; 
                                            myprint1("  > Add    : "+str(character_name)+" "+speech)
                                            breakdown.append({"scene_id":current_scene_id,
                                                                        "character_raw":character_name,
                                                                        "line_idx":line_idx,"speech":speech,"type":"SPEECH", "character":character_name })    
                                    multiline_current_character_text = None
                                    multiline_current_lines_of_dialog = []
                                    multiline_in_pattern = False
                                else:
                                    multiline_current_lines_of_dialog.append('')
                            else:
                                multiline_current_lines_of_dialog.append(trimmed_line)
                    elif character_mode=="NUM_TIMECODE_ARROW_TIMECODE_NEWLINE_MULTILINEDIALOG":
                        multiline_current_character_text="PERSONNAGE"
                        isTimecode=is_NUM_TIMECODE_ARROW_TIMECODE(trimmed_line)
                        if not isTimecode and not isEmptyLine:
                            myprint1(f"wasTimecode={wasTimecode} empty={isEmptyLine}")
                            speech=filter_speech(trimmed_line)
                            myprint1(f"Add line {multiline_current_character_text} {speech} ")
                            breakdown.append({"scene_id":current_scene_id,
                                                    "character_raw":multiline_current_character_text,
                                                    "line_idx":line_idx,"speech":speech,"type":"SPEECH", "character":multiline_current_character_text })    
                        if isEmptyLine:
                            myprint1(f"reset Timecode ")
                            wasTimecode=False

                        wasTimecode=is_NUM_TIMECODE_ARROW_TIMECODE(trimmed_line)
                        if wasTimecode:
                            myprint1(f"Timecode ")
                    
                    elif character_mode=="NUM_SEMICOLON_TIMECODE_SPACE_TIMECODE_SPACE_TIME":
                        multiline_current_character_text="PERSONNAGE"
                        isTimecode=is_NUM_SEMICOLON_TIMECODE_SPACE_TIMECODE_SPACE_TIME(trimmed_line)
                        if not isTimecode and not isEmptyLine:
                            speech=filter_speech(trimmed_line)
                            myprint1(f"Add line {multiline_current_character_text} {speech} ")
                            breakdown.append({"scene_id":current_scene_id,
                                                    "character_raw":multiline_current_character_text,
                                                    "line_idx":line_idx,"speech":speech,"type":"SPEECH", "character":multiline_current_character_text })    
                        if isEmptyLine:
                            wasTimecode=False

                        wasTimecode=is_NUM_SEMICOLON_TIMECODE_SPACE_TIMECODE_SPACE_TIME(trimmed_line)
                        
                    elif character_mode==    "TIMECODE_HYPHEN_TIMECODE_NEWLINE_CHARACTER_NEWLINE_DIALOG":
                        if isEmptyLine:
                            myprint1(f"reset waschar")
                            wasChar=False
                            wasTimecode=False
                        if wasChar:
                            wasTimecode=False
                            speech=filter_speech(trimmed_line)
                            myprint1(f"Add line {multiline_current_character_text} {speech} ")
                            breakdown.append({"scene_id":current_scene_id,
                                                    "character_raw":multiline_current_character_text,
                                                    "line_idx":line_idx,"speech":speech,"type":"SPEECH", "character":multiline_current_character_text })    
                        if wasTimecode:
                            multiline_current_character_text=trimmed_line
                            multiline_current_character_text=filter_character_name(multiline_current_character_text)     
                            myprint1(f"Set char {multiline_current_character_text} ")
                            wasChar=True
                        wasTimecode=is_TIMECODE_HYPHEN_TIMECODE(trimmed_line)
                        
                    elif character_mode==    "TIMECODE_NEWLINE_CHARACTERINBRACKETS_DIALOG_NEWLINE_NEWLINE":
                        myprint1(" > line2"+trimmed_line)
                        isTimecode=detect_timecodes(trimmed_line)
                        if not isTimecode:
                            has_character=detect_text_in_brackets(trimmed_line)
                            candidate_char=extract_text_in_brackets(trimmed_line)
                            if has_character  and not is_character_didascalie(candidate_char):
                                dialog= extract_text_not_in_brackets(trimmed_line) 
                                if dialog and len(dialog.strip())>0:
                                    multiline_current_character_text=candidate_char
                                    multiline_current_character_text=filter_character_name(multiline_current_character_text)
                                    speech=filter_speech(dialog)
                                    character_name=multiline_current_character_text
                                    
                                    myprint1(f"Add > ch={character_name} dialen={len(dialog.strip())} dia='{dialog}' spee={speech}")
                                    breakdown.append({"scene_id":current_scene_id,
                                                                            "character_raw":character_name,
                                                                            "line_idx":line_idx,"speech":speech,"type":"SPEECH", "character":character_name })    
                                    
                            else:
                                if len(trimmed_line)>0 and multiline_current_character_text!=None:
                                    speech=filter_speech(trimmed_line)
                                    character_name=multiline_current_character_text
                                    myprint1(f"Add > ch={character_name} dia={speech}")
                                    breakdown.append({"scene_id":current_scene_id,
                                                                            "character_raw":character_name,
                                                                            "line_idx":line_idx,"speech":speech,"type":"SPEECH", "character":character_name })    

                    elif character_mode==    "TIMECODE_HYPHEN_UPPERCASECHARACTER_SPACE_SEMICOLON_DOUBLESPACE_TEXT_NEWLINE_TEXT":
                        if wasChar:
                            speech=trimmed_line
                            myprint1(f"add  {multiline_current_character_text} speech={speech}")
                            breakdown.append({"scene_id":current_scene_id,
                                                                                "character_raw":multiline_current_character_text,
                                                                                "line_idx":line_idx,"speech":speech,"type":"SPEECH", "character":multiline_current_character_text })    
                            wasChar=False
                        else:
                            if is_characterline_TIMECODE_HYPHEN_UPPERCASECHARACTER_SPACE_SEMICOLON_DOUBLESPACE_TEXT_NEWLINE_TEXT(trimmed_line):
                                
                                multiline_current_character_text=extract_character_name_TIMECODE_HYPHEN_UPPERCASECHARACTER_SPACE_SEMICOLON_DOUBLESPACE_TEXT_NEWLINE_TEXT(trimmed_line)
                                myprint1(f"set char=  {multiline_current_character_text}")
                                wasChar=True
                            else:
                                myprint1(f"not a char line")
                            

                    elif character_mode==    "TIMECODE_HYPHEN_TIMECODE_NEWLINE_CHARACTER_SEMICOLON_NEWLINE_DIALOG_NEWLINE":
                        if isEmptyLine:
                            myprint1(f"reset waschar")
                            wasChar=False
                        if wasChar:
                            speech=filter_speech(trimmed_line)
                            myprint1(f"Add line {multiline_current_character_text} {speech} ")
                            breakdown.append({"scene_id":current_scene_id,
                                                                                "character_raw":multiline_current_character_text,
                                                                                "line_idx":line_idx,"speech":speech,"type":"SPEECH", "character":multiline_current_character_text })    
                        if trimmed_line.endswith(":"):
                            multiline_current_character_text=trimmed_line[:-1]
                            multiline_current_character_text=filter_character_name(multiline_current_character_text)     
                            myprint1(f"Set char {multiline_current_character_text} ")
                            wasChar=True
                    elif character_mode==    "TIMECODE_ARROW_TIMECODE_NEWLINE_BRACKETS_CHARACTER_DIALOG_NEWLINE_DIALOG":
                        if isEmptyLine:
                            myprint1(f"reset waschar")
                            wasChar=False
                            wasTimecode=False
                        if wasChar:
                            speech=filter_speech(trimmed_line)
                            myprint1(f"Add line {multiline_current_character_text} {speech} ")
                            breakdown.append({"scene_id":current_scene_id,
                                                                                "character_raw":multiline_current_character_text,
                                                                                "line_idx":line_idx,"speech":speech,"type":"SPEECH", "character":multiline_current_character_text })    
#                        if wasTimecode:
                        if is_text_with_brackets_pattern(trimmed_line):
                            multiline_current_character_text=extract_text_between_brackets(trimmed_line)
                            multiline_current_character_text=filter_character_name(multiline_current_character_text)     
                            myprint1(f"Set char {multiline_current_character_text} ")

                            speech=extract_text_after_brackets(trimmed_line)
                            breakdown.append({"scene_id":current_scene_id,
                                                                                "character_raw":multiline_current_character_text,
                                                                                "line_idx":line_idx,"speech":speech,"type":"SPEECH", "character":multiline_current_character_text })    
#                       
                            wasChar=True
                        else:
                            myprint1(f"No between  ")

                        wasTimecode=is_timecode_arrow_timecode_format(trimmed_line)
                        myprint1(f"wasTimecode {wasTimecode}")


                    elif character_mode==    "LINE_NEWLINE_TIMECODE_ARROW_TIMECODE_NEWLINE_TEXT_ITAG":
                        noccurences=count_pattern_occurrences2(trimmed_line)
                        myprint1("nocc"+str(noccurences))

                        if isEmptyLine:
                            myprint1("resetDialog")
                            wasDialog=False
                        else:
                            if wasDialog:
                                myprint1("has character")
                                
                                if noccurences>0:
                                    for k in m:
                                        character_name=k['character']
                                        speech=k['dialog']
                                        speech=filter_speech(speech)

                                        character_name=filter_character_name(character_name)       
                                        multiline_current_character_text=character_name                              
                                        myprint1(f"Add {character_name} {speech}")
                                        breakdown.append({"scene_id":current_scene_id,
                                                                                "character_raw":character_name,
                                                                                "line_idx":line_idx,"speech":speech,"type":"SPEECH", "character":character_name })    
                                else:
                                    character_name=multiline_current_character_text
                                    speech=trimmed_line
                                    myprint1(f"Add {character_name} {speech}")
                                    breakdown.append({"scene_id":current_scene_id,
                                                                                "character_raw":character_name,
                                                                                "line_idx":line_idx,"speech":speech,"type":"SPEECH", "character":character_name })    

                        if noccurences>0:  
                            m=extract_pattern_occurrences2(trimmed_line)
                            myprint1("nocc"+str(m))
                                                    
                            if len(m)>0:
                                myprint1("aa"+str(m))
                                for k in m:
                                    myprint1(k)
                                    character_name=k['character']
                                    speech=k['dialog']
                                    speech=filter_speech(speech)

                                    character_name=filter_character_name(character_name)       
                                    multiline_current_character_text=character_name                              
                                    myprint1(f"Add {character_name} {speech}")
                                    breakdown.append({"scene_id":current_scene_id,
                                                                            "character_raw":character_name,
                                                                            "line_idx":line_idx,"speech":speech,"type":"SPEECH", "character":character_name })    

                                    wasDialog=True

            wasEmptyLine=isEmptyLine
            line_idx=line_idx+1

    if character_sep_type=="CHARACTER_MODE_MULTILINE":
        if character_mode=="CHARACTER_NEWLINE_DIALOG_NEWLINE_NEWLINE":
            # Handle the case where the file ends while still in pattern
            if multiline_current_character_text and multiline_current_lines_of_dialog:
                myprint1("  > Flush final")
                character_name=filter_character_name(multiline_current_character_text)
                for k in multiline_current_lines_of_dialog:
                    #speech='\n'.join(multiline_current_lines_of_text).strip()       
                    
                    speech=filter_speech(k)           
                    if k=="TRANSLATION":
                            break; 
                        
                    #speech='\n'.join(multiline_current_lines_of_text).strip()          
                    myprint1("  > Add    : "+str(character_name)+" "+speech)
                    breakdown.append({"scene_id":current_scene_id,
                                            "character_raw":character_name,
                                            "line_idx":line_idx,"speech":speech,"type":"SPEECH", "character":character_name })    
                        

    myprint1("breakdown"+str(breakdown))

    all_characters=get_all_characters(breakdown)
    breakdown,replaceList=merge_breakdown_character_talking_to(breakdown,all_characters)
    all_characters=get_all_characters(breakdown)
    #myprint1("all_characters"+str(all_characters))
    #myprint1("replacelist"+str(replaceList))

    replace_map=map_semi_duplicates(all_characters)
   # myprint1("replace_map"+str(replace_map))

    breakdown=merge_breakdown_character_by_replacelist(breakdown,replace_map)
    breakdown=split_AND_character(breakdown)
    all_characters=get_all_characters(breakdown)


    for item in breakdown:
        if item['type']=="SPEECH":
            character_name=item['character']
            #add to character order if new character
            if character_name not in character_order_map:
                character_order_map[character_name]=character_count
                character_count=character_count+1


            if character_name not in character_linecount_map:
                character_linecount_map[character_name]=1
            else:
                character_linecount_map[character_name]=character_linecount_map[character_name]+1

            spoken_text=item['speech']
            le=compute_length(spoken_text,countingMethod)
            if character_name not in character_textlength_map:
                character_textlength_map[character_name]=le
            else:
                character_textlength_map[character_name]=character_textlength_map[character_name]+le


            scene_id=item['scene_id']
            if character_name not in scene_characters_map:
                scene_characters_map[character_name] = set()
            scene_characters_map[character_name].add(scene_id)
            
            #add character to scene if not existing
            if scene_id not in character_scene_map:
                character_scene_map[scene_id] = set()
            character_scene_map[scene_id].add(character_name)

    if len(character_order_map)>0:                                    
        csv_file_path =output_path+script_name+"-comptage-detail.csv"
        data = [
        ]
        for key in character_order_map:
            data.append([
                str(character_order_map[key])+" - "+str(key),str(character_linecount_map[key]),
                str(character_textlength_map[key]),
                str(math.ceil(character_textlength_map[key]/40))])


        myprint1("  > Convert to csv.")
        with open(csv_file_path, mode='w', newline='',encoding=encoding_used) as file:
            writer = csv.writer(file, delimiter=',', quotechar='"', quoting=csv.QUOTE_MINIMAL)
            
            # Write data to the CSV file
            for row in data:
                writer.writerow(row)


        myprint1("  > Convert to xslx.")
        convert_csv_to_xlsx(output_path+script_name+"-comptage-detail.csv",output_path+script_name+"-comptage-detail.xlsx", script_name,encoding_used)
    myprint1("  > Parsing done.")
    success_result=breakdown, character_scene_map,scene_characters_map,character_linecount_map,character_order_map,character_textlength_map

    info ={
        "character_mode":character_mode,
        "scene_separator":scene_separator,
        "character_sep_type":character_sep_type,
        "success":True,
        "success_result":success_result
    }
    return info
    return info,breakdown, character_scene_map,scene_characters_map,character_linecount_map,character_order_map,character_textlength_map


#process_script("scripts/COMPTAGE/20012020.txt","20012020/","20012020","ALL")    
#process_script("scripts/examples/YOU CAN'T RUN FOREVER_SCRIPT_VO.txt","YOUCANRUNFOREVER_SCRIPT_VOe/","YOU CAN'T RUN FOREVER_SCRIPT_VO","ALL")    
#process_script("190421-1.txt","190421-1/","190421-1.txt")
#process_script("scripts/examples/LATENCY.docx","LATENCY/","LATENCY","ALL")    
#process_script("scripts/examples/LATENCY.docx","LATENCY/","LATENCY","ALL")    
#process_script("scripts/examples/Gods of the deep - CCSL.docx","GODS/","GODS","ALL")    
#process_script("scripts/examples/Blackwater Lane.docx","Blackwater Lane/","Blackwater Lane","ALL")    

            