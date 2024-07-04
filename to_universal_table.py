from pdfplumber.pdf import PDF
import pdfplumber          
import pandas as pd
from docx import Document
import logging                                                   

def myprint7(s):
    s="ui: "+str(s)
    logging.debug(s)
    #print(s.encode('utf-8'))

def convert_pdf_to_universaltables(file_path):
        res=[]
        lastTable=None
        with pdfplumber.open(file_path) as pdf:
            for page_number, page in enumerate(pdf.pages, start=1):
                myprint7("----\nreadpdf page"+str(page_number))
                tables = page.extract_tables()
                
                for table_number, page_table in enumerate(tables, start=1):
                    
                    row_count = len(page_table)
                    column_count = len(page_table[0]) if row_count > 0 else 0
                    
                    myprint7(f"Page {page_number} Table {table_number} Size: {row_count} x {column_count}")
                    
                    page_cells=[]
                    for row_index, row in enumerate(page_table):
                        #myprint7(f"Add row : {row_index} {row}")
                        resrow=[]
                        for col_index, cell in enumerate(row):
                            #myprint7(f"myCell({row_index}, {col_index}): {cell}")
                            if cell!=None:
                                resrow.append(cell)                               
                            else:
                                resrow.append("")
                        page_cells.append(resrow)
                    page_table={
                        "row_count":row_count,
                        "col_count":column_count,
                        "cells":page_cells
                    }
                         
                    if lastTable!=None  and column_count==lastTable['col_count']:
                        myprint7(f"append")
                        rc=lastTable["row_count"]
                        prc=page_table["row_count"]
                        myprint7(f"add lastrc= {rc} to prc={prc} ")
                        lastTable["cells"].extend(page_cells)     
                        lastTable["row_count"]=rc+prc
                        nrc=lastTable["row_count"]
                        myprint7(f"after newrc= {nrc} ")
                    else:
                        #flush last
                        if lastTable!=None:
                            nrc=lastTable["row_count"]
                            myprint7(f"ADD TO RES LASTTABLE= {nrc} ")
                            res.append(lastTable)
                            lastTable=None
                        prc=page_table["row_count"]
 #                       myprint7(f"ADD TO RES CURRENT {prc}")
#                        res.append(page_table)
                    if lastTable==None:
                        lastTable=page_table
   
        res.append(lastTable)
   
        myprint7(f"RES: {res}")
        
        for idx,k in enumerate(res):
            myprint7("------------- "+str(idx)+" ------------------")
            rc=k["row_count"]
            cc=k["col_count"]
            cells=k["cells"]
            myprint7(f"total Row count: {rc}")
            myprint7(f"total Column count: {cc}")
            myprint7(f"total Cells: {cells}")
        return res





def convert_docx_to_universaltables(file_path):
        myprint7("convert_docx_to_universaltables")
        res=[]
        table_dimensions=[]
        doc = Document(file_path)
        for table in doc.tables:
            num_rows = len(table.rows)
            num_columns = len(table.columns)
            table_dimensions.append((num_rows, num_columns))


        for k,table in enumerate(doc.tables):
                rescells=[]
                row_count = len(table.rows)
                column_count = len(table.columns)
                # Get the first table   
                for row_index,row in enumerate(table.rows):
                    resrow=[]
                    for col_index,cell in enumerate(row.cells):
                        t=cell.text.strip()
                        resrow.append(t)                                
                        #myprint7(f"myCell({row_index}, {col_index}): {t}")
                    rescells.append(resrow)
                table={
                    "row_count":row_count,
                    "col_count":column_count,
                    "cells":rescells
                }
                res.append(table)
        return res

def access_cell(df,i, j):
    try:
        cell_value = df.iloc[i, j]
        
        return cell_value
    except IndexError:
        myprint7(f"ERR Index out of range {i} {j}")
        return "Index out of range"

def convert_xlsx_to_universaltables(file_path):
        res=[]
        table_dimensions=[]
        df = pd.read_excel(file_path,header=None)



        rescells=[]
        row_count = df.shape[0]
        column_count = df.shape[1]
        myprint7(f"convert_xlsx_to_universaltables {row_count}  x {column_count}")

        # Get the first table   
        for row_index in range(0,row_count):
            row_data = df.iloc[row_index, :]
            myprint7(f"convert_xlsx_to_universaltables {row_index}")

            resrow=[]
            # Loop over the cells of the row
            for col_index, cell in enumerate(row_data):   
                t=access_cell(df,row_index, col_index)
                is_nan = pd.isna(df.at[row_index, col_index])
                if is_nan:
                    t=""
                myprint7(f"convert_xlsx_to_universaltables {col_index} : {t}")
                resrow.append(t)                                
                #myprint7(f"myCell({row_index}, {col_index}): {cell}")
            rescells.append(resrow)
        table={
            "row_count":row_count,
            "col_count":column_count,
            "cells":rescells
        }
        res.append(table)
        return res