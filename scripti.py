import os
import tkinter as tk
from tkinter import ttk, filedialog, Text,Menu,Toplevel,Scrollbar, Scale, HORIZONTAL, VERTICAL
from script_parser import process_script,get_pdf_page_blocks,detect_word_table,run_convert_pdf_to_txt,split_elements, get_pdf_text_elements, is_supported_extension,convert_word_to_txt,convert_xlsx_to_txt,convert_rtf_to_txt,convert_pdf_to_txt,filter_speech
from utils import get_intial_treeview_folder_path,get_setting_ini_path,copy_folder_contents,get_excel_column_name,help_word_table,help_pdf_text,help_merge,make_dpi_aware,detect_file_encoding,get_os,convert_csv_to_xlsx,get_encoding,get_log_file_path,get_temp_folder_path,get_recentfiles_file_path,save_string_to_file
import pandas as pd

import io
import tkinter.font as tkFont
import subprocess
#import pypdfium2 as pdfium
import traceback
import tkinter.font as tkfont
from tkinter import Canvas, PhotoImage

#import pytesseract
import platform 
import re
from docx import Document
import sys
import csv
# from PyPDF2 import PdfWriter, PdfReader
from pdfplumber.pdf import PDF
import pdfplumber                                                             
import math
import webbrowser
import time
import logging
from tkinter import ttk, messagebox
from PIL import Image, ImageTk
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.converter import TextConverter, PDFPageAggregator
from pdfminer.layout import LAParams
from pdfminer.pdfpage import PDFPage
from pdfminer.layout import LTImage
import threading
if False:
    class UTFStreamHandler(object):
        def __init__(self, stream):
            self.stream = stream

        def write(self, data):
            if not isinstance(data, str):
                data = str(data)
            self.stream.buffer.write(data.encode('utf-8'))
            self.stream.buffer.flush()

        def flush(self):
            self.stream.flush()

    # Only replace sys.stdout if it's not already wrapped
    if not isinstance(sys.stdout, io.TextIOWrapper) or sys.stdout.encoding.lower() != 'utf-8':
        sys.stdout = UTFStreamHandler(sys.stdout)


#sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

app_log_path=get_log_file_path()
logging.basicConfig(filename=app_log_path,encoding='utf-8',level=logging.DEBUG)
logging.debug("Script starting...")


last_row_id = None

countingMethods=[
#    "LINE_COUNT",
 #   "WORD_COUNT",
    "ALL",
    "BLOCKS_50",
    #"BLOCKS_40",
   # "ALL_NOSPACE",
   # "ALL_NOPUNC",
  #  "ALL_NOSPACE_NOPUNC",
 #   "ALL_NOAPOS",    
]

countingMethodNames={
  #  "LINE_COUNT":"Lines",
   # "WORD_COUNT":"Words",
    "ALL":"Caractères",
    "BLOCKS_50":"Répliques",
#    "BLOCKS_40":"Blocks (40)",
#    "ALL_NOSPACE":"No space",
 #   "ALL_NOPUNC":"No punctuation",
  #  "ALL_NOSPACE_NOPUNC":"No space, no punctuation",
   # "ALL_NOAPOS":"No apostrophe",
}

countingMethod="ALL"
currentOutputFolder=""
currentFilePath=""
currentScriptFilename=""
currentConvertedFilePath=""
do_debug=True

RECENT_FILES_PATH = get_recentfiles_file_path()
MAX_RECENT_FILES = 10

aoutputFolder=get_temp_folder_path()
currentRightclickRowId=None
currentXlsxPath=""
currentDialogPath=""
currentTimelinePath=""
currentHasImportTableTab=True
currentBreakdown=None
currentFig=None
currentCanvas=None
currentPDFPages=None
currentPDFPageIdx=None
currentCharacterMergeFromName=None
currentCharacterSelectRowId=None
currentCharacterMultiSelectRowIds=None
currentDisabledCharacterNames=[]
currentResultCharacterOrderMap=None
currentResultEnc=""
currentResultName=""
currentResultLinecountMap=None
currentResultSceneCharacterMap=None
currentMergePopupTable=None
currentMergedCharacters={}
currentMergedCharactersTo={}
currentMergePopupWindow=None
currentBlockSize=50
currentCharactersSelectedRowIds=[]
currentConvertedFileIgnoreBeginning=0
currentConvertedFileIgnoreEnd=0
def myprint2(s):
    logging.debug(s)
#    print(s)
def myprint7(s):
    s="ui: "+str(s)
    logging.debug(s)
    #print(s.encode('utf-8'))
def myprint7a(s):
    s="ui: "+str(s)
    logging.debug(s)
    print(s.encode('utf-8'))
def myprint4(s):
    logging.debug(s)
 #   print(s)
def myprint5(s):
    logging.debug(s)
  #  print(s)

myprint7a(f"Tmp                : {aoutputFolder}")
myprint7a(f"Recent files       : {RECENT_FILES_PATH}")
myprint7a(f"App log            : {app_log_path}")

# Only call this function if your application is running on Windows
if sys.platform.startswith('win32'):
    make_dpi_aware()
myprint7a(f"Make folder            : {aoutputFolder}")

if not os.path.exists(aoutputFolder):
    os.mkdir(aoutputFolder)
myprint7a(f"Make folder   done         : {aoutputFolder}")

def load_recent_files():
    if os.path.exists(RECENT_FILES_PATH):
        with open(RECENT_FILES_PATH, 'r') as f:
            return [line.strip() for line in f.readlines()]
    return []
myprint7a(f"Make 1")
recent_files = load_recent_files()
myprint7a(f"Make 2")


def save_recent_files(recent_files):
    with open(RECENT_FILES_PATH, 'w') as f:
        for file in recent_files:
            f.write(file + '\n')


def update_recent_files(file_path):
    if file_path in recent_files:
        recent_files.remove(file_path)
    recent_files.insert(0, file_path)
    if len(recent_files) > MAX_RECENT_FILES:
        recent_files.pop()
    save_recent_files(recent_files)
    update_recent_files_menu()

def compute_length_by_method(line,method):
    global currentBlockSize
    res=0
    if method=="ALL":
        res=len(line)
    elif method=="LINE_COUNT":
        res=1
    elif method=="ALL_NOSPACE":
        res= len(line.replace(" ",""))
    elif method=="ALL_NOPUNC":
        res= len(line.replace(",","").replace("?","").replace(".","").replace("!",""))
    elif method=="ALL_NOSPACE_NOPUNC":
        res= len(line.replace(" ","").replace(",","").replace("?","").replace(".","").replace("!",""))
    elif method=="ALL_NOAPOS":
        res= len(line.replace("'",""))
    elif method=="WORD_COUNT":
        res= len(line.split(" "))
    elif method=="BLOCKS_50":
        res= len(line)/currentBlockSize
    elif method=="BLOCKS_40":
        res= len(line)/40
    else:
        res= -1
    #print("compute_length_by_method METHOD "+method+" "+str(res))
    return res


def load_tree(parent, root_path):
    print(f"load_tree {root_path}")
    os.makedirs(root_path, exist_ok=True)
    print(f"load_tree created if not {root_path}")

    # Clear the tree view if root_path is the starting directory
    if parent == "":
        folders.delete(*folders.get_children())
        parent = folders.insert('', 'end', text=os.path.basename(root_path), open=True, values=["","",])

    # List all entries in the directory
    try:
        dir_entries = os.listdir(root_path)
    except PermissionError:
        return  # Skip directories without permission

    # Separate and sort directories and files
    dirs = sorted([entry for entry in dir_entries if os.path.isdir(os.path.join(root_path, entry))])
    files = sorted([entry for entry in dir_entries if not os.path.isdir(os.path.join(root_path, entry))])

    # Insert directories first
    for entry in dirs:
        entry_path = os.path.join(root_path, entry)

        dir_id = folders.insert(parent, 'end', text=" "+entry, image=folder_icon, open=False, values=[entry_path,"Dossier",],tags=('folder'))
        #folders.insert(dir_id, 'end', text=os.path.basename(entry_path), open=True, values=[entry_path])
        folders.insert(dir_id, 'end', text="Loading...", values=["dummy"])  # Dummy node

        #load_tree(dir_id, entry_path)  # Recursively load subdirectories

    # Insert files
    for entry in files:
        entry_path = os.path.join(root_path, entry)
        supported_tag="not_supported"
        
        name, extension = os.path.splitext(entry)
        extension_without_dot=extension
        if extension.startswith("."):
            extension_without_dot=extension_without_dot[1:]
            
        if is_supported_extension(extension):
            supported_tag="supported" 
            if extension==".docx" or extension==".doc":
                folders.insert(parent, 'end', text=" "+entry,image=docx_icon, values=[entry_path, extension_without_dot,],tags=(supported_tag))
            else:
                folders.insert(parent, 'end', text=" "+entry,image=txt_icon, values=[entry_path, extension_without_dot,],tags=(supported_tag))
        else:        
            folders.insert(parent, 'end', text=" "+entry, values=[entry_path, extension_without_dot,],tags=(supported_tag))

def on_motion(event):
    # Identify the row on which the mouse is currently hovering
    row_id = folders.identify_row(event.y)
    if row_id:
        # Retrieve current tags and add 'hover' tag
        current_tags = set(folders.item(row_id, 'tags'))
        
        current_tags.add('hover')
        folders.item(row_id, tags=list(current_tags))

    # Reset the background color of previously hovered rows
    global last_row_id
    if last_row_id and last_row_id != row_id:
        current_tags = set(folders.item(last_row_id, 'tags'))
        current_tags.discard('hover')  # Remove the hover tag
        folders.item(last_row_id, tags=list(current_tags))
    last_row_id = row_id

def on_leave(event):
    # When the mouse leaves the Treeview, reset the background of the last hovered row
    global last_row_id
    if last_row_id:
        current_tags = set(folders.item(last_row_id, 'tags'))
        current_tags.discard('hover')  # Remove the hover tag
        folders.item(last_row_id, tags=list(current_tags))
    last_row_id = None


def reset_tables(): 
    myprint7("reset_tables")
    
    #disable_merge_button()
    for item in breakdown_table.get_children():
        breakdown_table.delete(item)
    for item in character_list_table.get_children():
        character_list_table.delete(item)
    for item in character_table.get_children():
        character_table.delete(item)
    for item in stats_table.get_children():
        stats_table.delete(item)
    for item in character_stats_table.get_children():
        character_stats_table.delete(item)
    
def open_recent_file(file_path):
    runJob(file_path,"ALL")

def update_recent_files_menu():
    recent_files_menu.delete(0, tk.END)
    for file_path in recent_files:
        recent_files_menu.add_command(label=file_path, command=lambda path=file_path: open_recent_file(path))

def reset_tabs():
    if importTab != None:
        importTab.destroy()    
def runJobPreprocessing(file_path,enc,params={}):
    myprint7("runJobPreprocessing")
    global currentFilePath
    global currentScriptFilename
    global currentBreakdown
    global currentOutputFolder
    global currentTimelinePath

    global currentResultCharacterOrderMap
    global currentResultEnc
    global currentResultName
    global currentResultLinecountMap
    global currentResultSceneCharacterMap
    global importTab
    global currentPDFPages
    global currentPDFPageIdx

    file_name = os.path.basename(file_path)
    currentScriptFilename=file_name
    name, extension = os.path.splitext(file_name)
    #DOCX
    if extension==".docx" or extension==".doc":
        myprint7("runJobPreprocessing > Conversion Word to txt")
        #if importTab != None:
         #   importTab.destroy()       
            #importTab.reset(file_path)
        myprint7("runJobPreprocessing> open"+str(file_path))
        doc = Document(file_path)
        myprint7("opened")
        forceMode=""
        forceCols={}

        if 'param_type' in params and params['param_type']=="WORD":
            if 'character' in params and 'dialog' in params:
                char=params['character']
                dial=params['dialog']
                if char!=dial:
                    forceMode="DETECT_CHARACTER_DIALOG"
                    forceCols={
                        "CHARACTER":char,
                        "DIALOG":dial
                    }    
        myprint7("runJobPreprocessing nbTables="+str(len(doc.tables)))

        if len(doc.tables) > 0:
            myprint7("has table, show_importtable_tab")
            if len(params)==0:
                importTab=WordTableColumnSelector(tab_import,file_path)
#            show_importtable_tab()
        else:
            myprint7("no table, hide_importtable_tab")
 #           hide_importtable_tab()

        converted_file_path=convert_word_to_txt(file_path,os.path.abspath(currentOutputFolder),forceMode=forceMode,forceCols=forceCols)
        if len(converted_file_path)==0:
            myprint7("ui Conversion docx to txt failed")
            myprint7("ui Failed")
            hide_loading()
            return 
        file_path=converted_file_path
        myprint7("ui Converted file path :"+file_path)
    
    #XLSX
    if extension==".xlsx":
        myprint7("Conversion Excel to txt")

        forceMode=""
        forceCols={}
        
            
        converted_file_path=convert_xlsx_to_txt(file_path,os.path.abspath(currentOutputFolder),forceMode=forceMode,forceCols=forceCols)
        if len(converted_file_path)==0:
            myprint7("Conversion xlsx to txt failed")
            myprint7("Failed")
            hide_loading()
            return 
        file_path=converted_file_path
        myprint7("Conversion file path :"+file_path)
        
    #RTF
    if extension==".rtf":
        myprint7("Conversion RTF to txt")
        converted_file_path,txt_encoding=convert_rtf_to_txt(file_path,os.path.abspath(currentOutputFolder),enc)
        if len(converted_file_path)==0:
            myprint7("Conversion rtf to txt failed")
            myprint7("Failed")
            hide_loading()
            return
        enc=txt_encoding
        file_path=converted_file_path
        myprint7("Conversion file path :"+file_path)
    
    #PDF
    if extension==".pdf":
        with open(file_path, 'rb') as file:
            currentPDFPageIdx=10
            currentPDFPages = list(PDFPage.get_pages(file))
            pdf_viewer = PDFViewer(tab_import_pdf, file_path,os.path.abspath(currentOutputFolder),enc)
            return; 
        myprint7("Conversion PDF to txt")
        converted_file_path,txt_encoding=convert_pdf_to_txt(file_path,os.path.abspath(currentOutputFolder),enc)
        if len(converted_file_path)==0:
            myprint7("Conversion pdf to txt failed")
            myprint7("Failed")
            hide_loading()
            return
        enc=txt_encoding
        file_path=converted_file_path
        myprint7("Conversion file path :"+file_path)
    myprint7("runJobPreprocessing done")
    return file_path

def runJob(file_path,method,params={}):
    global currentFilePath
    global currentConvertedFilePath
    global currentScriptFilename
    global currentBreakdown
    global currentOutputFolder
    global currentTimelinePath
    global currentResultCharacterOrderMap
    global currentResultEnc
    global currentResultName
    global currentResultLinecountMap
    global currentResultSceneCharacterMap
    global importTab
    global currentConvertedFileIgnoreBeginning
    global currentConvertedFileIgnoreEnd
    global currentPDFPages
    global currentPDFPageIdx
    myprint7("------------------- JOB ---------------------------")
    myprint7("File path               : "+str(file_path))
    myprint7("Params               : "+str(params))
    update_recent_files(file_path)
    currentFilePath=file_path
    currentConvertedFilePath=file_path
    reset_tables()
    # Check if the selected item is a file and display its content
    if os.path.isfile(file_path):
        try:
            file_name = os.path.basename(file_path)
            currentScriptFilename=file_name
            name, extension = os.path.splitext(file_name)
            app.title(f"Scripti - {file_name}")
            myprint7("Name                : "+name)
            myprint7("Extension           : "+extension)
            if is_supported_extension(extension):
                myprint7("Supported           : YES")
                encoding_info = detect_file_encoding(file_path)
                encoding=encoding_info['encoding']
                myprint7("Encoding detection   : "+str(encoding_info))
                myprint7("Encoding detected   : "+str(encoding))
                myprint7("Encoding confidence : "+str(encoding_info['confidence']))
                enc=get_encoding(encoding)
                myprint7("Encoding used       : "+str(enc))

                currentOutputFolder=aoutputFolder+"/"+name+"/"
                if not os.path.exists(currentOutputFolder):
                    os.mkdir(currentOutputFolder)
                extension=extension.lower()

                file_path=runJobPreprocessing(file_path,enc,params)
                if file_path==None:
                    #failed
                    return
                currentConvertedFilePath=file_path
                myprint7(" > Opening "+file_path)
                with open(file_path, 'r', encoding=enc) as file:
                    myprint7(" > Opened")
                    content = file.read()
                    myprint7(" > Read")
                    file_preview.text_widget.delete(1.0, tk.END)
                    file_preview.text_widget.insert(tk.END, content)
                    
                myprint7(f" > Process, ignoring {currentConvertedFileIgnoreBeginning}  {currentConvertedFileIgnoreEnd}")
                res=            process_script(file_path,currentOutputFolder,name,method,enc,"",{},currentConvertedFileIgnoreBeginning,currentConvertedFileIgnoreEnd)
                if res==None:
                    return
                
                breakdown,character_scene_map,scene_characters_map,character_linecount_map,character_order_map,character_textlength_map=res
                myprint7(" > Processed")

                if breakdown==None:
                    myprint7(" > Failed")
                    hide_loading()
                else:
                    myprint7(" > OK")
                    currentBreakdown=breakdown

                    png_output_file=currentOutputFolder+name+"_timeline.png"
                    currentTimelinePath=png_output_file
                    currentResultCharacterOrderMap=character_order_map
                    currentResultEnc=enc
                    currentResultName=name
                    currentResultLinecountMap=character_linecount_map
                    currentResultSceneCharacterMap=scene_characters_map
                    postProcess(breakdown,character_order_map,enc,name,character_linecount_map,scene_characters_map,png_output_file)
                        
            else:
                myprint7(" > Not supported")
                #stats_label.config(text=f"Format {extension} not supported")
            hide_loading()

        except Exception as e:
            myprint7(f"Error opening file: {e} tried with encoding={ enc}")
            file_preview.text_widget.delete(1.0, tk.END)
            file_preview.text_widget.insert(tk.END, f"Error opening file: {e} tried with encoding{ enc} {traceback.format_exc()}")
            hide_loading()
def append_file_content(source_file, destination_file):
    try:
        # Open the source file in read mode
        with open(source_file, 'r') as file1:
            # Read the content of the source file
            content = file1.read()

        # Open the destination file in append mode
        with open(destination_file, 'a') as file2:
            # Append the content to the destination file
            file2.write(content)

        myprint7(f"Content of {source_file} appended to {destination_file} successfully.")
    except FileNotFoundError:
        myprint7(f"File not found: {source_file} or {destination_file}")
    except IOError:
        myprint7(f"An error occurred while appending content to {destination_file}")

def runGroupJob(file_paths,method):
    global currentFilePath
    global currentScriptFilename
    global currentBreakdown
    global currentOutputFolder
    global currentTimelinePath
    global currentResultCharacterOrderMap
    global currentResultEnc
    global currentResultName
    global currentResultLinecountMap
    global currentResultSceneCharacterMap
    global importTab
    global currentPDFPages
    global currentPDFPageIdx

    myprint7("runGroupJob > file_paths="+str(file_paths))
    first_path=file_paths[0]
    firstname, firstextension = os.path.splitext(first_path)
    first_enc=None
    grouped_path=firstname+"-grouped.txt"

#    update_recent_files(file_path)
    for file_path in file_paths:
        myprint7("runGroupJob > filepath="+str(file_path))
        currentFilePath=file_path
        reset_tables()
        # Check if the selected item is a file and display its content
        if os.path.isfile(file_path):
            try:
                file_name = os.path.basename(file_path)
                currentScriptFilename=file_name
                name, extension = os.path.splitext(file_name)
                app.title(f"Scripti - {file_name}")
                
                myprint7("Name                : "+name)
                myprint7("Extension           : "+extension)
                if is_supported_extension(extension):
                    myprint7("Supported           : YES")
                    encoding_info = detect_file_encoding(file_path)
                    encoding=encoding_info['encoding']
                    myprint7("Encoding detected   : "+str(encoding))
                    myprint7("Encoding confidence : "+str(encoding_info['confidence']))
                    enc=get_encoding(encoding)
                    myprint7("Encoding used       : "+str(enc))
                    if first_enc==None:
                        first_enc=enc
                    currentOutputFolder=aoutputFolder+"/"+name+"/"
                    if not os.path.exists(currentOutputFolder):
                        os.mkdir(currentOutputFolder)
                    extension=extension.lower()

                    file_path=runJobPreprocessing(file_path,enc)
                    myprint7(" > Opening "+file_path)
                    append_file_content(file_path,grouped_path)
                  
                else:
                    myprint7(" > Not supported")
                    #stats_label.config(text=f"Format {extension} not supported")
                    hide_loading()

            except Exception as e:
                file_preview.text_widget.delete(1.0, tk.END)
                file_preview.text_widget.insert(tk.END, f"Error opening file: {e} tried with encoding{ enc}")
                hide_loading()

    try:
        enc=first_enc
        file_path=grouped_path
        myprint7(" > Opening "+file_path)
        with open(file_path, 'r', encoding=enc) as file:
            myprint7(" > Opened")
            content = file.read()
            myprint7(" > Read")
            file_preview.text_widget.delete(1.0, tk.END)
            file_preview.text_widget.insert(tk.END, content)
            
        myprint7(" > Process")
        breakdown,character_scene_map,scene_characters_map,character_linecount_map,character_order_map,character_textlength_map=process_script(file_path,currentOutputFolder,name,method,enc)
        myprint7(" > Processed")

        if breakdown==None:
            myprint7(" > Failed")
            hide_loading()
        else:
            myprint7(" > OK")
            currentBreakdown=breakdown

            png_output_file=currentOutputFolder+name+"_timeline.png"
            currentTimelinePath=png_output_file
            currentResultCharacterOrderMap=character_order_map
            currentResultEnc=enc
            currentResultName=name
            currentResultLinecountMap=character_linecount_map
            currentResultSceneCharacterMap=scene_characters_map
            postProcess(breakdown,character_order_map,enc,name,character_linecount_map,scene_characters_map,png_output_file)
            

    except Exception as e:
        file_preview.text_widget.delete(1.0, tk.END)
        file_preview.text_widget.insert(tk.END, f"Error opening file: {e} tried with encoding{ enc}")
        hide_loading()

def postProcess(breakdown,character_order_map,enc,name,character_linecount_map,scene_characters_map,png_output_file):

    global currentResultCharacterOrderMap
    myprint7(f"Postprocess nbreakdown={len(breakdown)}")
    myprint7(f"Postprocess ncharacter_order_map={len(character_order_map)}")
    fill_breakdown_table(breakdown)
    save_dialog_csv(breakdown,enc,"")
    fill_stats_table(breakdown)
    if len(character_order_map)>0:
        fill_character_list_table(character_order_map)
        fill_character_stats_table(character_order_map,breakdown,enc)
        fill_character_table(character_order_map, breakdown,character_linecount_map,scene_characters_map)
        for char in character_order_map:
            save_dialog_csv(breakdown,enc,char)
    

def trim_filename_if_too_long(file_path, max_length=255):
    # Extract directory, filename, and extension
    directory, filename = os.path.split(file_path)
    name, ext = os.path.splitext(filename)
    
    # Determine the maximum allowable length for the filename without extension
    max_name_length = max_length - len(ext)
    
    # If the filename is too long, trim it and add '-cropped'
    if len(name) > max_name_length:
        cropped_name = name[:max_name_length - len('-cropped')] + '-cropped'
        new_filename = cropped_name + ext
    else:
        new_filename = filename
    
    # Reconstruct the file path
    new_file_path = os.path.join(directory, new_filename)
    return new_file_path
def save_dialog_csv(breakdown,enc,char):
    global currentDialogPath
    haschar=char!=None and len(char)>0
    totalcsvpath=currentOutputFolder+"/"+currentScriptFilename+"-dialog.csv"
    if haschar:
        if not os.path.exists(currentOutputFolder+"dialogs/"):
            os.mkdir(currentOutputFolder+"dialogs/")
        safechar=char.replace("/","_")
        totalcsvpath=currentOutputFolder+"dialogs/"+currentScriptFilename+"-dialog-"+safechar+".csv"

    currentDialogPath=totalcsvpath
    
    data=[]
    for item in breakdown:
        line_idx=item['line_idx']
        type_=item['type']
        if(type_=="SPEECH"):
            speech=item['speech']
            character=item['character']
            character_raw=item['character_raw']
            if not haschar or (haschar and character==char):
                filtered_speech=filter_speech(speech)
                datarow=[str(line_idx),character,filtered_speech];
                for m in countingMethods: 
                            #myprint4("add"+str(m))
                            le=compute_length_by_method(filtered_speech,m)
                            datarow.append(le)
                    
                data.append(datarow)
    
            
    myprint7("Saving dialog csv encoding="+enc)
    myprint7("Saving dialog csv path="+totalcsvpath)
    #myprint7("data"+str(data))
    totalcsvpath=trim_filename_if_too_long(totalcsvpath,150)
    with open(totalcsvpath, mode='w', newline='',encoding=enc) as file:
        writer = csv.writer(file, delimiter=',', quotechar='"', quoting=csv.QUOTE_MINIMAL)
        
        # Write data to the CSV file
        for row in data:
            #myprint7("Write "+str(row))
            writer.writerow(row)
    xlsxpath=totalcsvpath.replace(".csv",".xlsx")
    myprint7("xlsx"+xlsxpath)
    if len(data)>0:
        convert_dialog_csv_to_xlsx2(totalcsvpath,xlsxpath,enc)

def on_treeview_folder_select(event):
    global currentScriptFilename
    global currentOutputFolder
    logging.debug("on_folder_select")
    #myprint7("FOLDER SELECT")
    selected_item = folders.selection()[0]
    file_path = folders.item(selected_item, 'values')[0]
    if os.path.isfile(file_path):
        logging.debug(">>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>")
        currentScriptFilename=file_path
        loading_label_txt.config(text="Analyse de "+file_path)
        show_loading()
        app.update_idletasks()  # Force the UI to update
        #for item in disabled_character_list_table.get_children():
         #   disabled_character_list_table.delete(item)
        threading.Thread(target=runJob,args=(file_path,countingMethod)).start()


# Function to remove all items
def remove_all_tree_items():
    for item in folders.get_children():
        folders.delete(item)


def open_folder():
    myprint7("Change folder")
    remove_all_tree_items()
    directory = filedialog.askdirectory(initialdir=os.getcwd())
    if directory:
        update_ini_settings_file("SCRIPT_FOLDER",directory)
        load_tree("",directory)

def open_script():
    myprint7("Open script")
    file_path = filedialog.askopenfilename()
    if file_path:
        myprint7(f"Selected file: {file_path}")
        runJob(file_path,"ALL")

def open_script_group():
    myprint7("Open script group")
    file_paths = filedialog.askopenfilenames()
    sorted_file_paths = sorted(file_paths)
    runGroupJob(sorted_file_paths,"ALL")

def open_folder_firsttime():
    myprint7("Change folder")
    directory = filedialog.askdirectory(initialdir=os.getcwd(),title="Choose your scripts folder")
    if directory:
        update_ini_settings_file("SCRIPT_FOLDER",directory)


def exit_app():
    app.quit()


def center_window():
    app.update_idletasks()
    width = app.winfo_width()
    height = app.winfo_height()
    screen_width = app.winfo_screenwidth()
    screen_height = app.winfo_screenheight()
     # Calculate width and height as 80% of screen dimensions
    width = int(screen_width * 0.9)
    height = int(screen_height * 0.8)
    x = int((screen_width - width) / 2)
    y = 20#int((screen_height - height) / 2)
    app.geometry(f'{width}x{height}+{x}+{y}')

def center_panel(panel):
    panel.update_idletasks()
    width = panel.winfo_width()
    height = panel.winfo_height()
    screen_width = panel.winfo_screenwidth()
    screen_height = panel.winfo_screenheight()
     # Calculate width and height as 80% of screen dimensions
    width = int(screen_width * 0.8)
    height = int(screen_height * 0.7)
    x = int((screen_width - width) / 2)
    y = 20#int((screen_height - height) / 2)
    panel.geometry(f'{width}x{height}+{x}+{y}')


def reduce_width():
    current_width = app.winfo_width()
    current_height = app.winfo_height()
    # Reduce width by 1 pixel
    new_width = current_width - 1
    # Set the new geometry
    app.geometry(f"{new_width}x{current_height}")


def stats_per_character(breakdown,character_name):
    global currentBlockSize
    line_count=0
    word_count=0
    character_count=0
    
    for item in breakdown:
        if item['type']=="SPEECH":
            if item['character']==character_name:
                t=item['speech']
                filtered_speech=filter_speech(t)
                line_count=line_count+1
                le=compute_length_by_method(filtered_speech,"ALL")
                character_count=character_count+le
                word_count=word_count+len(t.split(" "))
    #myprint7(character_name,character_count)
            
    replica_count=math.ceil(character_count/currentBlockSize)
    return line_count,word_count,character_count,replica_count

def fill_character_table(character_order_map, breakdown,character_linecount_map,scene_characters_map):
    global currentBlockSize
    order_idx=0
    for item in character_order_map:
        lines=character_linecount_map[item]
        line_count,word_count,character_count,replica_count=stats_per_character(breakdown,item)
        scenes=scene_characters_map[item]
        scenes=', '.join(scenes)
        status="VISIBLE"
        if item in currentDisabledCharacterNames:
            status="MASQUÉ"
            character_table.insert('','end',values=(" - ",item,status,str(character_count),str(math.ceil(character_count/currentBlockSize)),scenes),tags=("hidden"))
        elif item in currentMergedCharacters:
            status="FUSIONNÉ (avec "+str(currentMergedCharacters[item])+")"
            character_table.insert('','end',values=(" - ",item,status,str(character_count),str(math.ceil(character_count/currentBlockSize)),scenes),tags=("hidden"))
        else:
            order_idx=order_idx+1
            character_table.insert('','end',values=(str(order_idx),item,status,str(character_count),str(math.ceil(character_count/currentBlockSize)),scenes))
        
        

def compute_length(method,line):
    if method=="ALL":
        return len(line);
    return len(line);

def fill_character_list_table(character_order_map):
    #myprint7("fill_character_list_table")

    for character_name in character_order_map:
        #myprint7("CHAR add"+character_name)
        if (not character_name in currentDisabledCharacterNames) or (not character_name in currentMergedCharacters):
            character_named = character_name 
            #myprint4("CHAR add"+character_named)
            character_list_table.insert('','end',values=(character_named,))
        

def fill_character_stats_table(character_order_map, breakdown,encoding_used):
    myprint4("fill_character_stats_table")

    total_by_character_by_method={}
    for character_name in character_order_map:
        #myprint4("CHAR add"+character_name)

        character_named = character_name 
        #myprint4("CHAR add"+character_named)
    #    character_list_table.insert('','end',values=(character_named,))
        
        
        character_order=character_order_map[character_name]
        #myprint4("CHAR"+str(character_name))
        rowtotal=("-",character_name,"-","TOTAL")
        total_by_method={}
        for m in countingMethods:
            total_by_method[m]=0
        
        for item in breakdown:
            line_idx=item['line_idx']
            type_=item['type']
            if(type_=="SPEECH"):
                speech=item['speech']
                character=item['character']
                character_raw=item['character_raw']
                
                filtered_speech=filter_speech(speech)

                if character==character_name:
                    #myprint4("    MATCH"+str(speech))

                    row=(str(line_idx),character,character_raw, speech)
                    for m in countingMethods: 
                        #myprint4("add"+str(m))
                        le=compute_length_by_method(filtered_speech,m)
                        row=row+(str(le),)
                        total_by_method[m]=total_by_method[m]+le
                    #myprint4("add"+str(row))
                    character_stats_table.insert('','end',values=row)
        
        #round for BLOCKS
        for m in countingMethods:
            if m.startswith("BLOCKS"):
                total_by_method[m]=math.ceil(total_by_method[m])

        for m in countingMethods:
            rowtotal=rowtotal+(total_by_method[m],)
        character_stats_table.insert('','end',values=rowtotal,tags=['total'])
        
        total_by_character_by_method[character_name]=total_by_method
#        total_by_character_by_method[str(character_order)+" - "+character_name]=total_by_method
    totalcsvpath=currentOutputFolder+"/"+currentScriptFilename+"-comptage.csv"
    generate_total_csv(total_by_character_by_method,totalcsvpath,encoding_used,character_order_map)
    for item in character_stats_table.get_children():
        character_stats_table.delete(item)



def convert_csv_to_xlsx2(csv_file_path, xlsx_file_path, n,encoding_used):
    myprint4("convert_csv_to_xlsx2 "+csv_file_path+ " to "+xlsx_file_path+" "+encoding_used)

    # Read the CSV file
    df = pd.read_csv(csv_file_path,header=None,encoding=encoding_used)
    myprint4("convert_csv_to_xlsx2 4")
    # Convert columns explicitly to numeric where appropriate
    for col in df.columns[1:]:  # Skip the first column if it's non-numeric (e.g., names)
        df[col] = pd.to_numeric(df[col], errors='ignore')

    # Write the DataFrame to an Excel file
    myprint4(" > Write to "+xlsx_file_path)

    header1=["SCRIPT_NAME"]
    header2=["Role"]
    for i in countingMethods:
        if i != "ALL":
            header1.append("?")
            txt=countingMethodNames[i]
            if i=="BLOCKS_50":
                txt="Répliques"
            header2.append(txt)
        
    header_rows = pd.DataFrame([
        header1,
        header2
    ])
    myprint4("convert_csv_to_xlsx2 3")
    # Concatenate the header rows and the original data
    # The ignore_index=True option reindexes the new DataFrame
    df = pd.concat([header_rows, df], ignore_index=True)

    # Write the DataFrame to an Excel file
    with pd.ExcelWriter(xlsx_file_path, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name='Sheet1')

        # Load the workbook and sheet for modification
        workbook = writer.book
        sheet = workbook['Sheet1']

        # Merge cells in the first and second new rows
        # Assuming you want to merge from the first to the last column
        col=get_excel_column_name(n)
        sheet.merge_cells("A1:"+col+"1")  # Modify range according to your number of columns
        sheet.merge_cells("A2:"+col+"2")  # Modify range according to your number of columns
 #       sheet.merge_cells('A2:D2')  # Modify this as needed
        sheet['A1'] = currentScriptFilename
        sheet['A2'] = "Length: "
        sheet.column_dimensions['A'].width = 50 
 # Load the workbook and get the active sheet
    myprint4("convert_csv_to_xlsx2 4")
  
def convert_dialog_csv_to_xlsx2(csv_file_path, xlsx_file_path, encoding_used):
    myprint4("convert_dialog_csv_to_xlsx2 "+csv_file_path+ " to "+xlsx_file_path+" "+encoding_used)

    # Read the CSV file
    df = pd.read_csv(csv_file_path,header=None,encoding=encoding_used)
    # Convert columns explicitly to numeric where appropriate
    for col in df.columns[1:]:  # Skip the first column if it's non-numeric (e.g., names)
        df[col] = pd.to_numeric(df[col], errors='ignore')

    # Write the DataFrame to an Excel file
    myprint4(" > Write to "+xlsx_file_path)

    header1=["Line","Character","Dialog"]
    header2=[]
    for i in countingMethods:
            txt=countingMethodNames[i]
            if i=="BLOCKS_50":
                txt="Lines"
            header1.append(txt)
        
    header_rows = pd.DataFrame([
        header1,
    ])
    myprint4("convert_csv_to_xlsx2 3")
    # Concatenate the header rows and the original data
    # The ignore_index=True option reindexes the new DataFrame
    df = pd.concat([header_rows, df], ignore_index=True)

    # Write the DataFrame to an Excel file
    with pd.ExcelWriter(xlsx_file_path, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name='Sheet1')

  



def generate_total_csv(total,csv_path,encoding_used,character_order_map):
    global currentXlsxPath
    #myprint4("Total csv path          : "+csv_path)
    #header

    s=""
    showHeader=False
    if showHeader:
        s="Role,"
        for m in countingMethods:
            if m!="ALL":                   
                s=s+str(m)+","
        s=s[0:len(s)-1]
        s=s+"\n"
    #myprint4("Total csv path          : 1")
    
    data = []
    order_idx=0

    for character in total:
        if (not character in currentDisabledCharacterNames) and (not character in currentMergedCharacters):
            order_idx=order_idx+1
            datarow=[str(order_idx)+" - " +str(character)];
            for method in total[character]:
                if method!="ALL":
                    #check merged characters and add eventually
                    #myprint4("char="+str(character))
                    if character in currentMergedCharactersTo:
                        #myprint4("in merged"+str(currentMergedCharactersTo))
                        mergedwith=currentMergedCharactersTo[character]
                        for k in mergedwith:
                            #myprint4("in merged add "+str(k))
                            total[character][method]=total[character][method]+total[k][method]
                    #else:
                        #myprint4("not merged")
                    #myprint4(str(character)+": Add method "+method+" = "+str(total[character][method]))
                    datarow.append(str(total[character][method]))

            data.append(datarow)
        else:
            myprint4("generate_total_csv skip"+character)
    
            
    #myprint4("data"+str(data))
    with open(csv_path, mode='w', newline='',encoding=encoding_used) as file:
        writer = csv.writer(file, delimiter=',', quotechar='"', quoting=csv.QUOTE_MINIMAL)
        
        # Write data to the CSV file
        for row in data:
            #myprint4("Write "+str(row))
            writer.writerow(row)

    
    xlsx_path=csv_path.replace(".csv",".xlsx")
    currentXlsxPath=xlsx_path
    n=len(countingMethods)+1
    #myprint4("Total xlsx path          : "+xlsx_path)
    convert_csv_to_xlsx2(csv_path,xlsx_path,n,encoding_used)

def on_button_click():
    myprint4("Button clicked!")
def export_csv():
    myprint4("Export")

def is_AND_character(character):
    return " AND " in character
def fill_breakdown_table(breakdown):
    myprint7("breakdown"+str(breakdown))
    for item in breakdown:
        type_=item['type']
        line_idx=item['line_idx']
        if(type_=="SCENE_SEP"):
            scene_id=item['scene_id']
            breakdown_table.insert('','end',values=("","","",""), tags=('border'))
            breakdown_table.insert('','end',values=(str(line_idx),"New scene",scene_id,""), tags=('scene','bold'))
        elif(type_=="SPEECH"):
            speech=item['speech']
            character=item['character']

            if not character in currentDisabledCharacterNames:
                breakdown_table.insert('','end',values=(str(line_idx),"Speech",character,speech))
        elif(type_=="NONSPEECH"):         
            text=item['text']
            breakdown_table.insert('','end',values=(str(line_idx),"Other",text,""), tags=('nonspeech',))
    #myprint4("NB ROWS = "+str(len(breakdown_table.get_children())))
    breakdown_table.update_idletasks()

def fill_stats_table(breakdown):
    for item in breakdown:
        type_=item['type']
        line_idx=item['line_idx']
        if(type_=="SPEECH"):
            speech=item['speech']
            filtered_speech=filter_speech(speech)
            character=item['character']
            tout=len(filtered_speech)
            if not character in currentDisabledCharacterNames:
                stats_table.insert('','end',values=(str(line_idx),character,speech,str(tout)))
    myprint4("NB ROWS = "+str(len(breakdown_table.get_children())))
    breakdown_table.update_idletasks()


def clear_table(treeview):
    """
    Clears all rows from the given Treeview table.
    
    Args:
        treeview (ttk.Treeview): The Treeview widget instance.
    """
    for item in treeview.get_children():
        treeview.delete(item)


###################################################################################################
## .INI FILE
def check_settings_ini_exists():
    # Get the absolute path of the directory where the script is located
#    script_folder = os.path.abspath(os.path.dirname(__file__))
    
    # Define the path to the settings ini file in the same directory as the script
    ini_file_path = get_setting_ini_path()

    # Check if the settings ini file exists
    if os.path.isfile(ini_file_path):
        myprint4(f"settings.ini file exists at: {ini_file_path}")
        return True
    else:
        myprint4(f"settings.ini file does not exist in the directory: {ini_file_path}")
        return False


def write_settings_ini():
    # Get the absolute path of the directory where the script is located
    script_folder = get_intial_treeview_folder_path()#  os.path.abspath(os.path.dirname(__file__))
    
    # Define the content to write to the settings ini file
    content = f"SCRIPT_FOLDER = {script_folder}"
    
    # Define the path to the settings.ini file in the same directory as the script
#    ini_file_path = os.path.join(script_folder, 'settings.ini')
    ini_file_path = get_setting_ini_path()
    # Write the content to the settings ini file
    with open(ini_file_path, 'w') as ini_file:
        ini_file.write(content)
    
    myprint4(f"settings.ini file created at: {ini_file_path}")


def read_settings_ini():
    
    # Define the path to the settings ini file in the same directory as the script
    ini_file_path = get_setting_ini_path()
    
    # Check if the settings ini file exists
    if not os.path.isfile(ini_file_path):
        raise FileNotFoundError(f"settings.ini file does not exist in the directory: {ini_file_path}")
    
    # Read the settings ini file and store settings in a dictionary
    settings = {}
    with open(ini_file_path, 'r') as ini_file:
        for line in ini_file:
            line = line.strip()
            if line and '=' in line:  # Ensure the line contains an '=' character
                key, value = line.split('=', 1)
                settings[key.strip()] = value.strip()
    
    return settings
def update_ini_settings_file(field,new_folder):
    # Get the absolute path of the directory where the script is located
    
    # Define the path to the settings ini file in the same directory as the script
    ini_file_path = get_setting_ini_path()

    # Check if the settings ini file exists
    if not os.path.isfile(ini_file_path):
        raise FileNotFoundError(f"settings.ini file does not exist in the directory: {ini_file_path}")
    
    # Read the current settings and store them in a dictionary
    settings = {}
    with open(ini_file_path, 'r') as ini_file:
        for line in ini_file:
            line = line.strip()
            if line and '=' in line:
                key, value = line.split('=', 1)
                settings[key.strip()] = value.strip()
    
    # Update the SCRIPT_FOLDER field
    settings[field] = new_folder
    
    # Write the updated settings back to the settings.ini file
    with open(ini_file_path, 'w') as ini_file:
        for key, value in settings.items():
            ini_file.write(f"{key} = {value}\n")
    
    myprint4(f"settings.ini file updated with SCRIPT_FOLDER = {new_folder}")

def on_folder_open(event):
    myprint4("on_folder_open")
    # Find the node that was opened
    oid = folders.focus()  # Get the ID of the focused item
    values = folders.item(oid, "values")
    myprint4("on_folder_open 1")

    if len(values) > 0 and values[0] == "dummy":
        myprint4("values>0 ignore")
        # Ignore the dummy nodes (if the first value in the tuple is "dummy")
        return
    myprint5("on_folder_open 2")

    # Check if the node has the dummy child indicating it hasn't been loaded
    children = folders.get_children(oid)
    if len(children) == 1 and folders.item(children[0], "values")[0] == "dummy":
        myprint5("on_folder_open 3")

        # Remove the dummy and load actual content
        folders.delete(children[0])
        myprint5("Load "+str(folders.item(oid, "values")))
        load_tree(oid, folders.item(oid, "values")[0])
    else:
        myprint5("on_folder_open 4 oid="+str(oid)+" val="+str( folders.item(oid, "values")))
        load_tree(oid, folders.item(oid, "values")[0])

def toggle_folder(event):
    myprint5("Toggle folder")
# Identify the item on which the click occurred
    x, y, widget = event.x, event.y, event.widget
    row_id = widget.identify_row(y)
    if not row_id:
        return  # Exit if the click didn't happen on a row
    
    # Toggle the open state of the node
    if widget.tag_has('folder', row_id):  # Check if the item has the 'folder' tag
        
        if widget.item(row_id, 'open'):  # If the folder is open, close it
            myprint5("opened, close")
            widget.item(row_id, open=False)

        else:  # If the folder is closed, open it
            myprint5("not opened, open")
            widget.item(row_id, open=True)
#            on_folder_open(event)
            children = folders.get_children(row_id)
            if len(children) == 1 and folders.item(children[0], "values")[0] == "dummy":
                myprint5("on_folder_open 3")

                # Remove the dummy and load actual content
                folders.delete(children[0])
            load_tree(row_id, folders.item(row_id, "values")[0])


    
def open_xlsx_recap():
    global currentXlsxPath
    os_=get_os()
    myprint5("Open"+currentXlsxPath)
    if os_=="Windows":
        currentOutputFolderAbs = os.path.abspath(currentXlsxPath)
        myprint5("Absolute path          : "+currentOutputFolderAbs)

    # Check if the folder exists
        if not os.path.exists(currentOutputFolderAbs):
            myprint5(f"Folder does not exist: {currentOutputFolderAbs}")
            return
        try:
            os.startfile(currentOutputFolderAbs)
        except Exception as e:
            myprint5(f"Failed to open file: {e}")
    else:
        try:
            subprocess.run(['open', currentXlsxPath], check=True)
        except subprocess.CalledProcessError as e:
            myprint5(f"Failed to open file: {e}")
def open_dialog_recap():
    global currentDialogPath
    os_=get_os()
    myprint5("Open"+currentDialogPath)
    if os_=="Windows":
        currentOutputFolderAbs = os.path.abspath(currentDialogPath)
        myprint5("Absolute path          : "+currentOutputFolderAbs)

    # Check if the folder exists
        if not os.path.exists(currentOutputFolderAbs):
            myprint5(f"Folder does not exist: {currentOutputFolderAbs}")
            return
        try:
            os.startfile(currentOutputFolderAbs)
        except Exception as e:
            myprint5(f"Failed to open file: {e}")
    else:
        try:
            subprocess.run(['open', currentDialogPath], check=True)
        except subprocess.CalledProcessError as e:
            myprint5(f"Failed to open file: {e}")

def open_file_in_system():
    global currentRightclickRowId
    myprint5("open_file_in_system" )
    selected_item = folders.focus() 
    file_path = folders.item(currentRightclickRowId, 'values')[0]
    myprint5(file_path)
    if os.path.isfile(file_path):
        
        myprint5(f"Opening file for item: {file_path}")
        abs_file_path=file_path#folders.item(selected_item)
        os_=get_os()
        myprint5("Open"+file_path)
        if os_=="Windows":
         
        # Check if the folder exists
            if not os.path.exists(abs_file_path):
                myprint5(f"Folder does not exist: {abs_file_path}")
                return
            try:
                os.startfile(abs_file_path)
            except Exception as e:
                myprint5(f"Failed to open file: {e}")
        else:
            try:
                subprocess.run(['open', abs_file_path], check=True)
            except subprocess.CalledProcessError as e:
                myprint5(f"Failed to open file: {e}")

def open_timeline():
    global currentTimelinePath
    os_=get_os()
    myprint5("Open"+currentTimelinePath)
    if os_=="Windows":
        currentOutputFolderAbs = os.path.abspath(currentTimelinePath)
        myprint5("Absolute path          : "+currentOutputFolderAbs)

    # Check if the folder exists
        if not os.path.exists(currentOutputFolderAbs):
            myprint5(f"Folder does not exist: {currentOutputFolderAbs}")
            return
        try:
            os.startfile(currentOutputFolderAbs)
        except Exception as e:
            myprint5(f"Failed to open file: {e}")
    else:
        try:
            subprocess.run(['open', currentTimelinePath], check=True)
        except subprocess.CalledProcessError as e:
            myprint5(f"Failed to open file: {e}")

def set_counting_method(i):
    myprint5("set method "+i)
    global countingMethod
    countingMethod=i


def show_popup_line_size():
    global currentBlockSize
    popupBlocksize = tk.Toplevel()
    popupBlocksize.title("Taille des répliques") 

    # Create a label
    label = ttk.Label(popupBlocksize, text="Saisir la taille des répliques (par ex. 50):")
    label.pack(pady=10)
    # Create a StringVar to hold the default value


    def set_line_size():
        global currentBlockSize
        global currentFilePath
        entry_value = input_blocksize.get()
        try:
            int_value = int(entry_value)
            myprint7("set line size "+str(entry_value))
            currentBlockSize=int_value
            popupBlocksize.destroy()
            if len(currentFilePath)>0:
                threading.Thread(target=runJob,args=(currentFilePath,countingMethod)).start()
        except ValueError:
                myprint7("Invalid input: not an integer")


    # Create a single-line text entry widget
    entry = ttk.Entry(popupBlocksize, width=30, textvariable=input_blocksize)
    entry.pack(pady=10)
    button = ttk.Button(popupBlocksize, text="Appliquer", command=set_line_size)
    button.pack(side=tk.TOP, fill=tk.X)

def show_popup_counting_method():
    popup = tk.Toplevel()
    popup.title("Popup") 

    for i in countingMethods:
        button = ttk.Button(popup, text=i, command=set_counting_method(i))
        button.pack(side=tk.TOP, fill=tk.X)

    dropdown = ttk.Combobox(popup, values=countingMethods)
    dropdown.pack(pady=20)
    dropdown.current(0)
 #   dropdown.bind('<<ComboboxSelected>>', on_value_change)

def resizechart(self, event=None):
        # Resize the figure to match the dimensions of its container
        width, height = event.width, event.height
        if width > 0 and height > 0:  # Check to prevent initial null-dimension error
            dpi = self.fig.get_dpi()
            self.fig.set_size_inches(width / dpi, height / dpi)
            self.canvas.draw()

def myprint5_frame_size(fr):
    # This function myprint5s the size of the frame
    # Ensure the frame has been rendered by Tkinter before calling this
    myprint5("Frame width:"+str( fr.winfo_width()))
    myprint5("Frame height:"+str( fr.winfo_height()))

def show_loading():
        #myprint5("SHOW LOADING")
        loading_label.pack()
        paned_window.pack_forget()
        app.update_idletasks()
        app.update()

def hide_loading():
        #myprint5("HIDE LOADING")
        myprint7("hide_loading")
        loading_label.pack_forget()
        paned_window.pack(fill='both', expand=True)
        app.update_idletasks()
        app.update()
        app.deiconify()

def disable_character():
    global currentCharacterSelectRowId
    global currentDisabledCharacterNames
    global currentBreakdown
    if currentCharacterSelectRowId:
        name = getCharacterTableNameByRowId(currentCharacterSelectRowId)
        currentCharacterSelectRowId=None
        myprint2("disable_character")
        myprint2("bd="+str(currentBreakdown))
        currentDisabledCharacterNames.append(name)
        reset_tables()
        postProcess(currentBreakdown,currentResultCharacterOrderMap,currentResultEnc,currentResultName,currentResultLinecountMap,currentResultSceneCharacterMap,currentTimelinePath)
        

def enable_character():
    global currentCharacterSelectRowId
    global currentDisabledCharacterNames
        
    if currentCharacterSelectRowId:
        name = getCharacterTableNameByRowId(currentCharacterSelectRowId)
        currentCharacterSelectRowId=None
     
        myprint2("enable_character")
        currentDisabledCharacterNames.remove(name)
        reset_tables()
        postProcess(currentBreakdown,currentResultCharacterOrderMap,currentResultEnc,currentResultName,currentResultLinecountMap,currentResultSceneCharacterMap,currentTimelinePath)

def deselect_characters():
    global currentCharacterSelectRowId
    currentCharacterSelectRowId=None
    

    global currentCharactersSelectedRowIds
    currentCharactersSelectedRowIds=[]
    # Deselect all selected rows
    selected_items = character_table.selection()
    character_table.selection_remove(selected_items)

def restore_characters():
    myprint2("restore_characters")
    global currentDisabledCharacterNames
    currentDisabledCharacterNames=[]
    reset_tables()
    postProcess(currentBreakdown,currentResultCharacterOrderMap,currentResultEnc,currentResultName,currentResultLinecountMap,currentResultSceneCharacterMap,currentTimelinePath)

def on_resize(event):
        return

class WordTableColumnSelector(tk.Toplevel):
    def __init__(self, parent, file_path):
        myprint7("TableColumnSelector init")
        self.parent = parent
        self.table_list = []
        self.doc = None
        self.file_path=file_path
        self.check_vars = []
        self.create_widgets()
        self.doc = Document(file_path)
        self.table_list = [table for table in self.doc.tables]
        self.update_table_listbox()
        myprint7("TableColumnSelector tablecount = "+str(len(self.table_list)))
        
    def reset(self,file_path):
        myprint7("TableColumnSelector RESET")  
        self.table_list = []
        self.doc = None
        self.check_vars = []
        for widget in self.list_frame.winfo_children():
            widget.destroy()
        self.check_vars = []

    def create_widgets(self):
        myprint7("TableColumnSelector create_widgets")
        for widget in self.parent.winfo_children():
            widget.destroy()

        # Frame for table preview
        self.table_frame = tk.Frame(self.parent, borderwidth=0)
        self.table_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

        # Frame for table list
        self.menu_frame = tk.Frame(self.parent, width=200)
        self.menu_frame.pack(side=tk.RIGHT, fill='y')
        self.menu_frame.pack_propagate(False) 
        self.left_canvas = tk.Canvas(self.menu_frame, borderwidth=0)
        self.left_canvas.pack(side=tk.TOP, fill='both', expand=True)

        self.list_frame = tk.Frame(self.left_canvas)
        self.left_canvas.create_window((0, 0), window=self.list_frame, anchor='nw')

        self.buttonframe2 = tk.Frame(self.menu_frame)
        self.buttonframe2.pack(side=tk.TOP, fill='x')

        # Create buttons to navigate pages
        self.open_button = tk.Button(self.buttonframe2, text="Lancer le traitement", command=self.run, height=10)
        self.open_button.pack(side=tk.TOP, fill='x', expand=True, padx=10, pady=10)

        # Add a canvas to allow scrolling
        self.canvas = tk.Canvas(self.table_frame)
        self.canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

        # Frame within the canvas to hold the table
        self.table_frame = tk.Frame(self.canvas)
        self.canvas.create_window((0, 0), window=self.table_frame, anchor="nw")


    def run(self):
        myprint7("WordTableColumnSelector run")
        dialog=self.get_dialog_col()
        character=self.get_character_col()
        if character>-1 and dialog>-1:
            myprint7(f"ch={character} di={dialog}")
            params={
                'param_type':'WORD'
                ,'character':character,
                'dialog':dialog
            }
            threading.Thread(target=runJob,args=(self.file_path,countingMethod,params)).start()

    def destroy(self):
        myprint7("WordTableColumnSelector destroy")
        self.menu_frame.destroy()
        self.table_frame.destroy()
        self.left_canvas.destroy()
        self.canvas.destroy()
        self.table_frame.destroy()

    def process():
        myprint7("process")

    def update_table_listbox(self):
        myprint7("WordTableColumnSelector update listbox")
        for widget in self.list_frame.winfo_children():
            widget.destroy()

        self.check_vars = []

        for i, _ in enumerate(self.table_list):
            var = tk.BooleanVar()
            if len(self.table_list) == 1:  # Check the checkbox by default if there's only one table
                var.set(True)
            chk = tk.Checkbutton(self.list_frame, variable=var)
            lbl = tk.Label(self.list_frame, text=f"Table {i+1}")
            chk.grid(row=i, column=0, sticky='w', padx=5, pady=2)
            lbl.grid(row=i, column=1, sticky='w', padx=5, pady=2)
            lbl.bind("<Button-1>", lambda e, idx=i: self.on_table_select(idx))
            self.check_vars.append(var)

        self.list_frame.update_idletasks()
        self.left_canvas.config(scrollregion=self.left_canvas.bbox("all"))

        if len(self.table_list)==1 :
            self.on_table_select(0)

    def get_dialog_col(self):
        for idx,k in (self.comboboxes.items()):
            if k.get() == "DIALOGUE" or k.get()=="LES DEUX":
                return idx
        return -1
    def get_character_col(self):
        for idx,k in (self.comboboxes.items()):
            if k.get() == "PERSONNAGE" or k.get()=="LES DEUX":
                return idx
        return -1
    def on_table_select(self, index):
        myprint7("WordTableColumnSelector on_table_select")
        table=self.table_list[index]
        myprint7("on_table_select idx="+str(index)+" table="+str(table))
        if len(self.detect_map)==0:
            success, mode, character,dialog,map_=detect_word_table(table,"",{})        
            if success:
                self.detect_map=map_
                self.show_table_preview(table,self.detect_map)
    comboboxes={}
    detect_map={}
    column_labels = []
    def show_table_preview(self,  table,map_):
        myprint7("WordTableColumnSelector show_table_preview")
        for widget in self.table_frame.winfo_children():
            widget.destroy()

        num_cols = len(table.rows[0].cells)
        self.column_labels = [[] for _ in range(num_cols)]
        options = ["-", "PERSONNAGE", "DIALOGUE", "LES DEUX"]
        col_widths = [0] * num_cols
          # Calculate the max width for each column based on the content
        for row in table.rows[:3]:
            sumcolwidth=0
            for col_idx, cell in enumerate(row.cells):
                cell_text = cell.text
                cell_text="\n".join(cell_text.split(" ")) 
                cell_width = tkFont.Font().measure(cell_text)
                if cell_width > col_widths[col_idx]:
                    col_widths[col_idx] = cell_width
                sumcolwidth=sumcolwidth+col_widths[col_idx]
            for col_idx, cell in enumerate(row.cells):
                if sumcolwidth>600:        
                  col_widths[col_idx] = int(col_widths[col_idx]*0.7)
                else:
                  col_widths[col_idx] = int(col_widths[col_idx])
            myprint7("sumcolwidth"+str(sumcolwidth))
        myprint7("colwidth"+str(col_widths))

        myprint7("set col val")
        for col_idx in range(num_cols):
            myprint7("combobox gen"+str(col_idx))
            combobox = ttk.Combobox(self.table_frame, values=options,width=col_widths[col_idx] // 8)
            mapval=map_[col_idx]
            myprint7("set col val"+str(mapval))
            mapvaltype=mapval['type']
            myprint7("set col val"+str(mapvaltype))

            if mapvaltype=='CHARACTER':
                combobox.current(1)         # Set default value to "-"
            elif mapvaltype=='DIALOG':
                        combobox.current(2) # Set default value to "-"
            elif mapvaltype=='BOTH':
                combobox.current(3) # Set default value to "-"
            else:            
                combobox.current(0)         # Set default value to "-"
            combobox.grid(row=0, column=col_idx, sticky='nsew')
            def create_on_combobox_change(col):
                def on_combobox_change(event):
                    myprint7("change col_idx" + str(col))
                    labels = self.column_labels[col]
                    bg="white"
                    fore="black"
                    headerbg="black"
                    headerfore="white"
                    val=self.comboboxes[col].get()
                    myprint7(val)
                    if val!='DIALOGUE' and val!='PERSONNAGE' and val!='LES DEUX':
                        fore="grey"
                        bg="#ddd"
                        headerbg="#555"
                        headerfore="#cccccc"

                    rowidx=0
                    for k in labels:
                        if rowidx==0:
                            k.config(background=headerbg, foreground=headerfore)
                        else:
                            k.config(background=bg, foreground=fore)
                        rowidx=rowidx+1
                return on_combobox_change

            combobox.bind("<<ComboboxSelected>>", create_on_combobox_change(col_idx))
            self.comboboxes[col_idx]=combobox

        myprint7("colwidth"+str(col_widths))
        for row_idx, row in enumerate(table.rows[:50]):
            for col_idx, cell in enumerate(row.cells):
                mapval=map_[col_idx]
                mapvaltype=mapval['type']
                bg="white"
                fore="black"
                headerbg="black"
                headerfore="white"
                val=self.comboboxes[col_idx].get()
                if val!='DIALOGUE' and val!='PERSONNAGE' and val!='LES DEUX':
                    fore="grey"
                    bg="#ddd"
                    headerbg="#555"
                    headerfore="#cccccc"

                cell_text = cell.text
                if row_idx==0:
                    cell_text="\n".join(cell_text.split(" ")) 
                if row_idx==0:
                    header_label = tk.Label(self.table_frame, text=cell_text, borderwidth=1, relief="solid", width=col_widths[col_idx] // 8, bg=headerbg, fg=headerfore)                
                else:
                    header_label = tk.Label(self.table_frame, text=cell_text, borderwidth=1, relief="solid", width=col_widths[col_idx] // 8, bg=bg,fg=fore)
                header_label.grid(row=row_idx + 2, column=col_idx, sticky='nsew')
                self.column_labels[col_idx].append(header_label)
        self.table_frame.update_idletasks()
        self.canvas.config(scrollregion=self.canvas.bbox("all"))

    def set_char_column(self, col):
        messagebox.showinfo("Character Column", f"Column {col+1} set as Character Column")

    def set_dialog_column(self, col):
        messagebox.showinfo("Dialog Column", f"Column {col+1} set as Dialog Column")

def open_table_selector(file_path):
    global importTab
    importTab=WordTableColumnSelector(tab_import,file_path)

class PDFCropper(tk.Toplevel):
    def __init__(self, parent):
        super().__init__(parent)
        self.title("PDF Cropper")

        self.canvas = tk.Canvas(self)
        self.canvas.pack(fill=tk.BOTH, expand=True)
        self.rect = None
        self.start_x = None
        self.start_y = None
        self.end_x = None
        self.end_y = None
        self.crop_coords = None

        self.select_button = tk.Button(self, text="Select", command=self.crop_pdf)
        self.select_button.pack(side=tk.LEFT, padx=5, pady=5)

        self.reset_button = tk.Button(self, text="Reset", command=self.reset_selection)
        self.reset_button.pack(side=tk.RIGHT, padx=5, pady=5)

        self.canvas.bind("<ButtonPress-1>", self.on_button_press)
        self.canvas.bind("<B1-Motion>", self.on_mouse_drag)
        self.canvas.bind("<ButtonRelease-1>", self.on_button_release)

    def open_pdf(self, file_path):
        self.pdf = pdfplumber.open(file_path)
        self.page = self.pdf.pages[9]  # Page 10 (index 9)
        self.render_page()

    def render_page(self):
        with self.page.to_image() as img:
            self.img = img.original
            self.tk_image = ImageTk.PhotoImage(self.img)
        
        self.canvas.create_image(0, 0, anchor="nw", image=self.tk_image)
        self.canvas.config(scrollregion=self.canvas.bbox(tk.ALL))
        self.canvas.config(width=self.img.width, height=self.img.height)

    def on_button_press(self, event):
        self.start_x = self.canvas.canvasx(event.x)
        self.start_y = self.canvas.canvasy(event.y)
        self.rect = self.canvas.create_rectangle(self.start_x, self.start_y, self.start_x, self.start_y, outline='red')

    def on_mouse_drag(self, event):
        curX, curY = (self.canvas.canvasx(event.x), self.canvas.canvasy(event.y))
        self.canvas.coords(self.rect, self.start_x, self.start_y, curX, curY)

    def on_button_release(self, event):
        self.end_x = self.canvas.canvasx(event.x)
        self.end_y = self.canvas.canvasy(event.y)
        self.crop_coords = (self.start_x, self.start_y, self.end_x, self.end_y)

    def crop_pdf(self):
        if not self.crop_coords:
            messagebox.showwarning("No Selection", "Please select an area to crop.")
            return

        x0, y0, x1, y1 = self.crop_coords
        x0, y0 = min(x0, x1), min(y0, y1)
        x1, y1 = max(x0, x1), max(y0, y1)

        # Convert canvas coordinates to PDF coordinates
        x0_pdf, y0_pdf = x0 / self.img.width * self.page.width, y0 / self.img.height * self.page.height
        x1_pdf, y1_pdf = x1 / self.img.width * self.page.width, y1 / self.img.height * self.page.height

        crop_box = (x0_pdf, y0_pdf, x1_pdf, y1_pdf)
        cropped_page = self.page.within_bbox(crop_box)

        save_path = filedialog.asksaveasfilename(defaultextension=".pdf", filetypes=[("PDF files", "*.pdf")])
        if save_path:
            with pdfplumber.open(self.pdf.stream) as pdf:
                new_pdf = pdfplumber.PDF()
                new_pdf.pages.append(cropped_page)
                new_pdf.save(save_path)
            messagebox.showinfo("Success", f"Cropped PDF saved to {save_path}")

    def reset_selection(self):
        if self.rect:
            self.canvas.delete(self.rect)
            self.rect = None
        self.crop_coords = None

def open_pdf_cropper():
    root = tk.Tk()
    root.withdraw()  # Hide the main root window

    file_path = filedialog.askopenfilename(filetypes=[("PDF files", "*.pdf")])
    if file_path:
        cropper = PDFCropper(root)
        cropper.open_pdf(file_path)
        cropper.mainloop()

def clear_chart():
    global currentFig
    myprint2("CLEAR CHART")
    currentFig.clf()
    currentFig.canvas.draw()

def merge_characters2():
    global currentCharactersSelectedRowIds
    global currentMergedCharactersTo
    global currentMergedCharacters
    global currentBreakdown

    myprint7("merge_character2")
    myprint7("merge_character2"+str(currentBreakdown))
    
    myprint7("merge_character2 currentCharactersSelectedRowIds"+str(currentCharactersSelectedRowIds))
    myprint7(str(currentMergedCharacters))
    myprint7(str(currentMergedCharactersTo))
    largestname=None
    largestnamerep=0
    for k in currentCharactersSelectedRowIds:

        name = character_table.item(k, 'values')[1]
        myprint7("merge_character2 Merge checklargest "+str(k)+str(name))
        
        line_count,word_count,character_count,replica_count=stats_per_character(currentBreakdown,name)
        rep=math.ceil(character_count/currentBlockSize)
        if rep>largestnamerep:
            largestname=name
            largestnamerep=rep
    
    for k in currentCharactersSelectedRowIds:
        name = character_table.item(k, 'values')[1]
        if name!=largestname:
            myprint7("merge_character2 Merge "+largestname+" <- "+str(name))
            mergeto = largestname
            mergefrom = name

            currentMergedCharacters[mergefrom]=mergeto
            if mergeto in currentMergedCharactersTo:
                existing=currentMergedCharactersTo[mergeto]
                myprint7("Merge already created" +str(existing))

                currentMergedCharactersTo[mergeto] = existing + (mergefrom,)
            else:
                myprint7("Merge new group")

                currentMergedCharactersTo[mergeto]=(mergefrom,)
    myprint7("currentMergedCharacters"+str(currentMergedCharacters))
    myprint7("currentMergedCharactersTo"+str(currentMergedCharactersTo))
    currentCharactersSelectedRowIds=[]
    reset_tables()
    postProcess(currentBreakdown,currentResultCharacterOrderMap,currentResultEnc,currentResultName,currentResultLinecountMap,currentResultSceneCharacterMap,currentTimelinePath)
def getCharacterTableNameByRowId(rowid):
    name = character_table.item(rowid, 'values')[1]
    return name

def merge_characters():
    myprint2("merge_characters")
    global currentCharacterMergeFromName
    name = getCharacterTableNameByRowId(currentCharacterSelectRowId)
    currentCharacterMergeFromName=name
    myprint2("Merge "+name)
    create_popup(currentResultCharacterOrderMap,name)

def hide_character():
    myprint2("hide_character")
    global currentDisabledCharacterNames
    name = getCharacterTableNameByRowId(currentCharacterSelectRowId)
    currentDisabledCharacterNames.append(name)
    #disabled_character_list_table.insert('','end',values=(name,))

    myprint2("Hide "+name)
    reset_tables()
    postProcess(currentBreakdown,currentResultCharacterOrderMap,currentResultEnc,currentResultName,currentResultLinecountMap,currentResultSceneCharacterMap,currentTimelinePath)


def on_right_click(event):
    myprint7("on right click")
    global currentRightclickRowId
    # Identify the row clicked
    try:
        row_id = folders.identify_row(event.y)
        if row_id:
            # Select the row under cursor
            #folders.selection_set(row_id)
            currentRightclickRowId=row_id
            menu.post(event.x_root, event.y_root)  # Show the context menu
    except Exception as e:
        myprint2(e)

def disable_merge_button():
   # btn_merge.config(state='disabled')
    myprint7("Button disabled")

def enable_merge_button():
    #btn_merge.config(state='normal')
    myprint7("Button enabled")
def on_character_table_click(event):
    global currentCharacterSelectRowId
    global currentCharactersSelectedRowIds


 # Identify the item clicked on
    item = character_table.identify_row(event.y)
    
    if item:
        currentCharacterSelectRowId=item
        
        if item in currentCharactersSelectedRowIds:
            # If the item is already in the list, remove it
            currentCharactersSelectedRowIds.remove(item)
            character_table.selection_remove(item)
            character_table.item(item, tags=())
        else:
            # If the item is not in the list, add it
            currentCharactersSelectedRowIds.append(item)
            character_table.selection_add(item)
            character_table.item(item, tags=('grouped',))

    myprint7(f"Number of selected rows: {currentCharactersSelectedRowIds}")
    



def open_result_folder():
    # Open a folder in Finder using the `open` command
    myprint2("Opening "+currentOutputFolder)
    currentOutputFolderAbs = os.path.abspath(currentOutputFolder)
    myprint2("Absolute path          : "+currentOutputFolderAbs)

  # Check if the folder exists
    if not os.path.exists(currentOutputFolderAbs):
        myprint2(f"Folder does not exist: {currentOutputFolderAbs}")
        return
    try:
        if sys.platform.startswith('darwin'):
            subprocess.run(['open', currentOutputFolderAbs], check=True)
        elif sys.platform.startswith('win32'):
            # Correct approach for Windows
            subprocess.run(['explorer', currentOutputFolderAbs], check=True)
        elif sys.platform.startswith('linux'):
            subprocess.run(['xdg-open', currentOutputFolderAbs], check=True)
    except Exception as e:
        myprint2(f"Error opening folder: {e}")

def open_conversion_file():
    # Open a folder in Finder using the `open` command
    myprint2("Opening "+currentOutputFolder)
    currentOutputFolderAbs = os.path.abspath(currentOutputFolder)
    myprint2("Absolute path          : "+currentOutputFolderAbs)

  # Check if the folder exists
    if not os.path.exists(currentOutputFolderAbs):
        myprint2(f"Folder does not exist: {currentOutputFolderAbs}")
        return
    try:
        if sys.platform.startswith('darwin'):
            subprocess.run(['open', currentOutputFolderAbs], check=True)
        elif sys.platform.startswith('win32'):
            # Correct approach for Windows
            subprocess.run(['explorer', currentOutputFolderAbs], check=True)
        elif sys.platform.startswith('linux'):
            subprocess.run(['xdg-open', currentOutputFolderAbs], check=True)
    except Exception as e:
        myprint2(f"Error opening folder: {e}")

def show_importtable_tab():
    global currentHasImportTableTab
    myprint7("show_importtable_tab")
    if not currentHasImportTableTab:
        notebook.insert(0,tab_import, text='Tables Word')
        notebook.update_idletasks()
        currentHasImportTableTab=True


def hide_importtable_tab():
    global currentHasImportTableTab    
    myprint7("hide_importtable_tab"+str(len(notebook.tabs())))
    if currentHasImportTableTab:#tab_import in notebook.tabs():
        myprint7("hide_importtable_tab has")
        
        notebook.forget(tab_import)
        notebook.update_idletasks()
        currentHasImportTableTab=False
    #else:
     #   myprint7("hide_importtable_tab no")
    

def create_popup(character_map, mergedchar):
    global currentMergedCharactersTo
    global currentMergePopupWindow
    global currentMergePopupTable
    popup = Toplevel(app)
    popup.title("Fusionner")
    popup.geometry("300x550")  # Size of the popup window
    currentMergePopupWindow = popup

    # Create a frame for the Treeview and Scrollbar
    tree_frame = tk.Frame(popup)
    tree_frame.pack(fill='both', expand=True)

    # Create the Treeview
    popup_character_list_table = ttk.Treeview(tree_frame, columns=('Character'), show='headings')
    # Define the column headings
    popup_character_list_table.heading('Character', text='')

    # Define the column width and alignment
    popup_character_list_table.column('Character', width=150, anchor='w')
    popup_character_list_table.pack(side='left', fill='both', expand=True)

    # Create a vertical scrollbar and associate it with the Treeview
    scrollbar = ttk.Scrollbar(tree_frame, orient='vertical', command=popup_character_list_table.yview)
    popup_character_list_table.configure(yscroll=scrollbar.set)
    scrollbar.pack(side='right', fill='y')

    for k in character_map:
        if k != mergedchar:
            if not (k in currentMergedCharactersTo):    
                popup_character_list_table.insert('', 'end', text=k, values=[k])

    currentMergePopupTable = popup_character_list_table

    # Frame to control the size of the button
    button_frame = tk.Frame(popup, height=40)  # Set the height to 40 pixels
    button_frame.pack(fill='x')  # Fill frame horizontally
    button_frame.pack_propagate(False)  # Prevent frame from resizing to fit contents

    close_btn = tk.Button(button_frame, text="Merge", command=merge_with)
    close_btn.pack(fill='x', expand=True)


def create_popup_old(character_map):
    global currentMergePopupWindow
    global currentMergePopupTable
    popup = Toplevel(app)
    popup.title("Merge with")
    popup.geometry("200x350")  # Size of the popup window
    currentMergePopupWindow=popup

    popup_character_list_table = ttk.Treeview(popup, columns=('Character'), show='headings')
    # Define the column headings
    popup_character_list_table.heading('Character', text='')

    # Define the column width and alignment
    popup_character_list_table.column('Character', width=50, anchor='w')
   # popup_character_list_table.bind('<Button-1>', merge_with)
    popup_character_list_table.pack(fill='both',expand=True)
    
    for k in character_map:
        popup_character_list_table.insert('', 'end', text=k, values=[k])

    currentMergePopupTable=popup_character_list_table


    # Frame to control the size of the button
    button_frame = tk.Frame(popup, height=40)  # Set the height to 40 pixels
    button_frame.pack(fill='x')  # Fill frame horizontally
    button_frame.pack_propagate(False)  # Prevent frame from resizing to fit contents

    close_btn = tk.Button(button_frame, text="Merge", command=merge_with)
    close_btn.pack(fill='x',expand=True)



def clear_character_stats():
    for item in character_stats_table.get_children():
        character_stats_table.delete(item)
def on_item_selected(event):
    tree = event.widget
    selection = tree.selection()
    item = tree.item(selection)
    record = item['values']
    # Do something with the selection, for example:
    if len(record)==0:
        return
    clear_character_stats()
    character_name=record[0]
    
    #character_named = character_name 
    #character_list_table.insert('','end',values=(character_named,))
   
    rowtotal=("",character_name,"","TOTAL")       
    total_by_method={}
    for m in countingMethods:
        total_by_method[m]=0

    for item in currentBreakdown:
        line_idx=item['line_idx']
        type_=item['type']
        if(type_=="SPEECH"):

            speech=item['speech']
            character=item['character']
            character_raw=item['character_raw']
            
            filtered_speech=filter_speech(speech)

            if character==character_name:
                #myprint2("    MATCH"+str(speech))

                row=(str(line_idx),character,character_raw, speech)
                for m in countingMethods: 
                    #myprint2("add"+str(m))
                    le=compute_length_by_method(filtered_speech,m)
                    row=row+(str(le),)
                    total_by_method[m]=total_by_method[m]+le
                #myprint2("add"+str(row))
                character_stats_table.insert('','end',values=row)
    for m in countingMethods:
        if m.startswith("BLOCKS"):
            total_by_method[m]=math.ceil(total_by_method[m])

    for m in countingMethods:
        rowtotal=rowtotal+(total_by_method[m],)
    character_stats_table.insert('',0,values=rowtotal,tags=['total'])

class PDFViewer:
    def __init__(self, root, file_path,currentOutputFolder,encoding):
        myprint7("pdf1")
        self.root = root
#        self.strategy="OCR"
        self.strategy="BLOCKS"
        for widget in self.root.winfo_children():
            widget.destroy()
        self.currentOutputFolder=currentOutputFolder
        self.encoding=encoding
        self.file_path = file_path
        self.page_number = 1
        self.scale=1
        self.left_threshold=125
        self.right_threshold=189
        self.top_threshold=68
        self.bottom_threshold=37
        if self.strategy=="OCR":
            self.left_threshold=90
            self.right_threshold=47
            self.top_threshold=50
            self.bottom_threshold=37
        elif self.strategy=="BLOCKS":
            self.left_threshold=125
            self.right_threshold=189
            self.top_threshold=68
            self.bottom_threshold=37
        
        self.text_elements=None
        self.canvas_height=0
        self.canvas_width=0

    
        self.top_offset=0
        self.left_offset=0
        # Initialize the vertical and horizontal line IDs
        self.vertical_line = None
        self.vertical_liner = None
        self.horizontal_line = None
        self.horizontal_liner = None

        myprint7("pdf2")
        # Open the PDF file
        self.pdf_document = pdfplumber.open(self.file_path)
        self.num_pages = len(self.pdf_document.pages)
        myprint7("Num pages : "+str(self.num_pages))
        self.input_firstpage = tk.StringVar()
        self.input_lastpage = tk.StringVar()
        if self.num_pages>10:
            self.input_firstpage.set(str(2))
            self.input_lastpage.set(str(self.num_pages))
        elif self.num_pages>2:
            self.input_firstpage.set(str(2))
            self.input_lastpage.set(str(self.num_pages))
        else:
            self.input_firstpage.set(str(1))
            self.input_lastpage.set(str(self.num_pages))
    
        self.pdf_page_width = self.pdf_document.pages[0].width
        self.pdf_page_height = self.pdf_document.pages[0].height
        myprint7("PDF Page width : "+str(self.pdf_page_width)+" x "+str(self.pdf_page_height))

        col="#ececec"

        # Create a frame for the canvas and scrollbars
        self.canvas_frame = tk.Frame(root)
        self.canvas_frame.pack(side='left', fill='both', expand=True)

        # Create a Canvas to display the PDF page
        self.canvas = Canvas(self.canvas_frame,bg="#ffffff")
        self.canvas.pack(side='left',fill='y', expand=False,padx=0, pady=0)
       

        # Create a frame for sliders
        self.slider_frame = tk.Frame(root,bg=col)
        self.slider_frame.pack(fill='both', expand=True,side='top')

        

            # Create a label# Create a bold font
        bold_font = tkfont.Font(weight="bold")

        # Create the label
        label = tk.Label(self.slider_frame, text="Pages à traiter",         font=bold_font, anchor="w", justify="left", bg=self.slider_frame.cget('bg'))
        label.pack(pady=0, fill="x", padx=(10, 0)) 

        label = tk.Label(self.slider_frame, text="Première page:", anchor="w", justify="left", bg=self.slider_frame.cget('bg'))
        label.pack(pady=0, fill="x", padx=(10, 0)) 
        # Create a StringVar to hold the default value

        # Create a single-line text entry widget
        entry = ttk.Entry(self.slider_frame, width=30, textvariable=self.input_firstpage)
        entry.pack(pady=10)

        label = tk.Label(self.slider_frame, text="Dernière page:", anchor="w", justify="left", bg=self.slider_frame.cget('bg'))
        label.pack(pady=0, fill="x", padx=(10, 0)) 

        # Create a single-line text entry widget
        entry = ttk.Entry(self.slider_frame, width=30, textvariable=self.input_lastpage)
        entry.pack(pady=10)

 # Create the label
        label = tk.Label(self.slider_frame, text="Zone à traiter",         font=bold_font, anchor="w", justify="left", bg=self.slider_frame.cget('bg'))
        label.pack(pady=0, fill="x", padx=(10, 0)) 
    # Create a label
        label = tk.Label(self.slider_frame, text="Ajuster les marges pour faire apparaitre le dialogue en bleu.", anchor="w", justify="left", bg=self.slider_frame.cget('bg'))
        label.pack(pady=0, fill="x", padx=(10, 0)) 

        # Create a horizontal slider for width threshold
        self.left_slider = Scale(self.slider_frame, from_=0, to=900, orient=HORIZONTAL, label='Seuil gauche', command=self.update_left_line)
        self.left_slider.set(self.left_threshold)
        self.left_slider.pack(side='top', fill='x', padx=10, pady=5)


        # Create a horizontal slider for width threshold
        self.right_slider = Scale(self.slider_frame, from_=0, to=1300, orient=HORIZONTAL, label='Seuil droite', command=self.update_right_liner)
        self.right_slider.set(self.right_threshold)
        self.right_slider.pack(side='top', fill='x', padx=10, pady=5)

        # Create a vertical slider for height threshold
        self.top_slider = Scale(self.slider_frame, from_=0, to=1200, orient=HORIZONTAL, label='Seuil haut', command=self.update_top_line)
        self.top_slider.set(self.top_threshold)
        self.top_slider.pack(side='top', fill='x', padx=5, pady=10)

        # Create a vertical slider for height threshold
        self.bottom_slider = Scale(self.slider_frame, from_=0, to=1200, orient=HORIZONTAL, label='Seuil bas', command=self.update_bottom_liner)
        self.bottom_slider.set(self.bottom_threshold)
        self.bottom_slider.pack(side='top', fill='x', padx=5, pady=10)



        if False:
            # Add scrollbars to the canvas
            self.h_scrollbar = Scrollbar(self.canvas_frame, orient='horizontal', command=self.canvas.xview)
            self.h_scrollbar.pack(side='bottom', fill='x')
            self.v_scrollbar = Scrollbar(self.canvas_frame, orient='vertical', command=self.canvas.yview)
            self.v_scrollbar.pack(side='right', fill='y')
            self.canvas.configure(xscrollcommand=self.h_scrollbar.set, yscrollcommand=self.v_scrollbar.set)


        # Create a frame for sliders
        self.slider_frame = tk.Frame(root)
        self.slider_frame.pack(side='top', fill='x')

        myprint7("pdf4")


        # Draw the initial vertical line
        self.draw_vertical_line(self.left_threshold)

        self.buttonframe0 = tk.Frame(root,bg=col)
        self.buttonframe0.pack(side='top', fill='x')

        self.buttonframe = tk.Frame(root,bg=col)
        self.buttonframe.pack(side='top', fill='x')

        self.buttonframe2 = tk.Frame(root,bg=col)
        self.buttonframe2.pack(side='top', fill='x')
        style = ttk.Style()
        # Configure the style
        style.configure("Custom.Horizontal.TProgressbar", 
                        troughcolor='#ececec',  # Light gray background
                        background='#4CAF50',   # Green progress color
                        thickness=20)  
        
        # Create and pack the progress bar
        self.progress_bar = ttk.Progressbar(self.buttonframe2, orient='horizontal', mode='determinate',  style="Custom.Horizontal.TProgressbar")
        self.progress_bar.pack(pady=0, fill='x', expand=True)
        # Create the label
        label = tk.Label(self.buttonframe0, text="Aperçu",         font=bold_font, anchor="w", justify="left", bg=self.slider_frame.cget('bg'))
        label.pack(pady=0, fill="x", padx=(10, 0)) 

        # Create buttons to navigate pages
        self.prev_button = tk.Button(self.buttonframe, text="Page précedente", command=self.prev_page)
        self.prev_button.pack(side=tk.LEFT, padx=10, pady=10)

        self.label = tk.Label(self.buttonframe0, text="Page: "+str(self.page_number+1)+" / "+str(self.num_pages), font=('Arial', 12), bg=self.slider_frame.cget('bg'))
        self.label.pack(side=tk.BOTTOM, fill=tk.X)

        self.progress_bar['value'] = 0
        self.max_value = self.num_pages
        self.progress_bar['maximum'] = self.max_value

        self.next_button = tk.Button(self.buttonframe, text=" Page suivante", command=self.next_page)
        self.next_button.pack(side=tk.RIGHT, padx=10, pady=10)

        self.fr=tk.Frame(self.buttonframe2,height=40)
        self.next_button.pack(side=tk.LEFT, padx=10, pady=10)

        # Create buttons to navigate pages
        self.open_button = tk.Button(self.buttonframe2, text="Lancer le traitement", command=self.run,height=10)
        self.open_button.pack(side=tk.LEFT, fill='x',expand=True, padx=10, pady=10)

     
        # Display the initial page
        self.display_page(self.page_number)
        self.current_image = None
        self.redraw()
    
    def draw_horizontal_line(self, y):
        # Remove the old horizontal line if it exists
        if self.horizontal_line is not None:
            self.canvas.delete(self.horizontal_line)

        # Draw the new horizontal line
        self.horizontal_line = self.canvas.create_line(0, y, self.canvas.winfo_width(), y, fill="black")
    def draw_horizontal_liner(self, y):
        # Remove the old horizontal line if it exists
        if self.horizontal_liner is not None:
            self.canvas.delete(self.horizontal_liner)

        # Draw the new horizontal line
        self.horizontal_liner = self.canvas.create_line(0, y, self.canvas.winfo_width(), y, fill="black")
    show_preview=True

    def display_page(self, page_number):
        return self.display_page_alt(page_number)
    
    def display_page_alt(self, page_number):
        # Render the page as an image
        myprint7("pdf8"+str(page_number))
        page = self.pdf_document.pages[page_number]
        myprint7("pdf8a")
        myprint7("pdf8a uses pdfium")
    
        notebook.select(1)
        app.update_idletasks()
            
        w=2.08
        img_width=self.pdf_page_width*w
        img_height=self.pdf_page_height*w

        canvasframe_width = self.canvas_frame.winfo_width()
        canvasframe_height = self.canvas_frame.winfo_height()
        new_width=canvasframe_height/img_height*img_width
        myprint7(f"RESIZE TO  {new_width} x {canvasframe_height}")
        self.canvas.config(width=new_width, height=canvasframe_height)

        canvas_width = new_width# self.canvas.winfo_width()
        canvas_height = canvasframe_height# self.canvas.winfo_height()
        myprint7(f"CANVAS   {canvas_width} y= {canvas_height}")
        
        self.canvas_height=canvas_height
        self.canvas_width=canvas_width
        # Calculate the scaling factor
        scale_factor = min(canvas_width / img_width, canvas_height / img_height)
        self.scale=scale_factor
        # Scale the image
        scaled_width = int(img_width * scale_factor)
        scaled_height = int(img_height * scale_factor)
#            myprint7(f"img_width     {img_width}")
#           myprint7(f"img_height    {img_height}")
        myprint7(f"canvas_height {canvas_height}")
        myprint7(f"canvas_width  {canvas_width}")
        myprint7(f"scale_factor  {scale_factor}")

        myprint7(f"scaled_width  {scaled_width}")
        myprint7(f"scaled_height {scaled_height}")

            
        
        myprint7("pdf8d")

        # Clear the canvas and display the image
        self.canvas.delete("all")
#        self.canvas.create_image(0, 0, anchor=tk.NW, image=nimg_tk)

        self.pdf_mode,self.text_elements,self.page_split_elements=get_pdf_text_elements(self.file_path,self.page_number,self.get_first_page(),self.get_last_page(),self.progress_bar)
        
        x1=self.left_threshold
        x2=self.canvas_width-self.right_threshold
        y1=self.canvas_height-self.top_threshold
        y2=self.bottom_threshold
        x2=self.pdf_page_width-self.right_threshold
        y1=self.pdf_page_height-self.top_threshold

        myprint7("run split with left={x1} top={y1} right={x2} bottom={y2} ")
        self.redraw()

    pdf_mode=None


    # def display_page_pdfium(self, page_number):
    #     # Render the page as an image
    #     myprint7("pdf8"+str(page_number))
    #     page = self.pdf_document.pages[page_number]
    #     myprint7("pdf8a")
    #     myprint7("pdf8a uses pdfium")
    #     if self.show_preview:
            
    #         image = page.to_image(resolution=150,)
    #         myprint7("pdf8b")
    #         img = image#Image.open(io.BytesIO(image.original))
    #         img.save("pdfpreview.png")
    #         myprint7("pdf8c")

    #     app.after(100, self.display_page_2_scale)


    def draw_vertical_line(self, x):
        # Remove the old vertical line if it exists
        if self.vertical_line is not None:
            self.canvas.delete(self.vertical_line)

        # Draw the new vertical line
        self.vertical_line = self.canvas.create_line(x, 0, x, self.canvas.winfo_height(), fill="black")
    def draw_vertical_liner(self, x):
        # Remove the old vertical line if it exists
        if self.vertical_liner is not None:
            self.canvas.delete(self.vertical_liner)

        # Draw the new vertical line
        self.vertical_liner = self.canvas.create_line(x, 0, x, self.canvas.winfo_height(), fill="black")

    def update_left_line(self, event):
        # Update the vertical line position based on the slider value
        self.left_threshold=self.left_slider.get()
        self.redraw()
    def update_right_liner(self, event):
        # Update the vertical line position based on the slider value
        self.right_threshold= self.right_slider.get()
        self.redraw()
    def update_top_line(self, event):
        # Update the vertical line position based on the slider value
        self.top_threshold= self.top_slider.get()
        self.redraw()
    def update_bottom_liner(self, event):
        # Update the vertical line position based on the slider value
        self.bottom_threshold= self.bottom_slider.get()
        self.redraw()
    def next_page(self):
        if self.page_number < self.num_pages - 1:
            self.page_number += 1
            self.display_page(self.page_number)
            self.label.config(text="Page: "+str(self.page_number)+" / "+str(self.num_pages))

    def margin_to_positionx(self,x):
        return self.left_offset + x
    def margin_to_positionx2(self,x):
        myprint7(f"rmargin {x} width={self.canvas_width} mar={self.left_offset}")
        return self.canvas_width+ self.left_offset - x
    def margin_to_positiony(self,x):
        return self.top_offset + x
    def margin_to_positiony2(self,x):
        return self.canvas_height+ self.top_offset - x

    def group_words_into_lines(self, words, y_tolerance=3):
        #myprint7("group_words_into_lines")
        myprint7(f"group_words_into_lines N input ={len(words)}")
        
        # Dictionary to hold words grouped by their y position
        lines_map = {}
        
        for word in sorted(words, key=lambda w: (w['bbox'][1], w['bbox'][0])):
            text = word['text']
            x0, y0, x1, y1 = word['bbox']

            # Find the appropriate y position group
            added = False
            for y in list(lines_map.keys()):
                if abs(y - y0) <= y_tolerance:
                    lines_map[y].append(word)
                    added = True
                    break
            
            if not added:
                lines_map[y0] = [word]
        
  
        myprint7(f"group_words_into_lines 1 map={lines_map.keys()}")
        for idx,k in enumerate(lines_map):
            myprint7(f"{idx} : {k} {lines_map[k]}")

        # Merge the words in the map and output as a sorted list of lines

        lines = []
        for y in sorted(lines_map.keys()):
            sorted_line = sorted(lines_map[y], key=lambda w: w['bbox'][0])
            full_text = ' '.join([word['text'] for word in sorted_line])
            bbox = sorted_line[0]['bbox']
            lines.append({
                'text': full_text,
                'bbox': bbox
            })
        myprint7(f"group_words_into_lines 1 done")

     #   myprint7(f"group_words_into_lines N out ={len(lines)}")
      #  myprint7(f"group_words_into_lines res={lines}")
     #   myprint7(f"group_words_into_lines map={lines_map}")
        return lines

    def redraw(self):
        self.canvas.delete("all")
        myprint7("---------- redraw --------------")
        x1=self.left_threshold
        x2=self.canvas_width-self.right_threshold
        y1=self.canvas_height-self.top_threshold
        y2=self.bottom_threshold
        x2=self.pdf_page_width-self.right_threshold
        y1=self.pdf_page_height-self.top_threshold
        
        myprint7(f"run split with left={x1} top={y1} right={x2} bottom={y2} ")

        res=split_elements(self.text_elements,x1,y1,x2,y2)



        left=res['left']
        center=res['center']
        right=res['right']
        top=res['top']
        bottom=res['bottom']
        self.centered_blocks=center
        left_margin=self.left_offset
        
        if self.pdf_mode=="TABLE":
#            myprint7("TABLE CENTER group "+str(res))
 #           myprint7("TABLE CENTER group "+str(self.centered_blocks))
            
            center=self.group_words_into_lines(self.centered_blocks,3)

        #blue
        self.canvas.create_rectangle(0,0,self.canvas_width,self.canvas_height, outline="#cccccc", width=2,fill="#cccccc")

        #white
        self.canvas.create_rectangle(self.margin_to_positionx(self.left_threshold),self.margin_to_positiony(self.top_threshold),self.margin_to_positionx2(self.right_threshold),self.margin_to_positiony2( self.bottom_threshold), outline="white", width=2,fill="white")
        
        s="#bbbbbb"
        f="#dddddd"
        for k in left:
            self.draw_bbox(self.canvas_height,150,self.scale,k['bbox'],k['text'],s,f,left_margin)
        for k in center:
            self.draw_bbox(self.canvas_height,150,self.scale,k['bbox'],k['text'],"#246fd6","#99b3d6",left_margin)
        for k in right:
            self.draw_bbox(self.canvas_height,150,self.scale,k['bbox'],k['text'],s,f,left_margin)
        for k in top:
            self.draw_bbox(self.canvas_height,150,self.scale,k['bbox'],k['text'],s,f,left_margin)
        for k in bottom:
            self.draw_bbox(self.canvas_height,150,self.scale,k['bbox'],k['text'],s,f,left_margin)
        self.draw_lines()
    def draw_bbox(self,canvas_height,dpi,scale_factor, bbox_points,text,color,fillcolor,left_margin):
        # Conversion factors
        #myprint7("draw bbox")
        #myprint7(bbox_points)
        # Convert points to pixels
        x0, y0, x1, y1 = [coord * dpi / 72 for coord in bbox_points]

        # Apply the scale factor
        x0, y0, x1, y1 = [coord * scale_factor for coord in (x0, y0, x1, y1)]
        y0=canvas_height-y0
        y1=canvas_height-y1
        
        x0=x0+left_margin
        x1=x1+left_margin
        #myprint7("res"+str(x0)+ " "+str(y0)+ " "+str(x1)+ " "+str(y1)+ " ")

        # Draw the rectangle on the canvas
        self.canvas.create_rectangle(x0, y0, x1, y1, outline=color, width=1,fill=fillcolor)
        self.canvas.create_text(x0,y1-7,text=text, fill=color, font=("Courier", 11, "normal"), anchor=tk.NW)
        
    # def display_page_2_scale(self):
    #     notebook.select(1)
    #     app.update_idletasks()
    #     if self.show_preview:
    #         nimg = Image.open("pdfpreview.png")
    #         # Get the dimensions of the image and the canvas
    #         img_width, img_height = nimg.size

    #         canvasframe_width = self.canvas_frame.winfo_width()
    #         canvasframe_height = self.canvas_frame.winfo_height()
    #         new_width=canvasframe_height/img_height*img_width
    #         myprint7(f"RESIZE TO  {new_width} x {canvasframe_height}")
    #         self.canvas.config(width=new_width, height=canvasframe_height)

    #         canvas_width = new_width# self.canvas.winfo_width()
    #         canvas_height = canvasframe_height# self.canvas.winfo_height()
    #         myprint7(f"CANVAS   {canvas_width} y= {canvas_height}")
            
    #         self.canvas_height=canvas_height
    #         self.canvas_width=canvas_width
    #         # Calculate the scaling factor
    #         scale_factor = min(canvas_width / img_width, canvas_height / img_height)
    #         self.scale=scale_factor
    #         # Scale the image
    #         scaled_width = int(img_width * scale_factor)
    #         scaled_height = int(img_height * scale_factor)
    #         myprint7(f"img_width     {img_width}")
    #         myprint7(f"img_height    {img_height}")
    #         myprint7(f"canvas_height {canvas_height}")
    #         myprint7(f"canvas_width  {canvas_width}")
    #         myprint7(f"scale_factor  {scale_factor}")

    #         myprint7(f"scaled_width  {scaled_width}")
    #         myprint7(f"scaled_height {scaled_height}")
    #         scaled_img = nimg.resize((scaled_width, scaled_height), Image.Resampling.LANCZOS)
    #         scaled_img.save("pdfpreview_scaled.png")
    #     app.after(100, self.display_page_3_open_preview)

    def get_first_page(self):
        return int(self.input_firstpage.get())
    def get_last_page(self):
        return int(self.input_lastpage.get())
    page_split_elements=[]
    def run_elements(self):
        myprint7("RUN PDF")
        global currentFilePath
        global currentScriptFilename
        global currentBreakdown
        global currentOutputFolder
        global currentTimelinePath

        global currentResultCharacterOrderMap
        global currentResultEnc
        global currentResultName
        global currentResultLinecountMap
        global currentResultSceneCharacterMap
        global importTab
        global currentPDFPages
        global currentPDFPageIdx

        self.pdf_mode,self.all_text_elements,self.page_split_elements=get_pdf_text_elements(self.file_path,-1,self.get_first_page(),self.get_last_page(),self.progress_bar)
        myprint7("all blocks: "+str(len(self.all_text_elements)))
        

        x1=self.left_threshold
        x2=self.canvas_width-self.right_threshold
        y1=self.canvas_height-self.top_threshold
        y2=self.bottom_threshold
        x2=self.pdf_page_width-self.right_threshold
        y1=self.pdf_page_height-self.top_threshold
        
        if self.pdf_mode=="TABLE":
            myprint7("MERGE_PAGE_WORDS")
            line_merged_elements=[]
            for k in self.page_split_elements:
                res=split_elements(k,x1,y1,x2,y2)
                
                grouped=self.group_words_into_lines(res['center'],3)
                grouped = sorted(grouped, key=lambda x: x['bbox'][1], reverse=True)

                line_merged_elements.extend(grouped)
            self.all_centered_blocks=line_merged_elements
        else:

 #           myprint7("TABLE CENTER group "+str(res))
  #          myprint7("TABLE CENTER group "+str(self.all_centered_blocks))
            
   #         self.all_centered_blocks=self.group_words_into_lines(self.all_centered_blocks,3)
    #        myprint7("TABLE CENTER group after "+str(self.all_centered_blocks))



            res=split_elements(self.all_text_elements,x1,y1,x2,y2)
            self.all_centered_blocks=res['center'];
        
#        if self.pdf_mode=="TABLE":
 #           myprint7("TABLE CENTER group "+str(res))
  #          myprint7("TABLE CENTER group "+str(self.all_centered_blocks))
            
   #         self.all_centered_blocks=self.group_words_into_lines(self.all_centered_blocks,3)
    #        myprint7("TABLE CENTER group after "+str(self.all_centered_blocks))


        myprint7("centered blocks: "+str(len(self.all_centered_blocks)))
        converted_file_path,enc=run_convert_pdf_to_txt(self.file_path,self.currentOutputFolder,self.all_centered_blocks, self.encoding)
        myprint7("converted "+str(converted_file_path))
        
        with open(converted_file_path, 'r', encoding=self.encoding) as file:
            file_path=converted_file_path
            file_name = os.path.basename(file_path)
            currentScriptFilename=file_name
            name, extension = os.path.splitext(file_name)
            myprint7("Opened")
            content = file.read()
            myprint7("Read")
            file_preview.text_widget.delete(1.0, tk.END)
            file_preview.text_widget.insert(tk.END, content)
            
            enc=self.encoding
            myprint7("Process")
            breakdown,character_scene_map,scene_characters_map,character_linecount_map,character_order_map,character_textlength_map=process_script(file_path,currentOutputFolder,name,"ALL",enc)

            myprint7("Processed")
            if breakdown==None:
                myprint7("Failed")
                hide_loading()
            else:
                myprint7("OK")
                currentBreakdown=breakdown
                png_output_file=currentOutputFolder+name+"_timeline.png"
                currentTimelinePath=png_output_file
                currentResultCharacterOrderMap=character_order_map
                currentResultEnc=self.encoding
                currentResultName=name
                currentResultLinecountMap=character_linecount_map
                currentResultSceneCharacterMap=scene_characters_map
                postProcess(breakdown,character_order_map,enc,name,character_linecount_map,scene_characters_map,png_output_file)
   




    # def crop_pdf_and_extract_text(self,input_pdf,  output_pdf, crop_box):
    #     cropped_text = []
    #     writer = PdfWriter()
    #     page_idx=self.get_first_page()    

    #     with pdfplumber.open(input_pdf) as pdf:
    #         for i, page in enumerate(pdf.pages[page_idx:], start=page_idx):
    #             cropped_page = page.within_bbox(crop_box)
    #             image = cropped_page.to_image().original
    #             image.save(f"ocr_{i}.png")
    #             myprint7(f"page {i}")

    #             # Perform OCR on the image
    #             #text = pytesseract.image_to_string(image)
    #             cropped_text.append("TODO")
                
    #             # Use PyPDF2 to add the cropped page to the writer
    #             reader = PdfReader(input_pdf)
    #             page_to_add = reader.pages[i]
    #             page_to_add.trimbox.lower_left = (crop_box[0], crop_box[1])
    #             page_to_add.trimbox.upper_right = (crop_box[2], crop_box[3])
    #             writer.add_page(page_to_add)
            
    #         # Save the cropped PDF
    #         with open(output_pdf, 'wb') as f:
    #             writer.write(f)
        
    #     return "\n".join(cropped_text)

    def run(self):
   #     if self.strategy=="OCR":
 #           self.run_crop()
  #      elif self.strategy=="BLOCKS":
            threading.Thread(target=self.run_elements).start()
    
    
    # def run_crop(self):
    #     myprint7("-------------------- RUN CROP PDF -------------")
    #     global currentFilePath
    #     global currentScriptFilename
        # global currentBreakdown
        # global currentOutputFolder
        # global currentTimelinePath
        # global currentResultCharacterOrderMap
        # global currentResultEnc
        # global currentResultName
        # global currentResultLinecountMap
        # global currentResultSceneCharacterMap
        # global importTab
        # global currentPDFPages
        # global currentPDFPageIdx

        # x1=self.left_threshold
        # y1=self.top_threshold
        # x2=self.pdf_page_width-self.right_threshold
        # y2=self.pdf_page_height-self.top_threshold
        
        # cropped_path=self.file_path.lower().replace(".pdf","-cropped.pdf")
        # converted_file_path=self.file_path.lower().replace(".pdf","-cropped.converted.txt")

        # cropbox=[x1,y1,x2,y2]
        # myprint7(f"cropbox {cropbox}")
        # text=self.crop_pdf_and_extract_text(self.file_path,cropped_path,cropbox)
        # myprint7("converted "+str(text))
        # save_string_to_file(text,converted_file_path)
        # myprint7("converted ")
        
        # with open(converted_file_path, 'r', encoding=self.encoding) as file:
        #     file_path=converted_file_path
        #     file_name = os.path.basename(file_path)
        #     currentScriptFilename=file_name
        #     name, extension = os.path.splitext(file_name)
        #     myprint7("Opened")
        #     content = file.read()
        #     myprint7("Read")
        #     file_preview.delete(1.0, tk.END)
        #     file_preview.insert(tk.END, content)
            
        #     enc=self.encoding
        #     myprint7("Process")
        #     breakdown,character_scene_map,scene_characters_map,character_linecount_map,character_order_map,character_textlength_map=process_script(file_path,currentOutputFolder,name,"ALL",enc)

        #     myprint7("Processed")
        #     if breakdown==None:
        #         myprint7("Failed")
        #         hide_loading()
        #     else:
        #         myprint7("OK")
        #         currentBreakdown=breakdown
        #         png_output_file=currentOutputFolder+name+"_timeline.png"
        #         currentTimelinePath=png_output_file
        #         currentResultCharacterOrderMap=character_order_map
        #         currentResultEnc=self.encoding
        #         currentResultName=name
        #         currentResultLinecountMap=character_linecount_map
        #         currentResultSceneCharacterMap=scene_characters_map
        #         postProcess(breakdown,character_order_map,enc,name,character_linecount_map,scene_characters_map,png_output_file)
                
#     def display_page_3_open_preview(self):
#         if self.show_preview:
#             nimg = Image.open("pdfpreview_scaled.png")
            
#         if self.show_preview:    
#             nimg_tk = ImageTk.PhotoImage(nimg)
#             self.current_image = nimg_tk
        
#         myprint7("pdf8d")

#         # Clear the canvas and display the image
#         self.canvas.delete("all")
# #        self.canvas.create_image(0, 0, anchor=tk.NW, image=nimg_tk)

#         self.text_elements=get_pdf_text_elements(self.file_path,self.page_number,self.get_first_page(),self.get_last_page(),self.progress_bar)
        
#         x1=self.left_threshold
#         x2=self.canvas_width-self.right_threshold
#         y1=self.canvas_height-self.top_threshold
#         y2=self.bottom_threshold
#         x2=self.pdf_page_width-self.right_threshold
#         y1=self.pdf_page_height-self.top_threshold

#         myprint7("run split with left={x1} top={y1} right={x2} bottom={y2} ")
#         self.redraw()
       
    def draw_lines(self):
        myprint7(f"Draw {self.left_offset+ self.right_slider.get()*self.scale} {self.left_offset+self.left_slider.get()*self.scale} {self.canvas_height-self.top_slider.get()*self.scale} {self.canvas_height- self.bottom_slider.get()*self.scale}")
        
        v=self.margin_to_positionx2(self.right_threshold)
        self.draw_vertical_liner(v)            
        
        v=self.margin_to_positionx(self.left_threshold)
        self.draw_vertical_line(v)            
        
        
        v=self.margin_to_positionx(self.top_threshold)
        self.draw_horizontal_line(v)   
        

        v=self.margin_to_positiony2(self.bottom_threshold)
        self.draw_horizontal_liner(v)   
    def prev_page(self):
        
        if self.page_number > 0:
            self.page_number -= 1
            self.display_page(self.page_number)
        app.update_idletasks()  # Force the UI to update
        self.label.config(text="Page: "+str(self.page_number+1)+" / "+str(self.num_pages-1))

def merge_with():
    global currentMergedCharacters
    global currentMergedCharactersTo
    global currentMergePopupTable
    
    first_selected_item = currentMergePopupTable.selection()[0]
    first_cell_value = currentMergePopupTable.item(first_selected_item, "values")[0]  
    myprint2("Merge with "+str(first_cell_value))
    
    if first_cell_value in currentMergedCharactersTo:
        first_cell_value=currentMergedCharactersTo[first_cell_value]
        myprint2("Merge with adjust to "+str(first_cell_value))
    else:
        myprint2("Merge with no adjustment"+str(first_cell_value))

    mergeto=first_cell_value
    mergefrom=currentCharacterMergeFromName

    currentMergedCharacters[mergefrom]=mergefrom
    if mergeto in currentMergedCharactersTo:
        currentMergedCharactersTo[mergeto].append(mergefrom)
    else:
        currentMergedCharactersTo[mergeto]=(mergefrom,)
    
    currentMergePopupWindow.destroy()

    reset_tables()
    postProcess(currentBreakdown,currentResultCharacterOrderMap,currentResultEnc,currentResultName,currentResultLinecountMap,currentResultSceneCharacterMap,currentTimelinePath)

class TextPreview:
    def __init__(self, root):
        self.root = root

        self.input_ignore_beginning = tk.StringVar()
        self.input_ignore_end = tk.StringVar()
        self.input_ignore_beginning.set(str(0))
        self.input_ignore_end.set(str(0))
        
        # Create the main frame
        self.main_frame = tk.Frame(self.root)
        self.main_frame.pack(fill=tk.BOTH, expand=True)

        # Create the left panel
        self.left_panel = tk.Frame(self.main_frame)
        self.left_panel.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

        # Add a Text widget to the left panel
        self.text_widget = tk.Text(self.left_panel, wrap=tk.NONE, 
                        yscrollcommand=v_scroll.set, 
                        xscrollcommand=h_scroll.set,relief=tk.FLAT, borderwidth=0)
        self.text_widget.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        # Create the right panel
        self.right_panel = tk.Frame(self.main_frame, width=200)
        self.right_panel.pack(side=tk.RIGHT, fill=tk.Y)

        # Create the label
        label = tk.Label(self.right_panel, text="Lignes à ignorer",         font=bold_font, anchor="w", justify="left", bg=self.right_panel.cget('bg'))
        label.pack(pady=0, fill="x", padx=(10, 0)) 

        label = tk.Label(self.right_panel, text="Début:", anchor="w", justify="left", bg=self.right_panel.cget('bg'))
        label.pack(pady=0, fill="x", padx=(10, 0)) 
        # Create a StringVar to hold the default value

        # Create a single-line text entry widget
        entry_n = ttk.Entry(self.right_panel, width=30, textvariable=self.input_ignore_beginning)
        entry_n.pack(pady=10)
        entry_n.bind('<KeyRelease>', self.update_lines_color)

        label = tk.Label(self.right_panel, text="Fin:", anchor="w", justify="left", bg=self.right_panel.cget('bg'))
        label.pack(pady=0, fill="x", padx=(10, 0)) 

        # Create a single-line text entry widget
        entry_end = ttk.Entry(self.right_panel, width=30, textvariable=self.input_ignore_end)
        entry_end.pack(pady=10)
        entry_end.bind('<KeyRelease>', self.update_lines_color)

        self.buttonframe2 = tk.Frame(self.right_panel)
        self.buttonframe2.pack(side=tk.TOP, fill='x')

        # Create buttons to navigate pages
        self.open_button = tk.Button(self.buttonframe2, text="Lancer le traitement", command=self.run, height=10)
        self.open_button.pack(side=tk.TOP, fill='x', expand=True, padx=10, pady=10)

        if currentConvertedFileIgnoreBeginning>0:
            self.entry_n.set(f"{currentConvertedFileIgnoreBeginning}")
        if currentConvertedFileIgnoreEnd>0:
            self.entry_n.set(f"{currentConvertedFileIgnoreEnd}")
        self.change_lines_color(currentConvertedFileIgnoreBeginning,currentConvertedFileIgnoreEnd)

    def update_lines_color(self, event):
        try:
            n = int(self.input_ignore_beginning.get())
        except ValueError:
            n = 0
        
        try:
            m = int(self.input_ignore_end.get())
        except ValueError:
            m = 0
        
        self.change_lines_color(n, m)
    def change_lines_color(self, n, m):
         # Clear any existing tags
        self.text_widget.tag_remove("grey", "1.0", tk.END)

        # Create and configure a tag for grey text
        self.text_widget.tag_configure("grey", foreground="grey")

        # Get the total number of lines
        total_lines = int(self.text_widget.index('end-1c').split('.')[0])

        # Change the color of the first n lines
        for i in range(n):
            self.text_widget.tag_add("grey", f"{i + 1}.0", f"{i + 1}.0 lineend")

        # Change the color of the last m lines
        for i in range(total_lines - m, total_lines):
            self.text_widget.tag_add("grey", f"{i}.0", f"{i}.0 lineend")
    def run(self):
        myprint7("run Text")
        global currentConvertedFileIgnoreBeginning
        global currentConvertedFileIgnoreEnd
        global currentConvertedFilePath
        currentConvertedFileIgnoreBeginning=int(self.input_ignore_beginning.get())
        currentConvertedFileIgnoreEnd=int(self.input_ignore_end.get())
        runJob(currentConvertedFilePath,"")

def show_info():
    global currentBlockSize
    global currentOutputFolder
    global currentFilePath
    global currentConvertedFilePath
    global currentScriptFilename
    data={
        'currentBlockSize':currentBlockSize,
'currentOutputFolder':currentOutputFolder,
'currentFilePath':currentFilePath,
'currentConvertedFilePath':currentConvertedFilePath,
'currentScriptFilename':currentScriptFilename

    }
    show_info_data(data)
def show_info_data(data):
    # Create a new top-level window
    popup = tk.Toplevel()
    popup.title("Info")
    popup.geometry(f'{900}x{300}+{100}+{100}')

    # Create a frame to hold the table
    frame = ttk.Frame(popup)
    frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

    # Create the treeview
    tree = ttk.Treeview(frame, columns=('Variable', 'Value'), show='headings')
    tree.heading('Variable', text='Variable')
    tree.heading('Value', text='Value')
    tree.pack(fill=tk.BOTH, expand=True)

    # Insert data into the treeview
    for variable, value in data.items():
        tree.insert('', tk.END, values=(variable, value))

    # Add a scrollbar
    scrollbar = ttk.Scrollbar(frame, orient=tk.VERTICAL, command=tree.yview)
    tree.configure(yscroll=scrollbar.set)
    scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
###################################################################################################
## MAIN
myprint7a(f"Make 3")
logging.debug("Checking settings ini file")
myprint7a(f"Make 4")

settings_ini_exists = check_settings_ini_exists()
if settings_ini_exists == False:
    logging.debug("Writing settings ini file")
    write_settings_ini()
    script_folder = get_intial_treeview_folder_path()#os.path.abspath(os.path.dirname(__file__))
    myprint7a(f"Initial treeview folder          :  {script_folder}")
    update_ini_settings_file("SCRIPT_FOLDER",script_folder)
myprint7a(f"Make 5")

settings = read_settings_ini()
myprint7a(f"Make 6")




try:
    myprint7a(f"Make 6a")
    app = tk.Tk(className="Scripti")
    myprint7a(f"Make 6b")
    app.title('Scripti')

    myprint7a(f"Make 7")


    app_dir = os.path.dirname(os.path.abspath(__file__))
    icons_dir =os.path.join(app_dir,"icons/")
    if "Contents/MacOS" in app_dir:
        parent_dir = os.path.dirname(app_dir)
        examples_dir= icons_dir =os.path.join(parent_dir,"Resources/examples/")
        icons_dir =os.path.join(parent_dir,"Resources/icons/")
        copy_folder_contents(examples_dir,get_intial_treeview_folder_path())
    else:
        examples_dir= os.path.join(app_dir,"examples/")
        icons_dir =os.path.join(app_dir,"icons/")
        copy_folder_contents(examples_dir,get_intial_treeview_folder_path())
         
    myprint7a("Icon dir         :"+icons_dir)
    #icons_dir =app_dir+"/icons/"
    myprint7a("App dir           :"+app_dir)
    app.iconbitmap(icons_dir+'app_icon.ico') 

    # Ensure the app name appears in the macOS menu bar
    if os.name == 'posix':  # This check is for macOS
        app.tk.call('wm', 'iconname', app._w, 'Scripti')
        app.tk.call('wm', 'iconphoto', app._w, '-default', tk.PhotoImage(file=icons_dir + 'app_icon.png'))

    # Use iconphoto for cross-platform icon setting
    if os.name == 'nt':  # This check is for Windows
        app.iconbitmap(icons_dir + 'app_icon.ico')
    else:
        icon_path = icons_dir + 'app_icon.png'
        app.iconphoto(True, tk.PhotoImage(file=icon_path))


    logging.debug("Creating app")

    #app.bind('<Configure>', on_resize)

    # Menu bar
    menu_bar = Menu(app)
    app.config(menu=menu_bar)

    folder_icon = tk.PhotoImage(file=icons_dir+"folder_icon.png")  # Adjust path to your icon file
    import_icon = tk.PhotoImage(file=icons_dir+"import.png")  # Adjust path to your icon file
    txt_icon = tk.PhotoImage(file=icons_dir+"txt_icon.png")  # Adjust path to your icon file
    docx_icon = tk.PhotoImage(file=icons_dir+"docx_icon.png")  # Adjust path to your icon file
    original_icon = tk.PhotoImage(file=icons_dir+"textd_icon.png")  # Adjust path to your icon file
    char_icon = tk.PhotoImage(file=icons_dir+"character_icon.png")  # Adjust path to your icon file
    timeline_icon = tk.PhotoImage(file=icons_dir+"timeline_icon.png")  # Adjust path to your icon file
    scene_icon = tk.PhotoImage(file=icons_dir+"scenes_icon.png")  # Adjust path to your icon file
    export_icon = tk.PhotoImage(file=icons_dir+"export_icon.png")  # Adjust path to your icon file
    chat_icon = tk.PhotoImage(file=icons_dir+"chat_icon.png")  # Adjust path to your icon file


    # File menu
    file_menu = Menu(menu_bar, tearoff=0)
    menu_bar.add_cascade(label="Fichier", menu=file_menu)
    file_menu.add_command(label="Ouvrir un dossier de travail...", command=open_folder)
    file_menu.add_command(label="Ouvrir un fichier de script...", command=open_script)
    file_menu.add_command(label="Ouvrir un groupe de fichiers de script...", command=open_script_group)
    #file_menu.add_command(label="Export csv...", command=export_csv)
    file_menu.add_separator()
    recent_files_menu = tk.Menu(file_menu, tearoff=0)
    file_menu.add_cascade(label="Fichiers récents", menu=recent_files_menu)
    update_recent_files_menu()
    file_menu.add_separator()

    file_menu.add_command(label="Quitter", command=exit_app)

    input_blocksize = tk.StringVar()
    input_blocksize.set(str(currentBlockSize))



    importTab = None



    character_menu = Menu(menu_bar, tearoff=0)
    menu_bar.add_cascade(label="Personnages", menu=character_menu)
    #settings_menu.add_command(label="Changer la methode de comptage counting method...", command=show_popup_counting_method)
    character_menu.add_command(label="Fusionner sélection", command=merge_characters2)
    character_menu.add_separator()
    #settings_menu.add_command(label="Set block length...", command=open_folder)
    character_menu.add_command(label="Désactiver sélection", command=disable_character)
    character_menu.add_command(label="Activer sélection", command=enable_character)
    character_menu.add_command(label="Activer tout", command=restore_characters)
    character_menu.add_separator()
    character_menu.add_command(label="Déselectionner tout", command=deselect_characters)


    settings_menu = Menu(menu_bar, tearoff=0)
    menu_bar.add_cascade(label="Paramètres", menu=settings_menu)
    #settings_menu.add_command(label="Changer la methode de comptage counting method...", command=show_popup_counting_method)
    settings_menu.add_command(label="Change la taille des répliques...", command=show_popup_line_size)
    settings_menu.add_command(label="Info...", command=show_info)


    help_menu = Menu(menu_bar, tearoff=0)
    menu_bar.add_cascade(label="Aide", menu=help_menu)
    #settings_menu.add_command(label="Changer la methode de comptage counting method...", command=show_popup_counting_method)
    help_menu.add_command(label="Word > Tableaux...", command=help_word_table)
    help_menu.add_command(label="PDF > Texte...", command=help_pdf_text)
    help_menu.add_command(label="Fusionner personnages...", command=help_merge)

    loading_label = ttk.Frame(app)
    #loading_label.pack(side=tk.TOP, fill=tk.X,expand=True)
    # Load folder button
    #load_button = ttk.Button(loading_label, text="Hide ", command=hide_loading)
    #load_button.pack(side=tk.TOP, fill=tk.X,padx=20,pady=20)

    loading_label_txt = ttk.Label(loading_label, text="Analyze", font=('Arial', 12),padding="20 100 20 20")
    loading_label_txt.pack(side=tk.TOP, fill=tk.X,expand=True)


    # Create a PanedWindow widget
    paned_window = ttk.PanedWindow(app, orient=tk.HORIZONTAL)
    paned_window.pack(fill=tk.BOTH, expand=True,padx=0,pady=0)


    # Create two frames for the left and right panels
    left_frame = ttk.Frame(paned_window, width=200, height=400, relief=tk.FLAT)
    right_frame = ttk.Frame(paned_window, width=400, height=400, relief=tk.FLAT)

    # Add frames to the PanedWindow
    paned_window.add(left_frame, weight=1)  # The weight determines how additional space is distributed
    paned_window.add(right_frame, weight=2)

    menu = tk.Menu(app, tearoff=0)
    menu.add_command(label="Open File", command=open_file_in_system)

    char_menu = tk.Menu(app, tearoff=0)
    char_menu.add_command(label="Merge with...", command=merge_characters)
    char_menu.add_command(label="Hide", command=hide_character)

    #######################################################################################
    # Folder tree
    folders = ttk.Treeview(left_frame, columns=("Path","Extension",))
    folders.heading("#0", text="Fichier")
    folders.heading("Extension", text="Type")
    folders.heading("Path", text="Path")
    folders.column("#0", width=240)  # Adjust as needed
    folders.column("Path", width=0, stretch=tk.NO)
    folders.column("Extension", width=60, stretch=tk.NO)

    folders.tag_configure('not_supported', foreground='#cccccc')
    folders.tag_configure('supported', foreground='#444444')
    #folders.tag_configure('folder', foreground='#6666cc')
    bold_font = tkFont.Font( weight="bold")
    folders.tag_configure('bold', font=bold_font)

    # Default tag with normal background
    folders.tag_configure('normal', background='white')

    folders.tag_configure('hover', background='#f4f4f4')
    style = ttk.Style()
    #style.configure('TNotebook.Tab', padding=[10,10,10,10])  # Adjust these values as needed

    style.configure("Treeview", rowheight=30)  # Increase the row height
    style.configure("Treeview.Item", padding=(3, 4, 3, 4))  # Top and bottom padding
    #bold_font = ('Arial', 10, 'bold')
    #style.configure("Treeview", font=bold_font)

    
    folders.pack(side=tk.TOP, fill=tk.BOTH, expand=True)
    folders.bind('<<TreeviewSelect>>', on_treeview_folder_select)
    folders.bind("<<TreeviewOpen>>", on_folder_open)
    # Bind motion event
    folders.bind('<Motion>', on_motion)
    folders.bind('<Leave>', on_leave)
    folders.bind('<Button-1>', toggle_folder)


        
    folders.bind('<Button-3>', on_right_click)  # Right click on Windows/Linux
    folders.bind('<Button-2>', on_right_click) 

    # Notebook (tabbed interface)
    notebook = ttk.Notebook(right_frame)
    notebook.pack(fill=tk.BOTH, expand=True, padx=0, pady=0)

    # Configure the style of the tab
    fontsize=14
    if platform.system() == 'Windows':
        style.configure('TNotebook.Tab', background='#f0f0f0', padding=(5, 3), font=('Helvetica', fontsize))
    #style.configure('TNotebook.Tab', background='#f0f0f0', padding=(5, 3), font=('Helvetica', fontsize))
    # Configure the tab area (optional, for better Windows look)
    style.configure('TNotebook', tabposition='nw', background='#f0f0f0')
    style.configure('TNotebook', padding=0)  # Removes padding around the tab area

    #def on_tab_selected(event):
        #myprint2("Tab selected:", event.widget.select())


    #notebook.bind("<<NotebookTabChanged>>", on_tab_selected)
    # File preview tab
    tab_import = ttk.Frame(notebook)

    notebook.add(tab_import, text='Tables Word',image=import_icon, compound=tk.LEFT)

    tab_import_pdf = ttk.Frame(notebook)
    #pdf_viewer = PDFViewer(tab_import_pdf, file_path)
    # Set the initial page number
    #current_page = 10
    #go_to_page(current_page)

    notebook.add(tab_import_pdf, text='Extracteur de dialogue PDF',image=import_icon, compound=tk.LEFT)


    # Create a frame for the 'Texte' tab
    tab_text = ttk.Frame(notebook)
    notebook.add(tab_text, text='Texte')

    # Create a Text widget with vertical and horizontal scrollbars
    text_frame = ttk.Frame(tab_text, borderwidth=0,relief=tk.FLAT,)
    text_frame.pack(fill=tk.BOTH, expand=True)

    # Create vertical scrollbar
    v_scroll = ttk.Scrollbar(text_frame, orient=tk.VERTICAL)
    v_scroll.pack(side=tk.RIGHT, fill=tk.Y)

    # Create horizontal scrollbar
    h_scroll = ttk.Scrollbar(text_frame, orient=tk.HORIZONTAL)
    h_scroll.pack(side=tk.BOTTOM, fill=tk.X)

    # Create Text widget

    file_preview=TextPreview(text_frame)
#    file_preview = tk.Text(text_frame, wrap=tk.NONE, 
 #                       yscrollcommand=v_scroll.set, 
  #                      xscrollcommand=h_scroll.set,relief=tk.FLAT, borderwidth=0)
#    file_preview.pack(fill=tk.BOTH, expand=True)

    # Configure scrollbars to work with the Text widget
 #   v_scroll.config(command=file_preview.yview)
  #  h_scroll.config(command=file_preview.xview)

    # Statistics tab
    tab_characters = ttk.Frame(notebook)

    def merge_together():
        myprint7("merge")

    #btn_merge = ttk.Button(tab_characters, text="Fusionner ...", command=merge_together)
    #btn_merge.pack(side=tk.TOP, fill=tk.X,padx=20,pady=20)

    # Create a Treeview widget within the stats_frame for the table


    character_table = ttk.Treeview(tab_characters, columns=('Order', 'Character','Status', 'Characters','Lines'), show='headings', selectmode='none')
    #character_table = ttk.Treeview(character_tab, columns=('Order', 'Character', 'Lines','Characters','Words','Blocks (50)','Scenes'), show='headings')
    # Define the column headings
    character_table.heading('Order', text='Order')
    character_table.heading('Character', text='Personnage')
    character_table.heading('Status', text='Status')
    #character_table.heading('Lines', text='Lines')
    character_table.heading('Characters', text='Caractères')
    #character_table.heading('Words', text='Words')
    character_table.heading('Lines', text='Lignes')
    #character_table.heading('Blocks (40)', text='Blocks (40)')
    #character_table.heading('Scenes', text='Scènes')

    # Define the column width and alignment
    character_table.column('Order', width=25, anchor='center')
    character_table.column('Character', width=200, anchor='w')
    character_table.column('Status', width=50, anchor='w')
    #character_table.column('Lines', width=50, anchor='w')
    character_table.column('Characters', width=50, anchor='w')
    #character_table.column('Words', width=50, anchor='w')
    character_table.column('Lines', width=50, anchor='w')
    #character_table.column('Scenes', width=50, anchor='w')

    # Pack the Treeview widget with enough space
    character_table.pack(fill='both', expand=True)
    notebook.add(tab_characters, text='Personnages',image=char_icon, compound=tk.LEFT)

    character_table.tag_configure('hidden', foreground='#999999')
    character_table.tag_configure('grouped', foreground='#000000',background='#995555')
    character_table.bind('<Button-3>', on_character_table_click)  # Right click on Windows/Linux
    character_table.bind('<Button-2>', on_character_table_click) 
    character_table.bind('<Button-1>', on_character_table_click) 












    breakdown_tab = ttk.Frame(notebook)
    # Create a Treeview widget within the stats_frame for the table
    breakdown_table = ttk.Treeview(breakdown_tab, columns=('Line', 'Type', 'Character','Text'), show='headings')
    # Define the column headings
    breakdown_table.heading('Line', text='Ligne')
    breakdown_table.heading('Type', text='Type')
    breakdown_table.heading('Character', text='Personnage')
    breakdown_table.heading('Text', text='Dialogue')

    # Define the column width and alignment
    breakdown_table.column('Line', width=25, anchor='w')
    breakdown_table.column('Type', width=25, anchor='w')
    breakdown_table.column('Character', width=50, anchor='w')
    breakdown_table.column('Text', width=200, anchor='w')
    # Pack the Treeview widget with enough space
    breakdown_table.pack(fill='both', expand=True)
    # Configure the tag to change the background color
    breakdown_table.tag_configure('nonspeech', background='#fafafa')
    breakdown_table.tag_configure('scene', background='#fffec8')
    bold_font = tkFont.Font( weight="bold")
    breakdown_table.tag_configure('border', background='#444444')  # A lighter shade to simulate space
    breakdown_table.tag_configure('bold', font=bold_font)
    #notebook.add(breakdown_tab, text='Scenes',image=scene_icon, compound=tk.LEFT)
            











    tab_dialog = ttk.Frame(notebook)
    # Create a Treeview widget within the stats_frame for the table
    stats_table = ttk.Treeview(tab_dialog, columns=('Line number',  'Character','Text','Characters'), show='headings')
    # Define the column headings
    stats_table.heading('Line number', text='Ligne')
    stats_table.heading('Character', text='Personnage')
    stats_table.heading('Text', text='Dialogue')
    stats_table.heading('Characters', text='Caractères')

    # Define the column width and alignment
    stats_table.column('Line number', width=25, anchor='center')
    stats_table.column('Character', width=100, anchor='w')
    stats_table.column('Text', width=200, anchor='w')
    stats_table.column('Characters', width=50, anchor='w')

    # Pack the Treeview widget with enough space
    stats_table.pack(fill='both', expand=True)
    # Configure the tag to change the background color
    stats_table.tag_configure('nonspeech', background='#fafafa')
    stats_table.tag_configure('scene', background='#fffec8')
    bold_font = tkFont.Font( weight="bold")
    stats_table.tag_configure('border', background='#444444')  # A lighter shade to simulate space
    stats_table.tag_configure('bold', font=bold_font)

    notebook.add(tab_dialog, text="Dialogue dans l'ordre",image=chat_icon, compound=tk.LEFT)


    # Statistics tab
    tab_dialog_by_character = ttk.Frame(notebook)

    # Create a Treeview widget within the stats_frame for the table
    cols=('Line #', 'Character','Character (raw)','Line')
    for i in countingMethods:
        cols= cols+(countingMethodNames[i],)

    style.configure('CleftPanel.TFrame', background='#fafafa')
    style.configure('CrightPanel.TFrame', background='#fafafa')


    # Create left and right frames (panels) inside the tab
    cleft_panel = ttk.Frame(tab_dialog_by_character, borderwidth=0, relief="flat", width=200)
    cright_panel = ttk.Frame(tab_dialog_by_character, borderwidth=0, relief="flat")

    cleft_panel.configure(style='CleftPanel.TFrame')  # Apply the styled background
    cright_panel.configure(style='CrightPanel.TFrame')  # Apply the styled background

    # Pack the frames into the tab
    #cleft_panel.pack(side='left', fill='y', padx=(0, 20))
    #cright_panel.pack(side='right', fill='y', expand=True)

    # Configure column weights to make right panel flexible
    tab_dialog_by_character.grid_columnconfigure(1, weight=1)
    tab_dialog_by_character.grid_rowconfigure(0, weight=1)


    character_list_table = ttk.Treeview(cleft_panel, columns=('Character'), show='headings')
    # Define the column headings
    character_list_table.heading('Character', text='Personnage')
    # Define the column width and alignment
    character_list_table.column('Character', width=50, anchor='w')
    # Pack the Treeview widget with enough space
    character_list_table.pack(fill='both', expand=True)

    # Grid frames with padding
    cleft_panel.grid(row=0, column=0, sticky='nsew', padx=(0, 10))  # Add padding on the right side of left panel
    cright_panel.grid(row=0, column=1, sticky='nsew')  # Automatically spaced by the left panel's padding

    # Configure column weights to make right panel flexible
    tab_dialog_by_character.grid_columnconfigure(0, weight=0, minsize=200)  # Set minimum size for the left panel
    tab_dialog_by_character.grid_columnconfigure(1, weight=1) 
    tab_dialog_by_character.grid_rowconfigure(0, weight=1)

    character_list_table.bind('<ButtonRelease-1>', on_item_selected)



    character_stats_table = ttk.Treeview(cright_panel, columns=cols, show='headings')
    # Define the column headings
    character_stats_table.heading('Line #', text='Ligne #')
    character_stats_table.heading('Character', text='Personnage')
    character_stats_table.heading('Character (raw)', text='Personnage (brut)')
    character_stats_table.heading('Line', text='Réplique')
    for i in countingMethods:
        character_stats_table.heading(countingMethodNames[i], text=countingMethodNames[i])


    # Define the column width and alignment
    character_stats_table.column('Line #', width=25, anchor='center')
    character_stats_table.column('Character',  anchor='w', width=0, stretch=tk.NO)
    character_stats_table.column('Character (raw)',anchor='w', width=0, stretch=tk.NO)
    character_stats_table.column('Line', width=100, anchor='w')
    for i in countingMethods:
        character_stats_table.column(countingMethodNames[i], width=25, anchor='w')

    # Pack the Treeview widget with enough space
    character_stats_table.pack(fill='both', expand=True)
    character_stats_table.tag_configure('total', background='#444444',foreground="#ffffff")

    notebook.add(tab_dialog_by_character, text='Répliques par personnage',image=chat_icon, compound=tk.LEFT)

    ################################################################################
    tab_export = ttk.Frame(notebook)
    # Load folder button
    load_button = ttk.Button(tab_export, text="Ouvrir le dossier de résultats...", command=open_result_folder)
    load_button.pack(side=tk.TOP, fill=tk.X,padx=20,pady=20)
    load_button = ttk.Button(tab_export, text="Ouvrir le fichier de conversion ...", command=open_conversion_file)
    load_button.pack(side=tk.TOP, fill=tk.X,padx=20,pady=20)
    load_button = ttk.Button(tab_export, text="Test ...", command=hide_importtable_tab)
    load_button.pack(side=tk.TOP, fill=tk.X,padx=20,pady=20)

    #load_buttonb = ttk.Button(export_tab, text="Show loading ", command=show_loading)
    #load_buttonb.pack(side=tk.TOP, fill=tk.X,padx=20,pady=20)

    # Load folder button
    load_button = ttk.Button(tab_export, text="Ouvrir le comptage .xlsx...", command=open_xlsx_recap)
    load_button.pack(side=tk.TOP, fill=tk.X,padx=20,pady=20)

    load_button = ttk.Button(tab_export, text="Ouvrir le détail du dialogue...", command=open_dialog_recap)
    load_button.pack(side=tk.TOP, fill=tk.X,padx=20,pady=20)







    ################################################################################



    notebook.add(tab_export, text='Export',image=export_icon, compound=tk.LEFT)


    logging.debug("Launch")

    currentScriptFolder=settings['SCRIPT_FOLDER']
    load_tree("",currentScriptFolder)


    center_window()  # Center the window
    app.title('Scripti')
    app.wm_title = " Your title name "
    app.mainloop()
except Exception as e:
    print(f"Error initializing Tkinter: {e}")
    traceback.print_exc()